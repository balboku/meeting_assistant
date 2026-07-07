"""
Supplementary evidence analysis for existing meeting notes.

This module keeps the first version deliberately file-based: uploaded evidence is
stored next to generated output, analyzed by Gemini, and appended to the meeting
Markdown without changing the database schema.
"""

from __future__ import annotations

import mimetypes
import os
import re
import shutil
import time
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from google import genai
from google.genai import types

from backend.database import get_meeting
from backend.tasks import GEMINI_MODEL, MAX_UPLOAD_WAIT_SECONDS, POLLING_INTERVAL


ROOT_DIR = Path(__file__).parent.parent
load_dotenv(dotenv_path=ROOT_DIR / ".env")

ATTACHMENT_DIR = Path(os.getenv("MEETING_ATTACHMENT_DIR") or ROOT_DIR / "output" / "attachments")
SUPPORTED_EVIDENCE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".pdf",
    ".txt",
    ".md",
    ".csv",
    ".docx",
}
TEXT_EVIDENCE_EXTENSIONS = {".txt", ".md", ".csv"}
EVIDENCE_SECTION_HEADING = "## 📎 五、補充資料與佐證 (Supplementary Evidence)"
MAX_CONTEXT_CHARS = 24000
MAX_TEXT_EVIDENCE_CHARS = 20000


def _safe_filename(filename: str) -> str:
    name = Path(filename or "evidence").name.strip()
    if not name:
        name = "evidence"
    stem = Path(name).stem or "evidence"
    suffix = Path(name).suffix.lower()
    safe_stem = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", stem, flags=re.UNICODE).strip("._")
    if not safe_stem:
        safe_stem = "evidence"
    return f"{safe_stem[:80]}{suffix}"


def _unique_attachment_path(meeting_id: int, filename: str) -> Path:
    directory = ATTACHMENT_DIR / f"meeting_{meeting_id}"
    directory.mkdir(parents=True, exist_ok=True)

    safe_name = _safe_filename(filename)
    candidate = directory / safe_name
    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix
    for index in range(2, 1000):
        numbered = directory / f"{stem}_{index}{suffix}"
        if not numbered.exists():
            return numbered

    raise RuntimeError("補充資料檔名重複過多，請更換檔名後再試。")


def _strip_markdown_for_context(markdown: str) -> str:
    content = markdown.strip()
    transcript_marker = re.search(r"\n##\s*📝\s*四、", content)
    if transcript_marker:
        content = content[:transcript_marker.start()].strip()
    if len(content) > MAX_CONTEXT_CHARS:
        content = content[:MAX_CONTEXT_CHARS] + "\n\n[系統截斷：會議內容過長，僅提供前段摘要與決議作為比對依據]"
    return content


def _read_text_evidence(path: Path, suffix: str) -> Optional[str]:
    if suffix in TEXT_EVIDENCE_EXTENSIONS:
        return path.read_text(encoding="utf-8", errors="replace")[:MAX_TEXT_EVIDENCE_CHARS]

    if suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        text = "\n".join(paragraphs + table_rows)
        return text[:MAX_TEXT_EVIDENCE_CHARS]

    return None


def _build_evidence_prompt(
    meeting_record: dict,
    original_filename: str,
    attachment_path: Path,
    note: Optional[str],
    extracted_text: Optional[str],
) -> str:
    meeting_context = _strip_markdown_for_context(meeting_record.get("full_content", ""))
    note_text = note.strip() if note else "（無）"
    extracted_block = ""
    if extracted_text is not None:
        extracted_block = f"""

【補充資料文字內容】
{extracted_text}
""".rstrip()

    return f"""
你是會議紀錄助理，正在替既有會議紀錄補上「補充資料與佐證」。
請檢視補充資料，判斷它與會議內容的關聯性，並只輸出 Markdown。

【會議資訊】
- 會議 ID：{meeting_record.get("id")}
- 標題：{meeting_record.get("title")}
- 日期：{meeting_record.get("date")}
- 原始音檔：{meeting_record.get("source_audio")}

【既有會議紀錄（摘要/決議/待辦優先，逐字稿可能已省略）】
{meeting_context}

【補充資料】
- 原始檔名：{original_filename}
- 系統保存路徑：{attachment_path}
- 使用者備註：{note_text}{extracted_block}

【輸出要求】
請用繁體中文輸出以下 Markdown，並保持可直接貼入會議紀錄：
### 資料：{original_filename}
- 系統判斷：一句話說明關聯性（高度相關 / 部分相關 / 低相關 / 無法判斷）與理由。
- 擷取重點：列出補充資料中可佐證或補充會議紀錄的具體事實。
- 對會議記錄的影響：指出應補到摘要、決議或待辦事項的內容；沒有就寫「未發現需要更新」。
- 可能矛盾或待確認：若補充資料與逐字稿/摘要不同，清楚標示；沒有就寫「未發現明顯矛盾」。
- 來源註記：寫出原始檔名與系統保存路徑。

務必區分「逐字稿提到」、「補充資料顯示」、「系統推論」、「需人工確認」。
不要杜撰補充資料中不存在的數字、人名、日期或承諾。
""".strip()


def _wait_for_uploaded_file(client, uploaded):
    elapsed = 0
    while not uploaded.state or uploaded.state.name == "PROCESSING":
        if elapsed >= MAX_UPLOAD_WAIT_SECONDS:
            raise RuntimeError("補充資料上傳至 AI 後處理逾時")
        time.sleep(POLLING_INTERVAL)
        elapsed += POLLING_INTERVAL
        uploaded = client.files.get(name=uploaded.name)

    if uploaded.state.name == "FAILED":
        raise RuntimeError("補充資料上傳至 AI 後處理失敗")
    return uploaded


def generate_evidence_markdown(
    meeting_record: dict,
    attachment_path: Path,
    original_filename: str,
    note: Optional[str],
    model: Optional[str] = None,
) -> str:
    """Ask Gemini to inspect the uploaded evidence and return appendable Markdown."""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("未設定 GEMINI_API_KEY，無法分析補充資料。")

    suffix = attachment_path.suffix.lower()
    extracted_text = _read_text_evidence(attachment_path, suffix)
    prompt = _build_evidence_prompt(
        meeting_record=meeting_record,
        original_filename=original_filename,
        attachment_path=attachment_path,
        note=note,
        extracted_text=extracted_text,
    )

    client = genai.Client(api_key=api_key)
    uploaded = None
    try:
        contents = [prompt]
        if extracted_text is None:
            mime_type = mimetypes.guess_type(attachment_path.name)[0] or "application/octet-stream"
            uploaded = client.files.upload(
                file=str(attachment_path),
                config=types.UploadFileConfig(
                    display_name=original_filename,
                    mime_type=mime_type,
                ),
            )
            uploaded = _wait_for_uploaded_file(client, uploaded)
            contents = [uploaded, prompt]

        response = client.models.generate_content(
            model=model or GEMINI_MODEL,
            contents=contents,
            config=types.GenerateContentConfig(
                temperature=0.1,
                top_p=0.9,
                max_output_tokens=8192,
            ),
        )
        markdown = (response.text or "").strip()
        if not markdown:
            raise RuntimeError("AI 未回傳補充資料分析內容。")
        return markdown
    finally:
        if uploaded is not None:
            try:
                client.files.delete(name=uploaded.name)
            except Exception:
                pass


def _append_evidence(existing_markdown: str, evidence_markdown: str) -> str:
    cleaned_existing = existing_markdown.rstrip()
    cleaned_evidence = evidence_markdown.strip()
    if not cleaned_existing:
        return f"{EVIDENCE_SECTION_HEADING}\n\n{cleaned_evidence}\n"

    if EVIDENCE_SECTION_HEADING in cleaned_existing:
        return f"{cleaned_existing}\n\n{cleaned_evidence}\n"

    return f"{cleaned_existing}\n\n{EVIDENCE_SECTION_HEADING}\n\n{cleaned_evidence}\n"


def analyze_and_append_evidence(
    meeting_id: int,
    source_path: Path,
    original_filename: str,
    note: Optional[str] = None,
    model: Optional[str] = None,
) -> dict:
    """Copy an uploaded evidence file, analyze it, and append the result to Markdown."""
    source = Path(source_path)
    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_EVIDENCE_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EVIDENCE_EXTENSIONS))
        raise ValueError(f"不支援的補充資料格式：{suffix}。支援格式：{supported}")
    if not source.is_file():
        raise FileNotFoundError(f"找不到補充資料檔案：{source}")

    meeting_record = get_meeting(meeting_id)
    if not meeting_record:
        raise KeyError(f"找不到會議記錄：ID={meeting_id}")

    output_path = Path(meeting_record["output_path"])
    if not output_path.is_file():
        raise FileNotFoundError(f"找不到會議 Markdown 檔案：{output_path}")

    attachment_path = _unique_attachment_path(meeting_id, original_filename)
    shutil.copy2(source, attachment_path)

    evidence_markdown = generate_evidence_markdown(
        meeting_record=meeting_record,
        attachment_path=attachment_path,
        original_filename=original_filename,
        note=note,
        model=model,
    )

    current_markdown = output_path.read_text(encoding="utf-8")
    updated_markdown = _append_evidence(current_markdown, evidence_markdown)
    output_path.write_text(updated_markdown, encoding="utf-8")

    return {
        "status": "success",
        "meeting_id": meeting_id,
        "file_name": attachment_path.name,
        "attachment_path": str(attachment_path),
        "evidence_markdown": evidence_markdown.strip(),
        "full_content": updated_markdown,
    }
