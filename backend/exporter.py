from pathlib import Path
import logging
import os
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

logger = logging.getLogger("MeetingAssistant.Exporter")

TEMPLATE_PATH = Path(os.getenv("MEETING_DOCX_TEMPLATE_PATH", "4-QA-005 V01 會議紀錄.docx"))
ACTION_TABLE_HEADERS = {"#", "任務描述", "負責人", "期限", "優先級"}
ACTION_LINKED_TABLE_HEADERS = {"#", "關聯討論", "關聯決議", "任務描述", "負責人", "期限", "優先級"}
ACTION_TABLE_WIDTHS_DXA = [550, 5450, 1550, 1450, 850]
ACTION_LINKED_TABLE_WIDTHS_DXA = [500, 850, 850, 4000, 1250, 1300, 750]
DEFAULT_TABLE_WIDTH_DXA = sum(ACTION_TABLE_WIDTHS_DXA)
BODY_FONT = "Microsoft JhengHei"


def _template_path() -> Path:
    return Path(os.getenv("MEETING_DOCX_TEMPLATE_PATH", str(TEMPLATE_PATH)))


def export_meeting_to_docx(meeting_record: dict, output_filepath: str) -> bool:
    """
    將會議記錄內容填入 4-QA-005 V01 會議紀錄.docx 範本，並儲存到 output_filepath。
    保留原本的格式，僅在特定的儲存格加入文字。
    """
    template_path = _template_path()
    if not template_path.exists():
        logger.error(f"找不到範本檔案：{template_path}")
        return False

    try:
        doc = Document(template_path)

        if not doc.tables:
            logger.error("範本中沒有找到表格")
            return False

        table = doc.tables[0]

        # 填寫日期
        if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
            date_str = meeting_record.get("date", "").split(" ")[0]
            _set_cell_text(table.cell(0, 1), date_str)

        # 紀錄者
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            _set_cell_text(table.cell(2, 1), "AI 語音會議助理")

        # 會議主題
        if len(table.rows) > 3 and len(table.rows[3].cells) > 1:
            _set_cell_text(table.cell(3, 1), meeting_record.get("title", ""))

        # 討論內容 (Row 5, Cell 0)
        if len(table.rows) > 5 and len(table.rows[5].cells) > 0:
            content = meeting_record.get("full_content", "")
            _set_cell_markdown(table.cell(5, 0), content)

        doc.save(output_filepath)
        logger.info(f"成功匯出 Word 文件：{output_filepath}")
        return True

    except Exception as e:
        logger.error(f"匯出 Word 失敗：{e}")
        return False

def _set_cell_text(cell, text: str):
    """保留原本儲存格的樣式，將文字填入"""
    # 取得第一個段落
    if cell.paragraphs:
        p = cell.paragraphs[0]
        # 清除現有所有的 runs
        for run in p.runs:
            run.text = ""
        # 新增文字，這樣預設會繼承該段落的樣式
        p.add_run(text)
    else:
        cell.text = text


def _set_cell_markdown(cell, markdown: str) -> None:
    """將會議 Markdown 轉成較可讀的 Word 內容，特別處理待辦事項表格。"""
    _clear_cell_content(cell)
    lines = _strip_frontmatter(markdown).splitlines()
    first_paragraph_used = False
    in_transcript = False
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue
        if _is_horizontal_rule(line):
            index += 1
            continue

        table_rows, next_index = _consume_markdown_table(lines, index)
        if table_rows:
            _add_spacing_paragraph(cell, after=2)
            _add_word_table(cell, table_rows)
            _add_spacing_paragraph(cell, after=5)
            first_paragraph_used = True
            index = next_index
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            heading = heading_match.group(2).strip()
            in_transcript = "完整逐字稿" in heading or "Verbatim Transcript" in heading
            paragraph = _next_paragraph(cell, first_paragraph_used)
            first_paragraph_used = True
            _append_markdown_runs(
                paragraph,
                _clean_heading_text(heading),
                size=11,
                bold=True,
                color=RGBColor(31, 78, 121),
            )
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            index += 1
            continue

        role = "note" if line.startswith(">") else "transcript" if in_transcript else "body"
        text = _clean_body_line(line)
        if text:
            paragraph = _next_paragraph(cell, first_paragraph_used)
            first_paragraph_used = True
            _format_paragraph(paragraph, role)
            _append_markdown_runs(
                paragraph,
                text,
                size=9 if role == "transcript" else 10,
                italic=(role == "note"),
                color=RGBColor(89, 89, 89) if role == "note" else None,
            )
        index += 1


def _clear_cell_content(cell) -> None:
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is tc_pr:
            continue
        tc.remove(child)
    cell.add_paragraph()


def _strip_frontmatter(markdown: str) -> str:
    text = markdown or ""
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            return text[end + 4 :].lstrip()
    return text


def _is_horizontal_rule(line: str) -> bool:
    return bool(re.fullmatch(r"-{3,}", line.strip()))


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip().replace("<br>", "\n").replace("<br/>", "\n") for cell in stripped.split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _consume_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]] | None, int]:
    if start_index + 1 >= len(lines) or not _is_table_line(lines[start_index]):
        return None, start_index

    header = _split_table_row(lines[start_index])
    separator = _split_table_row(lines[start_index + 1])
    if not _is_separator_row(separator):
        return None, start_index

    rows = [header]
    index = start_index + 2
    while index < len(lines) and _is_table_line(lines[index]):
        row = _split_table_row(lines[index])
        if len(row) < len(header):
            row.extend([""] * (len(header) - len(row)))
        rows.append(row[: len(header)])
        index += 1

    return rows, index


def _next_paragraph(cell, first_paragraph_used: bool):
    if not first_paragraph_used and cell.paragraphs:
        return cell.paragraphs[0]
    return cell.add_paragraph()


def _add_spacing_paragraph(cell, after: int) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)


def _clean_heading_text(text: str) -> str:
    return re.sub(r"^[^\w\u4e00-\u9fff]+", "", text).strip()


def _clean_body_line(line: str) -> str:
    text = line.strip()
    text = re.sub(r"^>\s*", "", text)
    text = re.sub(r"^(\d+\.)\s{2,}", r"\1 ", text)
    return text


def _format_paragraph(paragraph, role: str) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(3 if role != "transcript" else 2)
    fmt.line_spacing = 1.1 if role != "transcript" else 1.05
    if role == "note":
        fmt.left_indent = Inches(0.15)
    elif re.match(r"^\d+\.\s", paragraph.text or ""):
        fmt.left_indent = Inches(0.15)


def _append_markdown_runs(
    paragraph,
    text: str,
    *,
    size: int | float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    _clear_paragraph_runs(paragraph)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        part_is_bold = part.startswith("**") and part.endswith("**")
        cleaned = part[2:-2] if part_is_bold else part.replace("**", "")
        if not cleaned:
            continue
        run = paragraph.add_run(cleaned)
        run.bold = bold or part_is_bold
        run.italic = italic
        _set_run_font(run, size=size, color=color)


def _clear_paragraph_runs(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:r"):
            p.remove(child)


def _set_run_font(run, *, size: int | float, color: RGBColor | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), BODY_FONT)


def _add_word_table(cell, rows: list[list[str]]) -> None:
    if not rows:
        return

    column_count = len(rows[0])
    table = cell.add_table(rows=len(rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    widths = _column_widths_for_table(rows[0])
    _apply_table_geometry(table, widths)

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            current_cell = table.cell(row_index, col_index)
            current_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_table_cell_text(current_cell, value, header=(row_index == 0), col_index=col_index)
            if row_index == 0:
                _set_cell_shading(current_cell, "D9EAF7")

    _repeat_table_header(table)


def _column_widths_for_table(header: list[str]) -> list[int]:
    normalized = {cell.strip() for cell in header}
    if normalized == ACTION_LINKED_TABLE_HEADERS:
        return ACTION_LINKED_TABLE_WIDTHS_DXA
    if normalized == ACTION_TABLE_HEADERS:
        return ACTION_TABLE_WIDTHS_DXA

    column_count = len(header)
    if column_count <= 0:
        return [DEFAULT_TABLE_WIDTH_DXA]
    base = DEFAULT_TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += DEFAULT_TABLE_WIDTH_DXA - sum(widths)
    return widths


def _apply_table_geometry(table, widths_dxa: list[int]) -> None:
    table_width = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    _set_width(tbl_pr, "w:tblW", table_width)
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for col_index, width in enumerate(widths_dxa):
        table.columns[col_index].width = Twips(width)

    for row in table.rows:
        for col_index, current_cell in enumerate(row.cells):
            width = widths_dxa[col_index]
            current_cell.width = Twips(width)
            tc_pr = current_cell._tc.get_or_add_tcPr()
            _set_width(tc_pr, "w:tcW", width)
            _set_cell_margins(current_cell, top=90, bottom=90, start=120, end=120)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def _set_cell_margins(cell, *, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = _ensure_child(tc_pr, "w:tcMar")
    for side, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        margin = _ensure_child(tc_mar, f"w:{side}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_table_cell_text(cell, text: str, *, header: bool, col_index: int) -> None:
    paragraph = cell.paragraphs[0]
    _append_markdown_runs(
        paragraph,
        text,
        size=9.5 if not header else 9,
        bold=header,
        color=RGBColor(31, 78, 121) if header else None,
    )
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
        if col_index == 1 and not header
        else WD_ALIGN_PARAGRAPH.CENTER
    )
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_table_header(table) -> None:
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)
