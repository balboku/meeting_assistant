from __future__ import annotations

import asyncio
import hashlib
import json
import sqlite3
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest import mock

import httpx
from docx import Document

from backend import previous_minutes, tasks


def docx_bytes(*paragraphs: str, with_table: bool = False) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "待辦"
        table.cell(0, 1).text = "負責人"
        table.cell(1, 0).text = "完成風險分析"
        table.cell(1, 1).text = "王小明"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


async def asgi_post(app, path: str, **kwargs):
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app),
        base_url="http://testserver",
    ) as client:
        return await client.post(path, **kwargs)


class PreviousMinutesParserTests(unittest.TestCase):
    def test_extracts_paragraphs_and_tables_with_traceable_hash(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "前次會議.docx"
            content = docx_bytes("前次決議：完成設計輸入。", with_table=True)
            path.write_bytes(content)

            context = previous_minutes.read_previous_minutes_context(path)

        self.assertEqual(context["filename"], "前次會議.docx")
        self.assertEqual(context["sha256"], hashlib.sha256(content).hexdigest())
        self.assertIn("前次決議：完成設計輸入。", context["text"])
        self.assertIn("完成風險分析 | 王小明", context["text"])
        self.assertFalse(context["text_truncated"])

    def test_rejects_renamed_non_docx_content(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "fake.docx"
            path.write_bytes(b"not a zip package")
            with self.assertRaisesRegex(
                previous_minutes.PreviousMinutesError,
                "不是有效的 Word",
            ):
                previous_minutes.read_previous_minutes_context(path)

    def test_rejects_docx_without_readable_text(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "empty.docx"
            path.write_bytes(docx_bytes())
            with self.assertRaisesRegex(
                previous_minutes.PreviousMinutesError,
                "沒有可讀取的文字",
            ):
                previous_minutes.read_previous_minutes_context(path)


class PreviousMinutesSummaryTests(unittest.TestCase):
    def setUp(self):
        self.context = {
            "filename": "前次會議.docx",
            "stored_path": "C:/records/previous.docx",
            "sha256": "a" * 64,
            "size_bytes": 1234,
            "text": "前次待辦：完成設計驗證。",
            "text_chars_extracted": 13,
            "text_truncated": False,
        }

    def test_prompt_treats_word_as_untrusted_background_and_requires_follow_up(self):
        prompt = tasks._build_summary_prompt(
            "[00:00] **[發言者 A]**：今天確認進度。",
            previous_minutes_context=self.context,
        )
        self.assertIn("未受信任的參考資料", prompt)
        self.assertIn("previous_meeting_follow_up", prompt)
        self.assertIn("本次決議、狀態與進度只能根據這份逐字稿判定", prompt)
        self.assertIn("前次待辦：完成設計驗證。", prompt)

    def test_missing_current_evidence_forces_not_discussed(self):
        payload = {
            "discussion_summary": [],
            "final_decisions": [],
            "action_items": [],
            "previous_meeting_follow_up": [
                {
                    "previous_item": "完成設計驗證",
                    "current_status": "completed",
                    "current_update": "已完成",
                    "evidence_timecodes": ["10:00"],
                }
            ],
        }
        normalized = tasks._normalize_summary_payload(
            payload,
            "[00:00] **[發言者 A]**：今天討論其他事項。",
            previous_minutes_context=self.context,
        )
        follow_up = normalized["previous_meeting_follow_up"][0]
        self.assertEqual(follow_up["current_status"], "not_discussed")
        self.assertEqual(follow_up["evidence_timecodes"], [])
        self.assertEqual(follow_up["current_update"], "本次逐字稿未找到可驗證提及")

    def test_current_evidence_preserves_supported_status_and_renders_source(self):
        payload = {
            "discussion_summary": [],
            "final_decisions": [],
            "action_items": [],
            "previous_meeting_follow_up": [
                {
                    "previous_item": "完成設計驗證",
                    "current_status": "completed",
                    "current_update": "本次確認已完成",
                    "evidence_timecodes": ["00:05"],
                }
            ],
        }
        normalized = tasks._normalize_summary_payload(
            payload,
            "[00:05] **[發言者 A]**：設計驗證已完成。",
            previous_minutes_context=self.context,
        )
        markdown = tasks._summary_json_to_markdown(normalized)
        self.assertEqual(
            normalized["previous_meeting_follow_up"][0]["current_status"],
            "completed",
        )
        self.assertIn("## 零、前次會議追蹤", markdown)
        self.assertIn("前次會議.docx", markdown)
        self.assertIn("a" * 64, markdown)
        self.assertIn("已完成", markdown)

    def test_status_is_not_confirmed_when_current_transcript_has_no_timecodes(self):
        normalized = tasks._normalize_summary_payload(
            {
                "discussion_summary": [],
                "final_decisions": [],
                "action_items": [],
                "previous_meeting_follow_up": [{
                    "previous_item": "完成設計驗證",
                    "current_status": "completed",
                    "current_update": "已完成",
                    "evidence_timecodes": ["00:00"],
                }],
            },
            "本次逐字稿沒有時間碼。",
            previous_minutes_context=self.context,
        )
        self.assertEqual(
            normalized["previous_meeting_follow_up"][0]["current_status"],
            "not_discussed",
        )

    def test_no_previous_context_keeps_existing_three_key_contract(self):
        normalized = tasks._normalize_summary_payload(
            {
                "discussion_summary": [],
                "final_decisions": [],
                "action_items": [],
                "previous_meeting_follow_up": [{"previous_item": "不應保留"}],
            },
            "",
        )
        self.assertNotIn("previous_minutes", normalized)
        self.assertNotIn("previous_meeting_follow_up", normalized)


class PreviousMinutesUploadTests(unittest.TestCase):
    def test_upload_retains_docx_and_enqueues_traceable_reference(self):
        from backend import main

        captured = {}
        media = b"ID3" + b"\0" * 64
        prior_docx = docx_bytes("前次待辦：完成設計驗證。")
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            prior_dir = root / "previous_minutes"
            output_dir = root / "output"
            source_dir.mkdir()
            output_dir.mkdir()
            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "PREVIOUS_MINUTES_DIR", prior_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(
                     main,
                     "enqueue_audio_job",
                     side_effect=lambda **kwargs: captured.update(kwargs),
                 ):
                response = asyncio.run(asgi_post(
                    main.app,
                    "/upload-media",
                    files={
                        "file": ("meeting.mp3", media, "audio/mpeg"),
                        "previous_minutes_file": (
                            "前次會議.docx",
                            prior_docx,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    },
                ))

            self.assertEqual(response.status_code, 202)
            self.assertEqual(captured["previous_minutes_filename"], "前次會議.docx")
            self.assertEqual(
                captured["previous_minutes_sha256"],
                hashlib.sha256(prior_docx).hexdigest(),
            )
            self.assertTrue(captured["previous_minutes_path"].is_file())
            self.assertEqual(captured["previous_minutes_path"].parent, prior_dir)

    def test_invalid_docx_is_rejected_and_new_media_is_cleaned(self):
        from backend import main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            prior_dir = root / "previous_minutes"
            output_dir = root / "output"
            source_dir.mkdir()
            output_dir.mkdir()
            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "PREVIOUS_MINUTES_DIR", prior_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "enqueue_audio_job") as enqueue:
                response = asyncio.run(asgi_post(
                    main.app,
                    "/upload-media",
                    files={
                        "file": ("meeting.mp3", b"ID3" + b"\0" * 64, "audio/mpeg"),
                        "previous_minutes_file": (
                            "fake.docx",
                            b"not a Word document",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    },
                ))

            self.assertEqual(response.status_code, 400)
            self.assertIn("不是有效的 Word", response.json()["detail"])
            self.assertFalse(list(source_dir.iterdir()))
            self.assertFalse(list(prior_dir.glob("*")))
            enqueue.assert_not_called()

    def test_frontend_exposes_optional_docx_picker_and_form_field(self):
        root = Path(__file__).resolve().parents[1]
        html = (root / "static" / "index.html").read_text(encoding="utf-8")
        script = (root / "static" / "previous_minutes.js").read_text(encoding="utf-8")
        self.assertIn('id="previous-minutes-input"', html)
        self.assertIn('accept=".docx,', html)
        self.assertIn('script src="/previous_minutes.js"', html)
        self.assertIn("previous_minutes_file", script)
        self.assertIn("本次狀態仍以本次逐字稿為準", script)


class PreviousMinutesQueueTests(unittest.TestCase):
    def test_queue_persists_and_worker_forwards_previous_minutes_reference(self):
        from backend import job_queue

        captured_job = {}
        prior_path = Path("C:/records/prior.docx")
        with mock.patch.object(
            job_queue,
            "create_job",
            side_effect=lambda job_id, **kwargs: captured_job.update(
                {"job_id": job_id, **kwargs}
            ),
        ):
            job_queue.enqueue_audio_job(
                "prior-context-job",
                audio_path=Path("C:/records/audio.mp3"),
                output_dir=Path("C:/records/output"),
                previous_minutes_path=prior_path,
                previous_minutes_filename="前次會議.docx",
                previous_minutes_sha256="b" * 64,
            )

        payload = captured_job["payload"]
        self.assertEqual(payload["previous_minutes_path"], str(prior_path))
        self.assertEqual(payload["previous_minutes_filename"], "前次會議.docx")
        self.assertEqual(payload["previous_minutes_sha256"], "b" * 64)

        captured_task = {}
        worker = job_queue.JobQueueWorker()
        with mock.patch.object(job_queue, "get_meeting_by_job_id", return_value=None), \
             mock.patch.object(
                 job_queue,
                 "process_audio_task",
                 side_effect=lambda **kwargs: captured_task.update(kwargs) or Path("done.md"),
             ), \
             mock.patch.object(worker, "_log_source_audio_retention"):
            worker._process_audio_job({
                "job_id": "prior-context-job",
                "task_type": "audio_processing",
                "payload": payload,
                "worker_id": "worker",
                "worker_generation": 1,
            })

        self.assertEqual(captured_task["previous_minutes_path"], prior_path)
        self.assertEqual(
            captured_task["previous_minutes_filename"],
            "前次會議.docx",
        )
        self.assertEqual(captured_task["previous_minutes_sha256"], "b" * 64)


class PreviousMinutesRecoveryTests(unittest.TestCase):
    def test_snapshot_restores_docx_and_rewrites_durable_references(self):
        from backend import database
        from backend.maintenance import (
            backup_database,
            create_record_snapshot,
            restore_record_snapshot,
            verify_record_snapshot,
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            previous_dir = root / "previous_minutes"
            backup_dir = root / "backups"
            source_dir.mkdir()
            previous_dir.mkdir()
            source = source_dir / "meeting.mp3"
            source.write_bytes(b"ID3-source")
            prior = previous_dir / "prior.docx"
            prior.write_bytes(docx_bytes("前次待辦：完成設計驗證。"))
            output = root / "meeting.md"
            output.write_text("# meeting", encoding="utf-8")
            prior_context = previous_minutes.read_previous_minutes_context(
                prior,
                original_filename="前次正式會議.docx",
            )
            quality_report = {
                "previous_minutes": previous_minutes.previous_minutes_metadata(
                    prior_context,
                    include_path=True,
                )
            }
            job_id = "previous-minutes-snapshot-job"
            with mock.patch.object(database, "DB_PATH", root / "meetings.db"):
                database.init_db()
                database.create_job(
                    job_id,
                    payload={
                        "audio_path": str(source),
                        "output_dir": str(root),
                        "previous_minutes_path": str(prior),
                        "previous_minutes_filename": "前次正式會議.docx",
                        "previous_minutes_sha256": prior_context["sha256"],
                    },
                )
                meeting_id = database.save_meeting(
                    title="快照前次紀錄",
                    date="2026/07/30",
                    source_audio=source.name,
                    output_path=str(output),
                    job_id=job_id,
                    quality_report=quality_report,
                )
                backup = backup_database(database.DB_PATH, backup_dir)
                snapshot = create_record_snapshot(
                    backup,
                    backup_dir,
                    source_media_dir=source_dir,
                    previous_minutes_dir=previous_dir,
                )

            verification = verify_record_snapshot(snapshot)
            self.assertTrue(verification["ok"], verification)
            self.assertEqual(verification["previous_minutes"], 1)
            restored = restore_record_snapshot(snapshot, root / "restore")
            restored_db = Path(restored["runtime"]["runtime_database"])
            conn = sqlite3.connect(restored_db)
            try:
                restored_report = json.loads(conn.execute(
                    "SELECT quality_report_json FROM meetings WHERE id=?",
                    (meeting_id,),
                ).fetchone()[0])
                restored_payload = json.loads(conn.execute(
                    "SELECT payload_json FROM jobs WHERE job_id=?",
                    (job_id,),
                ).fetchone()[0])
            finally:
                conn.close()

            restored_prior = Path(
                restored_report["previous_minutes"]["stored_path"]
            )
            self.assertTrue(restored_prior.is_file())
            self.assertEqual(restored_prior.read_bytes(), prior.read_bytes())
            self.assertEqual(
                Path(restored_payload["previous_minutes_path"]),
                restored_prior,
            )
            self.assertEqual(
                restored_payload["previous_minutes_sha256"],
                prior_context["sha256"],
            )


if __name__ == "__main__":
    unittest.main()
