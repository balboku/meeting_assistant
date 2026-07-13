import asyncio
import inspect
import json
import os
import subprocess
import sys
import tempfile
import unittest
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from unittest import mock

import httpx


ROOT = Path(__file__).resolve().parents[1]


def asgi_request(app, method: str, path: str, asgi_client=("127.0.0.1", 123), **kwargs):
    async def _request():
        transport = httpx.ASGITransport(app=app, client=asgi_client)
        async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
            return await client.request(method, path, **kwargs)

    return asyncio.run(_request())


class SecurityRegressionTests(unittest.TestCase):
    def test_remote_forwarded_requests_are_rejected_without_api_key(self):
        from backend.main import app

        response = asgi_request(app, "GET", "/health", headers={"X-Forwarded-For": "203.0.113.8"})

        self.assertEqual(response.status_code, 403)

    def test_favicon_is_public_and_served_from_logo_asset(self):
        from backend.main import app

        response = asgi_request(app, "GET", "/favicon.ico", headers={"X-Forwarded-For": "203.0.113.8"})

        self.assertEqual(response.status_code, 200)
        self.assertIn("image/x-icon", response.headers.get("content-type", ""))
        self.assertGreater(len(response.content), 0)

    def test_remote_browser_api_key_query_sets_cookie_for_followup_requests(self):
        import backend.main as main

        original_key = main.APP_API_KEY
        main.APP_API_KEY = "mobile-secret"
        try:
            first_response = asgi_request(
                main.app,
                "GET",
                "/history?api_key=mobile-secret",
                headers={"X-Forwarded-For": "203.0.113.8"},
            )
            self.assertIn(first_response.status_code, (307, 308))
            self.assertIn(
                "meeting_assistant_api_key=mobile-secret",
                first_response.headers.get("set-cookie", ""),
            )

            followup_response = asgi_request(
                main.app,
                "GET",
                "/health",
                headers={"X-Forwarded-For": "203.0.113.8"},
                cookies={"meeting_assistant_api_key": "mobile-secret"},
            )
        finally:
            main.APP_API_KEY = original_key

        self.assertEqual(followup_response.status_code, 200)

    def test_same_network_browser_requests_are_allowed_without_api_key(self):
        from backend.main import app

        for client_ip in ("192.168.1.50", "10.0.0.8", "100.84.193.112"):
            response = asgi_request(
                app,
                "GET",
                "/health",
                headers={"X-Forwarded-For": client_ip},
            )
            self.assertEqual(response.status_code, 200, msg=client_ip)

    def test_forwarded_for_is_ignored_from_non_loopback_clients(self):
        from backend.main import app

        response = asgi_request(
            app,
            "GET",
            "/health",
            asgi_client=("198.51.100.9", 456),
            headers={"X-Forwarded-For": "192.168.1.50"},
        )

        self.assertEqual(response.status_code, 403)

    def test_direct_same_network_client_is_allowed_without_forwarded_header(self):
        from backend.main import app

        response = asgi_request(
            app,
            "GET",
            "/health",
            asgi_client=("192.168.1.50", 456),
        )

        self.assertEqual(response.status_code, 200)

    def test_upload_rejects_files_larger_than_configured_limit(self):
        import backend.main as main

        original_limit = main.MAX_UPLOAD_BYTES
        main.MAX_UPLOAD_BYTES = 10
        try:
            response = asgi_request(
                main.app,
                "POST",
                "/upload-media",
                files={"file": ("large.mp3", BytesIO(b"0" * 12), "audio/mpeg")},
            )
        finally:
            main.MAX_UPLOAD_BYTES = original_limit

        self.assertEqual(response.status_code, 413)

    def test_markdown_rendering_is_sanitized_before_inner_html(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn("DOMPurify.sanitize", html)
        self.assertIn("function normalizeMeetingMarkdown", html)
        self.assertIn("DOMPurify.sanitize(marked.parse(normalizeMeetingMarkdown(md)))", html)

    def test_web_ui_uses_local_static_assets_without_cdn_dependencies(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertNotIn("cdn.jsdelivr.net", html)
        self.assertNotIn("fonts.googleapis.com", html)
        self.assertNotIn("fonts.gstatic.com", html)
        self.assertIn('/static/vendor/marked.min.js', html)
        self.assertIn('/static/vendor/purify.min.js', html)
        self.assertIn('<link rel="icon" type="image/png" href="/static/favicon.png" />', html)
        self.assertIn('<link rel="shortcut icon" href="/static/favicon.ico" />', html)
        self.assertIn('class="brand-logo"', html)
        self.assertIn('src="/static/favicon.png"', html)
        self.assertTrue((ROOT / "static" / "vendor" / "marked.min.js").is_file())
        self.assertTrue((ROOT / "static" / "vendor" / "purify.min.js").is_file())
        self.assertTrue((ROOT / "static" / "favicon.png").is_file())
        self.assertTrue((ROOT / "static" / "favicon.ico").is_file())


class ConfigRegressionTests(unittest.TestCase):
    def test_database_and_runtime_paths_can_be_overridden_by_environment(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {
                **os.environ,
                "DB_PATH": str(tmp_path / "custom.db"),
                "MEETING_TEMP_DIR": str(tmp_path / "custom-temp"),
                "MEETING_OUTPUT_DIR": str(tmp_path / "custom-output"),
                "MEETING_SOURCE_AUDIO_DIR": str(tmp_path / "custom-source-audio"),
            }
            script = (
                "import json; "
                "import backend.database as database; "
                "import backend.main as main; "
                "print(json.dumps({"
                "'db': str(database.DB_PATH), "
                "'temp': str(main.TEMP_DIR), "
                "'output': str(main.OUTPUT_DIR), "
                "'source_audio': str(main.SOURCE_AUDIO_DIR)"
                "}, ensure_ascii=False))"
            )
            result = subprocess.run(
                [sys.executable, "-c", script],
                cwd=ROOT,
                env=env,
                check=True,
                capture_output=True,
                text=True,
            )

        payload = json.loads(result.stdout)
        self.assertEqual(payload["db"], str(tmp_path / "custom.db"))
        self.assertEqual(payload["temp"], str(tmp_path / "custom-temp"))
        self.assertEqual(payload["output"], str(tmp_path / "custom-output"))
        self.assertEqual(payload["source_audio"], str(tmp_path / "custom-source-audio"))

    def test_config_endpoint_reports_runtime_upload_limit_and_formats(self):
        import backend.main as main

        original_mb = main.MAX_UPLOAD_MB
        original_bytes = main.MAX_UPLOAD_BYTES
        main.MAX_UPLOAD_MB = 123
        main.MAX_UPLOAD_BYTES = 123 * 1024 * 1024
        try:
            response = asgi_request(main.app, "GET", "/config")
        finally:
            main.MAX_UPLOAD_MB = original_mb
            main.MAX_UPLOAD_BYTES = original_bytes

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["max_upload_mb"], 123)
        self.assertEqual(payload["max_upload_bytes"], 123 * 1024 * 1024)
        self.assertEqual(payload["model"], main.GEMINI_MODEL)
        self.assertEqual(payload["transcription_model"], main.GEMINI_MODEL)
        self.assertEqual(payload["summary_model"], main.SUMMARY_MODEL)
        self.assertEqual(payload["summary_fallback_model"], main.SUMMARY_FALLBACK_MODEL)
        self.assertEqual(payload["summary_verifier_model"], main.SUMMARY_VERIFIER_MODEL)
        self.assertEqual(payload["source_media_archive_retention_days"], main.SOURCE_MEDIA_ARCHIVE_RETENTION_DAYS)
        self.assertFalse(payload["auth"]["enabled"])
        self.assertEqual(payload["recording_profiles"]["audio_standard"]["audio_bps"], 48000)
        self.assertEqual(payload["recording_profiles"]["video_balanced"]["video_fps"], 15)
        self.assertEqual(payload["recording_profiles"]["video_balanced"]["label"], "錄影平衡")
        self.assertIn(".mp3", payload["supported_extensions"])
        self.assertIn(".mp4", payload["supported_extensions"])

    def test_upload_media_is_primary_and_upload_audio_is_deprecated_alias(self):
        import backend.main as main

        response = asgi_request(main.app, "GET", "/openapi.json")

        self.assertEqual(response.status_code, 200)
        paths = response.json()["paths"]
        self.assertIn("/upload-media", paths)
        self.assertIn("/upload-audio", paths)
        self.assertEqual(paths["/upload-media"]["post"]["summary"], "上傳音訊或影片並觸發 AI 處理")
        self.assertFalse(paths["/upload-media"]["post"].get("deprecated", False))
        self.assertTrue(paths["/upload-audio"]["post"]["deprecated"])
        self.assertIn("相容舊路徑", paths["/upload-audio"]["post"]["summary"])


class AuthAuditRegressionTests(unittest.TestCase):
    def test_auth_feature_is_disabled_by_default_but_tables_and_helpers_exist(self):
        import backend.database as database
        from backend.auth import AUTH_FEATURE_ENABLED, ROLE_PERMISSIONS

        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "meetings.db"
            with mock.patch.object(database, "DB_PATH", db_path):
                database.init_db()
                user = database.upsert_app_user(
                    "USER@example.com",
                    display_name="測試同仁",
                    role="editor",
                )
                audit_id = database.record_audit_log(
                    action="meeting.transcript.update",
                    actor_email=user["email"],
                    resource_type="meeting",
                    resource_id="7",
                    detail={"source": "manual_transcript_edit"},
                )
                logs = database.list_audit_logs()

                import sqlite3
                conn = sqlite3.connect(db_path)
                try:
                    tables = {
                        row[0]
                        for row in conn.execute(
                            "SELECT name FROM sqlite_master WHERE type='table'"
                        ).fetchall()
                    }
                finally:
                    conn.close()

        self.assertFalse(AUTH_FEATURE_ENABLED)
        self.assertIn("meeting:write", ROLE_PERMISSIONS["editor"])
        self.assertEqual(user["email"], "user@example.com")
        self.assertGreater(audit_id, 0)
        self.assertEqual(logs[0]["action"], "meeting.transcript.update")
        self.assertEqual(logs[0]["detail"]["source"], "manual_transcript_edit")
        self.assertIn("app_users", tables)
        self.assertIn("audit_logs", tables)


class ProjectGovernanceRegressionTests(unittest.TestCase):
    def test_gitignore_excludes_runtime_artifacts(self):
        gitignore = (ROOT / ".gitignore").read_text(encoding="utf-8")

        for pattern in (
            "__pycache__/",
            "*.pyc",
            "meetings.db",
            "meetings.db-*",
            "temp/",
            "output/",
            "backups/",
            ".env",
        ):
            self.assertIn(pattern, gitignore)

    def test_verify_script_covers_core_local_checks(self):
        verify_script = ROOT / "scripts" / "verify.sh"

        self.assertTrue(verify_script.is_file())
        self.assertTrue(os.access(verify_script, os.X_OK))

        script = verify_script.read_text(encoding="utf-8")
        for command in (
            ".venv/bin/python -m unittest discover -v",
            ".venv/bin/python -m compileall -q backend gui tests meeting_assistant.py start.py test_regex.py test_gemini.py",
            ".venv/bin/python scripts/security_scan.py",
            ".venv/bin/python -m pip check",
            "node --check static/index.html",
        ):
            self.assertIn(command, script)

    def test_quality_benchmark_example_runs_without_ai_calls(self):
        result = subprocess.run(
            [
                sys.executable,
                "scripts/run_quality_benchmark.py",
                "benchmarks/meeting_quality/cases.example.json",
                "--min-score",
                "80",
            ],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
        )
        payload = json.loads(result.stdout)
        self.assertTrue(payload["passed"])
        self.assertEqual(payload["case_count"], 1)
        self.assertGreaterEqual(payload["average_score"], 80)

    def test_quality_benchmark_can_scan_generated_markdown_directory(self):
        sample = (
            "## 一、討論摘要 (Discussion Summary)\n\n"
            "### D1. 測試議題\n- 摘要：保留可追蹤摘要。\n\n"
            "## 二、最終決議 (Final Decisions)\n\n"
            "| # | 關聯討論 | 決議 | 依據 | 狀態 |\n"
            "|---|---|---|---|---|\n"
            "| R1 | D1 | 確認測試。 | [00:00] | confirmed |\n\n"
            "## 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---|---|---|---|---|---|\n"
            "| A1 | D1 | R1 | 完成測試。 | 發言者 A | 未提及 | 中 |\n\n"
            "## 四、完整逐字稿 (Verbatim Transcript)\n\n"
            "### 【第 1 段｜00:00 – 10:00】\n"
            "[00:00] **[發言者 A]**：測試逐字稿。\n"
        )
        with tempfile.TemporaryDirectory() as tmpdir:
            scan_dir = Path(tmpdir)
            (scan_dir / "meeting-one.md").write_text(sample, encoding="utf-8")
            (scan_dir / "ignore.txt").write_text("not markdown", encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    "scripts/run_quality_benchmark.py",
                    "--scan-dir",
                    str(scan_dir),
                    "--limit",
                    "1",
                    "--min-score",
                    "80",
                ],
                cwd=ROOT,
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
            )

        payload = json.loads(result.stdout)
        self.assertTrue(payload["passed"])
        self.assertEqual(payload["case_count"], 1)
        self.assertEqual(payload["results"][0]["id"], "meeting-one")
        self.assertGreaterEqual(payload["results"][0]["score"], 80)

    def test_quality_benchmark_summary_format_is_human_readable(self):
        result = subprocess.run(
            [
                sys.executable,
                "scripts/run_quality_benchmark.py",
                "benchmarks/meeting_quality/cases.example.json",
                "--min-score",
                "80",
                "--format",
                "summary",
            ],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
        )

        self.assertIn("passed=True", result.stdout)
        self.assertIn("golden_meeting_structure", result.stdout)
        self.assertIn("failed=-", result.stdout)

    def test_quality_benchmark_flags_omitted_or_repeated_transcripts(self):
        bad_sample = (
            "## 一、討論摘要 (Discussion Summary)\n\n"
            "### D1. 測試議題\n- 摘要：測試。\n\n"
            "## 二、最終決議 (Final Decisions)\n\n"
            "| # | 關聯討論 | 決議 | 依據 | 狀態 |\n"
            "|---|---|---|---|---|\n"
            "| R1 | D1 | 確認測試。 | [00:00] | confirmed |\n\n"
            "## 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---|---|---|---|---|---|\n"
            "| A1 | D1 | R1 | 完成測試。 | 發言者 A | 未提及 | 中 |\n\n"
            "## 四、完整逐字稿 (Verbatim Transcript)\n\n"
            "### 【第 1 段｜00:00 – 10:00】\n"
            "*(註：為節省篇幅，已省略逐字稿中重複內容)*\n"
            "[00:00] **[發言者 A]**：這一句不應該連續重複。\n"
            "[00:01] **[發言者 A]**：這一句不應該連續重複。\n"
            "[00:02] **[發言者 A]**：這一句不應該連續重複。\n"
            "[00:03] **[發言者 A]**：這一句不應該連續重複。\n"
        )
        with tempfile.TemporaryDirectory() as tmpdir:
            scan_dir = Path(tmpdir)
            (scan_dir / "bad-meeting.md").write_text(bad_sample, encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    "scripts/run_quality_benchmark.py",
                    "--scan-dir",
                    str(scan_dir),
                    "--min-score",
                    "95",
                ],
                cwd=ROOT,
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
            )

        payload = json.loads(result.stdout)
        failed_names = {item["name"] for item in payload["results"][0]["failed"]}
        self.assertEqual(result.returncode, 1)
        self.assertFalse(payload["passed"])
        self.assertIn("transcript_has_no_omission_notice", failed_names)
        self.assertIn("transcript_has_no_repeated_turn_loop", failed_names)

    def test_ci_runs_unit_tests_and_security_scan(self):
        ci = ROOT / ".github" / "workflows" / "ci.yml"
        security_scan = ROOT / "scripts" / "security_scan.py"

        self.assertTrue(ci.is_file())
        self.assertTrue(security_scan.is_file())

        workflow = ci.read_text(encoding="utf-8")
        self.assertIn("python scripts/security_scan.py", workflow)
        self.assertIn("python -m unittest discover -v", workflow)
        self.assertIn("python -m pip check", workflow)

    def test_docs_describe_operational_knobs_events_and_fts(self):
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        architecture = (ROOT / "ARCHITECTURE.md").read_text(encoding="utf-8")

        for env_name in (
            "DB_PATH",
            "MEETING_TEMP_DIR",
            "MEETING_OUTPUT_DIR",
            "MEETING_SOURCE_AUDIO_DIR",
            "MEETING_BACKUP_DIR",
            "DB_BACKUP_KEEP",
            "SOURCE_MEDIA_ARCHIVE_RETENTION_DAYS",
            "JOB_RETENTION_DAYS",
            "MEETING_ASSISTANT_NGROK",
            "MEETING_ASSISTANT_NGROK_URL",
        ):
            self.assertIn(env_name, readme)

        self.assertIn("job_events", architecture)
        self.assertIn("FTS5", architecture)


class MaintenanceRegressionTests(unittest.TestCase):
    def test_backup_database_creates_timestamped_sqlite_copy_and_prunes_old_backups(self):
        from backend.maintenance import backup_database

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            db_path = tmp_path / "meetings.db"
            backup_dir = tmp_path / "backups"
            db_path.write_bytes(b"sqlite-content")

            old_backup = backup_dir / "meetings_20250101_000000.db"
            backup_dir.mkdir()
            old_backup.write_bytes(b"old")

            created = backup_database(
                db_path=db_path,
                backup_dir=backup_dir,
                now=datetime(2026, 7, 5, 14, 0, 0),
                keep=1,
            )

            self.assertEqual(created.name, "meetings_20260705_140000.db")
            self.assertEqual(created.read_bytes(), b"sqlite-content")
            self.assertFalse(old_backup.exists())

    def test_cleanup_source_media_archives_prunes_only_expired_date_buckets(self):
        from backend.maintenance import cleanup_source_media_archives

        with tempfile.TemporaryDirectory() as tmpdir:
            archive_root = Path(tmpdir) / "source_media_deleted"
            expired = archive_root / "20260401"
            retained = archive_root / "20260710"
            ignored = archive_root / "manual-review"
            expired.mkdir(parents=True)
            retained.mkdir()
            ignored.mkdir()
            (expired / "old.webm").write_bytes(b"old-media")
            (expired / "note.txt").write_bytes(b"note")
            (retained / "fresh.webm").write_bytes(b"fresh")
            (ignored / "keep.webm").write_bytes(b"manual")

            result = cleanup_source_media_archives(
                archive_root=archive_root,
                retention_days=30,
                now=datetime(2026, 7, 13, 12, 0, 0),
            )

            self.assertTrue(result["enabled"])
            self.assertEqual(result["deleted_dirs"], 1)
            self.assertEqual(result["deleted_files"], 2)
            self.assertEqual(result["deleted_bytes"], len(b"old-media") + len(b"note"))
            self.assertFalse(expired.exists())
            self.assertTrue((retained / "fresh.webm").exists())
            self.assertTrue((ignored / "keep.webm").exists())

    def test_cleanup_source_media_archives_can_be_disabled(self):
        from backend.maintenance import cleanup_source_media_archives

        with tempfile.TemporaryDirectory() as tmpdir:
            archive_root = Path(tmpdir) / "source_media_deleted"
            expired = archive_root / "20260401"
            expired.mkdir(parents=True)
            (expired / "old.webm").write_bytes(b"old-media")

            result = cleanup_source_media_archives(
                archive_root=archive_root,
                retention_days=0,
                now=datetime(2026, 7, 13, 12, 0, 0),
            )

            self.assertFalse(result["enabled"])
            self.assertEqual(result["deleted_files"], 0)
            self.assertTrue((expired / "old.webm").exists())

    def test_maintain_database_runs_checkpoint_and_vacuum(self):
        from backend.maintenance import maintain_database

        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "meetings.db"
            import sqlite3
            conn = sqlite3.connect(db_path)
            try:
                conn.execute("CREATE TABLE sample (id INTEGER PRIMARY KEY, name TEXT)")
                conn.execute("INSERT INTO sample (name) VALUES ('demo')")
                conn.commit()
            finally:
                conn.close()

            result = maintain_database(db_path=db_path)

            self.assertTrue(result["wal_checkpoint"])
            self.assertTrue(result["vacuum"])

    def test_lifespan_runs_database_maintenance_before_worker_start(self):
        source = (ROOT / "backend" / "main.py").read_text(encoding="utf-8")
        lifespan_body = source[
            source.index("async def lifespan") :
            source.index("app = FastAPI(")
        ]

        self.assertIn("run_startup_maintenance", lifespan_body)
        self.assertIn("cleanup_source_media_archives", lifespan_body)
        self.assertLess(
            lifespan_body.index("run_startup_maintenance"),
            lifespan_body.index("cleanup_source_media_archives"),
        )
        self.assertLess(
            lifespan_body.index("cleanup_source_media_archives"),
            lifespan_body.index("job_worker.start()"),
        )


class StartupHealthRegressionTests(unittest.TestCase):
    def test_startup_health_reports_missing_api_key_and_unwritable_paths(self):
        from backend.maintenance import run_startup_health_checks

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            checks = run_startup_health_checks(
                temp_dir=tmp_path / "missing-temp",
                output_dir=tmp_path / "missing-output",
                static_vendor_dir=tmp_path / "missing-vendor",
                env={},
            )

        names = {check["name"]: check for check in checks}
        self.assertEqual(names["gemini_api_key"]["status"], "failed")
        self.assertEqual(names["temp_dir"]["status"], "failed")
        self.assertEqual(names["output_dir"]["status"], "failed")
        self.assertEqual(names["static_vendor"]["status"], "failed")

    def test_health_endpoint_includes_startup_checks(self):
        import backend.main as main

        response = asgi_request(main.app, "GET", "/health")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertIn("checks", payload)
        self.assertFalse(payload["auth"]["enabled"])
        self.assertTrue(any(check["name"] == "database" for check in payload["checks"]))


class MediaValidationRegressionTests(unittest.TestCase):
    def test_magic_sniff_accepts_mp3_wav_m4a_and_rejects_fake_mp3(self):
        from backend.media_validation import validate_media_magic

        self.assertIsNone(validate_media_magic(".mp3", b"ID3\x04\x00\x00\x00\x00\x00\x21audio"))
        self.assertIsNone(validate_media_magic(".wav", b"RIFF\x24\x00\x00\x00WAVEfmt "))
        self.assertIsNone(validate_media_magic(".m4a", b"\x00\x00\x00\x18ftypM4A \x00\x00\x00\x00"))
        self.assertIn("檔案內容", validate_media_magic(".mp3", b"<html>not audio</html>"))

    def test_upload_rejects_extension_that_does_not_match_magic(self):
        import backend.main as main

        response = asgi_request(
            main.app,
            "POST",
            "/upload-media",
            files={"file": ("fake.mp3", BytesIO(b"<html>not audio</html>"), "audio/mpeg")},
        )

        self.assertEqual(response.status_code, 415)


class TaskRegressionTests(unittest.TestCase):
    def test_audio_processing_task_is_sync_so_background_tasks_use_threadpool(self):
        from backend.tasks import process_audio_task

        self.assertFalse(inspect.iscoroutinefunction(process_audio_task))

    def test_backend_prompts_preserve_multilingual_transcript_policy(self):
        from backend import tasks

        main_prompt = tasks.MEETING_PROMPT
        segment_prompt = tasks._build_segment_prompt(0, 2)

        for prompt in (main_prompt, segment_prompt):
            self.assertIn("英文發言請保留英文原文", prompt)
            self.assertIn("中文國語發言請以繁體中文轉寫", prompt)
            self.assertIn("台語發言請標記為 `[台語]`", prompt)
            self.assertIn("摘要、決議與待辦事項仍統一使用繁體中文", prompt)

    def test_cli_prompt_preserves_multilingual_transcript_policy(self):
        import meeting_assistant

        prompt = meeting_assistant.build_meeting_prompt()

        self.assertIn("英文發言請保留英文原文", prompt)
        self.assertIn("中文國語發言請以繁體中文轉寫", prompt)
        self.assertIn("台語發言請標記為 `[台語]`", prompt)
        self.assertIn("摘要、決議與待辦事項仍統一使用繁體中文", prompt)

    def test_backend_prompts_require_anonymous_speaker_differentiation(self):
        from backend import tasks

        main_prompt = tasks.MEETING_PROMPT
        segment_prompt = tasks._build_segment_prompt(0, 2)

        for prompt in (main_prompt, segment_prompt):
            self.assertIn("目標是分辨「不同聲音」", prompt)
            self.assertIn("同一個聲音再次出現時必須沿用相同標籤", prompt)
            self.assertIn("不要把不同人的發言合併成同一位", prompt)
            self.assertIn("**[多人重疊]**", prompt)

    def test_cli_prompt_requires_anonymous_speaker_differentiation(self):
        import meeting_assistant

        prompt = meeting_assistant.build_meeting_prompt()

        self.assertIn("目標是分辨「不同聲音」", prompt)
        self.assertIn("同一個聲音再次出現時必須沿用相同標籤", prompt)
        self.assertIn("不要把不同人的發言合併成同一位", prompt)
        self.assertIn("若只能辨識到匿名發言者", prompt)

    def test_backend_prompts_use_medical_device_rnd_analysis_policy(self):
        from backend import tasks

        prompts = (
            tasks.MEETING_PROMPT,
            tasks._build_segment_prompt(0, 2),
            tasks._build_summary_prompt("佳世達 IEC 62304 SRS SDS traceability matrix"),
        )

        self.assertIn("「佳世達」為正確名稱", prompts[0])
        self.assertIn("請勿寫成「加斯達」、「嘉士達」或 Jasta", prompts[0])
        self.assertIn("IEC 62304", prompts[1])
        self.assertIn("討論摘要需依「專案/議題」分組", prompts[2])
        self.assertIn("最終決議只放已確認", prompts[2])
        self.assertIn("待辦事項只放可驗收行動", prompts[2])

    def test_cli_prompt_uses_medical_device_rnd_analysis_policy(self):
        import meeting_assistant

        prompt = meeting_assistant.build_meeting_prompt()

        self.assertIn("「佳世達」為正確名稱", prompt)
        self.assertIn("久方生技 / Maxima Biotech", prompt)
        self.assertIn("討論摘要需依「專案/議題」分組", prompt)
        self.assertIn("不要把追蹤目標、背景說明或教學內容列為決議", prompt)

    def test_domain_term_normalization_and_quality_notice(self):
        from backend.tasks import (
            _normalize_domain_terms,
            _prepend_transcript_quality_notice,
        )

        normalized = _normalize_domain_terms(
            "加斯達、嘉士達與 Jasta 文件提到 IEC 6304 與 IC6304，"
            "放電字句、平保、平寶、氣械老化、頻率政府與內型固定塊。"
            "Qisda (佳世達) 需要提升效率 $ ightarrow$ 降低功耗。"
        )

        self.assertIn("佳世達", normalized)
        self.assertNotIn("加斯達", normalized)
        self.assertNotIn("嘉士達", normalized)
        self.assertNotIn("Jasta", normalized)
        self.assertNotIn("佳世達 (佳世達)", normalized)
        self.assertNotIn("$ ightarrow$", normalized)
        self.assertIn("→", normalized)
        self.assertIn("佳世達", normalized)
        self.assertNotIn("IEC 6304", normalized)
        self.assertIn("IEC 62304", normalized)
        self.assertIn("放電治具", normalized)
        self.assertIn("品保", normalized)
        self.assertIn("機械老化", normalized)
        self.assertIn("頻率振幅", normalized)
        self.assertIn("內徑固定塊", normalized)
        self.assertNotIn("平保", normalized)
        self.assertNotIn("平寶", normalized)

        content = "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n## ✅ 二、最終決議 (Final Decisions)\n決議"
        transcript = "[系統提示：此處音檔包含無意義雜訊，已自動過濾後續重複內容]"
        with_notice = _prepend_transcript_quality_notice(content, transcript)

        self.assertIn("逐字稿品質註記", with_notice)
        self.assertIn("可能缺漏", with_notice)

    def test_replace_transcript_section_restores_verbatim_transcript(self):
        from backend.tasks import _extract_post_transcript_sections, _extract_transcript_section_body, _replace_transcript_section

        repaired_by_model = (
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n"
            "---\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "*(註：為節省篇幅，已過濾逐字稿中重複內容)*\n"
            "[38:17] **[發言者 A]**：只保留摘要化片段。\n"
            "\n## 📎 五、補充資料與佐證 (Supplementary Evidence)\n\n### 資料：spec.pdf\n- 系統判斷：保留。\n"
        )
        full_transcript = (
            "### 【第 1 段｜00:00 – 10:00】\n"
            "[00:00] **[發言者 A]**：完整第一句。\n"
            "[09:56] **[發言者 A]**：完整第一段結尾。\n\n"
            "### 【第 2 段｜10:00 – 20:00】\n"
            "[10:00] **[發言者 A]**：不可省略的第二段逐字稿。\n"
        )

        result = _replace_transcript_section(repaired_by_model, full_transcript)

        self.assertIn(full_transcript, result)
        self.assertNotIn("為節省篇幅", result)
        self.assertNotIn("只保留摘要化片段", result)
        self.assertEqual(result.count("## 📝 四、完整逐字稿 (Verbatim Transcript)"), 1)
        self.assertNotIn("---\n\n---", result)
        self.assertIn("## 📎 五、補充資料與佐證", result)
        self.assertIn("spec.pdf", _extract_post_transcript_sections(result))
        self.assertNotIn("補充資料", _extract_transcript_section_body(result))

    def test_transcript_integrity_rejects_omitted_transcript_section(self):
        from backend.tasks import _replace_transcript_section, _transcript_integrity_issues

        full_transcript = (
            "### 【第 1 段｜00:00 – 10:00】\n"
            "[00:00] **[發言者 A]**：完整第一句。\n"
            "[02:00] **[發言者 B]**：完整第二句。\n"
            "[09:56] **[發言者 A]**：完整第一段結尾。\n\n"
            "### 【第 2 段｜10:00 – 20:00】\n"
            "[10:00] **[發言者 A]**：不可省略的第二段逐字稿。\n"
            "[19:55] **[發言者 B]**：第二段結尾。"
        )
        omitted_transcript = (
            "*(註：為節省篇幅，已省略逐字稿中重複內容)*\n"
            "[00:00] **[發言者 A]**：只保留摘要化片段。"
        )
        content = _replace_transcript_section(
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要",
            omitted_transcript,
        )

        issues = _transcript_integrity_issues(content, full_transcript)

        self.assertTrue(any("省略" in issue for issue in issues))
        self.assertTrue(any("不一致" in issue for issue in issues))
        self.assertTrue(any("時間戳" in issue for issue in issues))

    def test_transcript_integrity_rejects_final_short_loop(self):
        from backend.tasks import _replace_transcript_section, _transcript_integrity_issues

        looped_transcript = "\n".join(
            f"[{80 + index // 30:02d}:{(index * 2) % 60:02d}] **[發言者 A]**：這兩段。"
            for index in range(30)
        )
        content = _replace_transcript_section(
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要",
            looped_transcript,
        )

        issues = _transcript_integrity_issues(content, looped_transcript)

        self.assertTrue(any("短句重複轉錄幻覺" in issue for issue in issues))

    def test_finalize_meeting_content_restores_transcript_and_validates(self):
        from backend.tasks import _finalize_meeting_content, _meeting_content_quality_issues

        summary_with_bad_transcript = (
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n決議\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|---------|---------|--------|------|--------|\n"
            "| A1 | D1 | R1 | 整理需求 | 發言者 A | 未提及 | 中 |\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "*(註：為節省篇幅，已過濾逐字稿中重複內容)*\n"
            "[00:00] **[發言者 A]**：摘要化片段。"
        )
        full_transcript = (
            "### 【第 1 段｜00:00 – 10:00】\n"
            "[00:00] **[發言者 A]**：完整逐字稿第一句。\n"
            "[09:56] **[發言者 B]**：完整逐字稿結尾。"
        )

        finalized = _finalize_meeting_content(summary_with_bad_transcript, full_transcript, "final-gate-job")

        self.assertIn(full_transcript, finalized)
        self.assertNotIn("為節省篇幅", finalized)
        self.assertEqual(_meeting_content_quality_issues(finalized), [])

    def test_summary_json_response_converts_to_readable_markdown(self):
        from backend.tasks import _summary_response_to_markdown

        response_text = json.dumps(
            {
                "discussion_summary": [
                    {
                        "id": "D1",
                        "topic": "佳世達測試",
                        "context": "會議討論 SRS 與 SDS 文件缺口。",
                        "summary": "**確認** SRS 與 SDS traceability matrix 需要補齊。",
                        "key_points": ["SRS 需求未完整對應", "SDS 設計文件需同步更新"],
                        "impact": "影響後續驗證追蹤。",
                        "evidence_timecodes": ["12:30"],
                    },
                    {
                        "id": "D2",
                        "topic": "驗證時程",
                        "summary": "測試排程需等缺口清單完成後再確認。",
                        "evidence_timecodes": ["18:20"],
                    },
                ],
                "final_decisions": [
                    {
                        "id": "R1",
                        "related_discussions": ["D1"],
                        "decision": "下次會議前先整理缺口清單。",
                        "basis": "逐字稿提到先做清單再回來討論。",
                        "status": "confirmed",
                    }
                ],
                "action_items": [
                    {
                        "id": "A1",
                        "related_discussions": ["D1"],
                        "related_decisions": ["R1"],
                        "task": "整理佳世達需求缺口",
                        "owner": "QA",
                        "due": "2026/07/10",
                        "priority": "高",
                    }
                ],
            },
            ensure_ascii=False,
        )

        markdown = _summary_response_to_markdown(response_text)

        self.assertIn("## 一、討論摘要 (Discussion Summary)", markdown)
        self.assertIn("## 二、最終決議 (Final Decisions)", markdown)
        self.assertIn("## 三、待辦事項 (Action Items)", markdown)
        self.assertIn("### D1. 佳世達測試", markdown)
        self.assertIn("### D2. 驗證時程", markdown)
        self.assertIn("| # | 關聯討論 | 決議 | 依據 | 狀態 |", markdown)
        self.assertIn("| R1 | D1 | 下次會議前先整理缺口清單。 | 逐字稿提到先做清單再回來討論。 | confirmed |", markdown)
        self.assertIn("| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |", markdown)
        self.assertIn("| A1 | D1 | R1 | 整理佳世達需求缺口 | QA | 2026/07/10 | 高 |", markdown)
        self.assertNotIn("**", markdown)

    def test_segment_transcript_timestamps_are_offset_to_global_time(self):
        from backend.tasks import _offset_transcript_timestamps

        transcript = "[00:00] **[發言者 A]**：第二段開始。\n[09:59] **[發言者 B]**：第二段結束。"

        adjusted = _offset_transcript_timestamps(transcript, offset_seconds=600)

        self.assertIn("[10:00] **[發言者 A]**：第二段開始。", adjusted)
        self.assertIn("[19:59] **[發言者 B]**：第二段結束。", adjusted)
        self.assertNotIn("[00:00]", adjusted)

    def test_recovered_transcript_blocks_are_sorted_by_timestamp(self):
        from backend.tasks import _sort_transcript_blocks_by_timestamp

        transcript = "\n".join(
            [
                "[30:00] **[發言者 A]**：第一小段開始。",
                "[30:33] **[發言者 A]**：第一小段略微越界。",
                "[30:30] **[發言者 B]**：第二小段開始。",
                "[30:46] **[發言者 A]**：後續內容。",
            ]
        )

        sorted_transcript = _sort_transcript_blocks_by_timestamp(transcript)

        self.assertLess(sorted_transcript.find("[30:30]"), sorted_transcript.find("[30:33]"))
        self.assertLess(sorted_transcript.find("[30:33]"), sorted_transcript.find("[30:46]"))

    def test_segment_transcript_quality_rejects_incomplete_nonterminal_segments(self):
        from backend.tasks import _segment_transcript_quality_issues

        incomplete = (
            "[40:00] **[發言者 A]**：只轉出段落開頭。\n"
            "[系統提示：此處音檔包含無意義雜訊，已自動過濾後續重複內容]"
        )
        complete = (
            "[40:00] **[發言者 A]**：段落開始。\n"
            "[48:30] **[發言者 B]**：段落接近結尾。"
        )

        issues = _segment_transcript_quality_issues(
            incomplete,
            segment_index=4,
            total_segments=10,
            segment_minutes=10,
        )

        self.assertTrue(any("自動過濾" in issue for issue in issues))
        self.assertTrue(any("未接近段尾" in issue for issue in issues))
        self.assertEqual(
            _segment_transcript_quality_issues(
                complete,
                segment_index=4,
                total_segments=10,
                segment_minutes=10,
            ),
            [],
        )

    def test_segment_transcript_quality_rejects_repeated_hallucinated_turns(self):
        from backend.tasks import _segment_transcript_quality_issues

        repeated = "\n".join(
            f"[{20 + index // 12:02d}:{(index * 5) % 60:02d}] **[發言者 A]**："
            "那如果我們在緊縮一級的情況下，它是不合格的，那分數就歸零。"
            for index in range(24)
        )

        issues = _segment_transcript_quality_issues(
            repeated,
            segment_index=2,
            total_segments=5,
            segment_minutes=10,
        )

        self.assertTrue(any("重複轉錄幻覺" in issue for issue in issues))

    def test_segment_transcript_quality_rejects_single_line_repetition_hallucination(self):
        from backend.tasks import _segment_transcript_quality_issues

        repeated_phrase = "那我們在這些範圍內，那我們就可以去對表看，對表看，因為"
        transcript = (
            "[30:00] **[發言者 A]**："
            + repeated_phrase * 80
            + "\n[39:55] **[發言者 B]**：這段已經接近段尾。"
        )

        issues = _segment_transcript_quality_issues(
            transcript,
            segment_index=3,
            total_segments=5,
            segment_minutes=10,
        )

        self.assertTrue(any("單句重複轉錄幻覺" in issue for issue in issues))

    def test_segment_transcript_quality_rejects_short_repeated_turn_hallucination(self):
        from backend.tasks import _segment_transcript_quality_issues

        transcript = "\n".join(
            f"[{80 + index // 30:02d}:{(index * 2) % 60:02d}] **[發言者 A]**：這兩段。"
            for index in range(30)
        )

        issues = _segment_transcript_quality_issues(
            transcript,
            segment_index=8,
            total_segments=14,
            segment_minutes=10,
        )

        self.assertTrue(any("短句重複轉錄幻覺" in issue for issue in issues))

    def test_segment_transcript_cache_round_trips_and_rejects_mismatched_context(self):
        from backend.tasks import (
            _load_segment_transcript_cache,
            _save_segment_transcript_cache,
            _segment_cache_context,
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")

            context = _segment_cache_context(
                audio_path=audio_path,
                model="gemini-test",
                total_segments=2,
                segment_minutes=10,
            )
            _save_segment_transcript_cache(
                output_dir=output_dir,
                job_id="resume-job",
                segment_index=0,
                context=context,
                transcript="[00:00] **[發言者 A]**：第一段開始。\n[09:30] **[發言者 A]**：已完成第一段。",
            )

            cached = _load_segment_transcript_cache(output_dir, "resume-job", 0, context)
            self.assertIn("已完成第一段", cached)

            mismatched = dict(context)
            mismatched["model"] = "another-model"
            self.assertIsNone(_load_segment_transcript_cache(output_dir, "resume-job", 0, mismatched))

    def test_segment_transcript_cache_rejects_incomplete_cached_segment(self):
        from backend.tasks import (
            _load_segment_transcript_cache,
            _save_segment_transcript_cache,
            _segment_cache_context,
            _segment_cache_file,
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")

            context = _segment_cache_context(
                audio_path=audio_path,
                model="gemini-test",
                total_segments=2,
                segment_minutes=10,
            )
            _save_segment_transcript_cache(
                output_dir=output_dir,
                job_id="resume-job",
                segment_index=0,
                context=context,
                transcript=(
                    "[00:00] **[發言者 A]**：只有開頭。\n"
                    "[系統提示：此處音檔包含無意義雜訊，已自動過濾後續重複內容]"
                ),
            )

            self.assertIsNone(_load_segment_transcript_cache(output_dir, "resume-job", 0, context))
            self.assertFalse(_segment_cache_file(output_dir, "resume-job", 0).exists())

    def test_segmented_audio_task_reuses_cached_transcripts(self):
        import backend.tasks as tasks

        summary_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "摘要。\n\n"
            "---\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "尚未決定。\n\n"
            "---\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|--------|------|--------|\n"
            "| 1 | 確認後續事項 | 發言者 A | 未提及 | 中 |\n"
        )

        class FakeModels:
            def generate_content(self, **kwargs):
                return type("Response", (), {"text": summary_content})()

        class FakeClient:
            def __init__(self, *args, **kwargs):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")
            seg1 = root / "_seg_meeting_000.mp3"
            seg2 = root / "_seg_meeting_001.mp3"
            seg1.write_bytes(b"seg1")
            seg2.write_bytes(b"seg2")

            context = tasks._segment_cache_context(
                audio_path=audio_path,
                model="gemini-test",
                total_segments=2,
                segment_minutes=tasks.SEGMENT_MINUTES,
            )
            tasks._save_segment_transcript_cache(
                output_dir=output_dir,
                job_id="resume-job",
                segment_index=0,
                context=context,
                transcript="[00:00] **[發言者 A]**：快取第一段。\n[09:30] **[發言者 A]**：第一段結束。",
            )

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", side_effect=FakeClient), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[seg1, seg2]), \
                 mock.patch.object(tasks, "_transcribe_segment", return_value="[00:00] **[發言者 B]**：新轉錄第二段。") as transcribe_mock, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="resume-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-test",
                )

            self.assertIsNotNone(output_path)
            transcribe_mock.assert_called_once()
            self.assertEqual(transcribe_mock.call_args.args[2], 1)
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("快取第一段", content)
            self.assertIn("[10:00] **[發言者 B]**：新轉錄第二段。", content)

    def test_segmented_audio_task_recovers_incomplete_segment_by_splitting_chunk(self):
        import backend.tasks as tasks

        summary_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "D1：補救後完成逐字稿。\n\n"
            "---\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "R1（關聯 D1）：尚未決定。\n\n"
            "---\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|---------|---------|--------|------|--------|\n"
            "| A1 | D1 | R1 | 確認後續事項 | 發言者 A | 未提及 | 中 |\n"
        )

        class FakeModels:
            def generate_content(self, **kwargs):
                return type("Response", (), {"text": summary_content})()

        class FakeClient:
            def __init__(self, *args, **kwargs):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")
            seg1 = root / "_seg_meeting_000.mp3"
            seg2 = root / "_seg_meeting_001.mp3"
            sub1 = root / "_sub__seg_meeting_000_300s_000.mp3"
            sub2 = root / "_sub__seg_meeting_000_300s_001.mp3"
            for path in (seg1, seg2, sub1, sub2):
                path.write_bytes(path.name.encode("utf-8"))

            incomplete = (
                "[00:00] **[發言者 A]**：只有開頭。\n"
                "[系統提示：此處音檔包含無意義雜訊，已自動過濾後續重複內容]"
            )
            recovered_first_half = (
                "[00:00] **[發言者 A]**：補救前半段。\n"
                "[04:30] **[發言者 A]**：前半段結束。"
            )
            recovered_second_half = (
                "[00:00] **[發言者 B]**：補救後半段。\n"
                "[04:30] **[發言者 B]**：後半段結束。"
            )
            second_segment = (
                "[00:00] **[發言者 C]**：第二段。\n"
                "[09:30] **[發言者 C]**：第二段結束。"
            )

            context = tasks._segment_cache_context(
                audio_path=audio_path,
                model="gemini-test",
                total_segments=2,
                segment_minutes=tasks.SEGMENT_MINUTES,
            )

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", side_effect=FakeClient), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[seg1, seg2]), \
                 mock.patch.object(tasks, "_split_audio_to_subsegments", return_value=[(sub1, 0, 300), (sub2, 300, 600)]), \
                 mock.patch.object(
                     tasks,
                     "_transcribe_segment",
                     side_effect=[incomplete, recovered_first_half, recovered_second_half, second_segment],
                 ) as transcribe_mock, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="recover-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-test",
                )

            self.assertIsNotNone(output_path)
            self.assertEqual(transcribe_mock.call_count, 4)
            self.assertEqual([call.args[1] for call in transcribe_mock.call_args_list], [seg1, sub1, sub2, seg2])

            content = output_path.read_text(encoding="utf-8")
            self.assertIn("[00:00] **[發言者 A]**：補救前半段。", content)
            self.assertIn("[05:00] **[發言者 B]**：補救後半段。", content)
            self.assertIn("[10:00] **[發言者 C]**：第二段。", content)

            cached = tasks._load_segment_transcript_cache(output_dir, "recover-job", 0, context)
            self.assertIsNotNone(cached)
            self.assertIn("[05:00] **[發言者 B]**：補救後半段。", cached)
            self.assertFalse(sub1.exists())
            self.assertFalse(sub2.exists())

    def test_single_segment_audio_task_uses_dual_model_pipeline(self):
        import backend.tasks as tasks

        summary_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "單段音檔也使用摘要模型。\n\n"
            "---\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "尚未決定。\n\n"
            "---\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|--------|------|--------|\n"
            "| 1 | 確認後續事項 | 發言者 A | 未提及 | 中 |\n"
        )

        class FakeModels:
            def __init__(self):
                self.calls = []

            def generate_content(self, **kwargs):
                self.calls.append(kwargs)
                return type("Response", (), {"text": summary_content})()

        class FakeClient:
            def __init__(self):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.mp3"
            audio_path.write_bytes(b"audio")

            fake_client = FakeClient()
            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", return_value=fake_client), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[audio_path]), \
                 mock.patch.object(tasks, "_transcribe_segment", return_value="[00:00] **[發言者 A]**：逐字稿內容。") as transcribe_mock, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="single-dual-model-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-transcribe-test",
                    summary_model="gemma-4-31b-it",
                    summary_fallback_model="gemini-transcribe-test",
                )

            self.assertIsNotNone(output_path)
            transcribe_mock.assert_called_once()
            self.assertEqual(transcribe_mock.call_args.args[5], "gemini-transcribe-test")
            self.assertEqual([call["model"] for call in fake_client.models.calls], ["gemma-4-31b-it"])
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("單段音檔也使用摘要模型", content)
            self.assertIn("### [Segment 1/1 | 00:00 - end]", content)
            self.assertIn("transcription_model: gemini-transcribe-test", content)
            self.assertIn("summary_model: gemma-4-31b-it", content)

    def test_audio_task_preserves_verbatim_transcript_after_repair(self):
        import backend.tasks as tasks

        malformed_summary = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "摘要。\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "R1（關聯 D1）：尚未決定。\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|\n"
            "| A1 | D1 | R1 | 整理需求 | 王經理 | 下週三 | 高 |\n"
        )
        repaired_with_short_transcript = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "摘要。\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "R1（關聯 D1）：尚未決定。\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|---------|---------|--------|------|--------|\n"
            "| A1 | D1 | R1 | 整理需求 | 王經理 | 下週三 | 高 |\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "*(註：為節省篇幅，已過濾逐字稿中重複內容)*\n"
            "[00:00] **[發言者 A]**：摘要化片段。\n"
        )
        full_transcript = (
            "[00:00] **[發言者 A]**：完整逐字稿第一句。\n"
            "[00:30] **[發言者 A]**：不可省略的第二句。"
        )

        class FakeModels:
            def __init__(self):
                self.calls = []

            def generate_content(self, **kwargs):
                self.calls.append(kwargs)
                text = malformed_summary if len(self.calls) == 1 else repaired_with_short_transcript
                return type("Response", (), {"text": text})()

        class FakeClient:
            def __init__(self):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.mp3"
            audio_path.write_bytes(b"audio")
            fake_client = FakeClient()

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", return_value=fake_client), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[audio_path]), \
                 mock.patch.object(tasks, "_transcribe_segment", return_value=full_transcript), \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="preserve-transcript-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-transcribe-test",
                    summary_model="gemma-4-31b-it",
                    summary_fallback_model="gemini-transcribe-test",
                )

            self.assertIsNotNone(output_path)
            self.assertEqual(len(fake_client.models.calls), 2)
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("[00:30] **[發言者 A]**：不可省略的第二句。", content)
            self.assertNotIn("為節省篇幅", content)
            self.assertNotIn("摘要化片段", content)
            self.assertEqual(content.count("## 📝 四、完整逐字稿 (Verbatim Transcript)"), 1)

    def test_segmented_audio_task_uses_summary_model_with_fallback(self):
        import backend.tasks as tasks

        summary_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "以備援模型完成摘要。\n\n"
            "---\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "尚未決定。\n\n"
            "---\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|--------|------|--------|\n"
            "| 1 | 確認後續事項 | 發言者 A | 未提及 | 中 |\n"
        )

        class FakeModels:
            def __init__(self):
                self.calls = []

            def generate_content(self, **kwargs):
                self.calls.append(kwargs)
                if kwargs["model"] == "gemma-4-31b-it":
                    raise RuntimeError("primary summary model unavailable")
                return type("Response", (), {"text": summary_content})()

        class FakeClient:
            def __init__(self):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")
            seg1 = root / "_seg_meeting_000.mp3"
            seg2 = root / "_seg_meeting_001.mp3"
            seg1.write_bytes(b"seg1")
            seg2.write_bytes(b"seg2")

            context = tasks._segment_cache_context(
                audio_path=audio_path,
                model="gemini-test",
                total_segments=2,
                segment_minutes=tasks.SEGMENT_MINUTES,
            )
            tasks._save_segment_transcript_cache(
                output_dir=output_dir,
                job_id="dual-model-job",
                segment_index=0,
                context=context,
                transcript="[00:00] **[發言者 A]**：第一段開始。\n[09:30] **[發言者 A]**：第一段結束。",
            )
            tasks._save_segment_transcript_cache(
                output_dir=output_dir,
                job_id="dual-model-job",
                segment_index=1,
                context=context,
                transcript="[10:00] **[發言者 B]**：第二段內容。",
            )

            fake_client = FakeClient()
            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", return_value=fake_client), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[seg1, seg2]), \
                 mock.patch.object(tasks, "_transcribe_segment") as transcribe_mock, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="dual-model-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-test",
                    summary_model="gemma-4-31b-it",
                    summary_fallback_model="gemini-test",
                )

            self.assertIsNotNone(output_path)
            transcribe_mock.assert_not_called()
            self.assertEqual(
                [call["model"] for call in fake_client.models.calls],
                ["gemma-4-31b-it", "gemini-test"],
            )
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("以備援模型完成摘要", content)
            self.assertIn("transcription_model: gemini-test", content)
            self.assertIn("summary_model: gemini-test", content)
            self.assertIn("summary_fallback_model: gemini-test", content)
            self.assertIn("recording_profile: legacy_upload", content)
            self.assertIn("source_audio_size_bytes:", content)

    def test_segmented_audio_task_rejects_fresh_incomplete_segment(self):
        import backend.tasks as tasks

        class FakeModels:
            def generate_content(self, **kwargs):
                raise AssertionError("不應在缺段時生成摘要")

        class FakeClient:
            def __init__(self, *args, **kwargs):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_dir = root / "output"
            output_dir.mkdir()
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")
            seg1 = root / "_seg_meeting_000.mp3"
            seg2 = root / "_seg_meeting_001.mp3"
            seg1.write_bytes(b"seg1")
            seg2.write_bytes(b"seg2")

            incomplete = (
                "[00:00] **[發言者 A]**：只有開頭。\n"
                "[系統提示：此處音檔包含無意義雜訊，已自動過濾後續重複內容]"
            )

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", side_effect=FakeClient), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[seg1, seg2]), \
                 mock.patch.object(tasks, "_split_audio_to_subsegments", side_effect=RuntimeError("split unavailable")), \
                 mock.patch.object(tasks, "_transcribe_segment", return_value=incomplete), \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting") as save_mock:
                output_path = tasks.process_audio_task(
                    job_id="incomplete-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    model="gemini-test",
                )

            self.assertIsNone(output_path)
            save_mock.assert_not_called()
            self.assertFalse(tasks._segment_cache_file(output_dir, "incomplete-job", 0).exists())

    def test_summary_preview_supports_discussion_summary_heading(self):
        from backend.tasks import _extract_summary_preview

        content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n"
            "這是摘要內容。\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n"
            "決議內容"
        )

        self.assertEqual(_extract_summary_preview(content), "這是摘要內容。")

    def test_hallucination_cleanup_preserves_markdown_table_separator_rows(self):
        from backend.tasks import clean_hallucinated_loops

        content = (
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|--------|------|--------|\n"
            "| 1 | 整理產品需求 | 王經理 | 下週三 | 高 |\n"
            "| 2 | 確認測試計畫 | 李小姐 | 未提及 | 高 |\n"
        )

        self.assertEqual(clean_hallucinated_loops(content), content)

    def test_meeting_content_quality_detects_missing_sections_and_broken_table(self):
        from backend.tasks import _meeting_content_quality_issues

        broken_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n決議\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|\n"
        )

        issues = _meeting_content_quality_issues(broken_content)

        self.assertIn("缺少完整逐字稿區塊", issues)
        self.assertIn("待辦事項表格分隔列不完整", issues)

    def test_meeting_content_repair_uses_gemini_once_when_quality_check_fails(self):
        from backend.tasks import _repair_meeting_content_if_needed, _meeting_content_quality_issues

        broken_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n決議\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|\n"
        )
        repaired_content = (
            "## 📋 一、討論摘要 (Discussion Summary)\n摘要\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n決議\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---------|--------|------|--------|\n"
            "| 1 | 整理需求 | 王經理 | 下週三 | 高 |\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n[00:00] **[發言者]**：內容。\n"
        )

        class FakeModels:
            def __init__(self):
                self.calls = []

            def generate_content(self, **kwargs):
                self.calls.append(kwargs)
                return type("Response", (), {"text": repaired_content})()

        class FakeClient:
            def __init__(self):
                self.models = FakeModels()

        client = FakeClient()

        result = _repair_meeting_content_if_needed(
            client=client,
            model="gemini-test",
            meeting_content=broken_content,
            job_id="quality-job",
        )

        self.assertEqual(result, repaired_content)
        self.assertEqual(len(client.models.calls), 1)
        self.assertEqual(_meeting_content_quality_issues(result), [])

    def test_golden_meeting_markdown_keeps_required_structure(self):
        from backend.tasks import _extract_summary_preview, _meeting_content_quality_issues

        content = (ROOT / "tests" / "fixtures" / "golden_meeting.md").read_text(encoding="utf-8")

        self.assertEqual(_meeting_content_quality_issues(content), [])
        self.assertEqual(_extract_summary_preview(content), "本次會議確認新版儀器驗證排程，並分派文件與測試準備工作。")


class DurableQueueRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database

    def test_jobs_table_has_durable_queue_columns(self):
        database = self._isolated_database()

        with database.get_db() as conn:
            columns = {row["name"] for row in conn.execute("PRAGMA table_info(jobs)").fetchall()}

        self.assertTrue(
            {
                "task_type",
                "source",
                "payload_json",
                "attempts",
                "max_attempts",
                "queued_at",
                "started_at",
                "updated_at",
                "cancel_requested",
                "progress_current",
                "progress_total",
            }.issubset(columns)
        )

    def test_startup_bumps_active_legacy_jobs_to_default_attempts(self):
        database = self._isolated_database()
        database.create_job("legacy-low-attempts-job", max_attempts=2)

        database.init_db()

        self.assertEqual(database.get_job("legacy-low-attempts-job")["max_attempts"], 5)

    def test_claim_next_pending_job_persists_payload_and_increments_attempts(self):
        database = self._isolated_database()
        database.create_job(
            "queue-job-1",
            task_type="audio_processing",
            source="upload",
            payload={"audio_path": "/tmp/audio.mp3", "model": "test-model"},
            max_attempts=3,
        )

        claimed = database.claim_next_pending_job()

        self.assertIsNotNone(claimed)
        self.assertEqual(claimed["job_id"], "queue-job-1")
        self.assertEqual(claimed["status"], "processing")
        self.assertEqual(claimed["attempts"], 1)
        self.assertEqual(claimed["payload"]["audio_path"], "/tmp/audio.mp3")
        self.assertEqual(database.get_job("queue-job-1")["status"], "processing")

    def test_failed_attempt_requeues_until_max_attempts_then_fails(self):
        database = self._isolated_database()
        database.create_job("queue-job-2", max_attempts=2)

        database.claim_next_pending_job()
        first_status = database.retry_or_fail_job("queue-job-2", "temporary failure")
        self.assertEqual(first_status, "pending")
        self.assertEqual(database.get_job("queue-job-2")["error_detail"], "temporary failure")

        database.claim_next_pending_job()
        final_status = database.retry_or_fail_job("queue-job-2", "permanent failure")
        job = database.get_job("queue-job-2")

        self.assertEqual(final_status, "failed")
        self.assertEqual(job["status"], "failed")
        self.assertEqual(job["error_detail"], "permanent failure")
        self.assertIsNotNone(job["completed_at"])

    def test_transient_503_failure_is_requeued_after_delay(self):
        database = self._isolated_database()
        database.create_job("queue-job-503", max_attempts=5)

        with mock.patch.dict(os.environ, {"JOB_QUEUE_TRANSIENT_RETRY_DELAY_SECONDS": "60"}, clear=False):
            database.claim_next_pending_job()
            status = database.retry_or_fail_job("queue-job-503", "503 UNAVAILABLE")
            job = database.get_job("queue-job-503")

            self.assertEqual(status, "pending")
            self.assertIn("服務暫時忙碌", job["message"])
            self.assertIsNone(database.claim_next_pending_job())

    def test_interrupted_processing_jobs_are_requeued_on_startup(self):
        database = self._isolated_database()
        database.create_job(
            "queue-job-3",
            payload={"audio_path": "/tmp/audio.mp3", "output_dir": "/tmp"},
            max_attempts=3,
        )
        database.claim_next_pending_job()

        requeued = database.requeue_interrupted_jobs()
        job = database.get_job("queue-job-3")

        self.assertEqual(requeued, 1)
        self.assertEqual(job["status"], "pending")
        self.assertIn("重新排入", job["message"])

    def test_interrupted_legacy_jobs_without_payload_are_not_requeued(self):
        database = self._isolated_database()
        database.create_job("queue-job-legacy", max_attempts=3)
        database.claim_next_pending_job()

        requeued = database.requeue_interrupted_jobs()
        job = database.get_job("queue-job-legacy")

        self.assertEqual(requeued, 0)
        self.assertEqual(job["status"], "failed")
        self.assertIn("缺少可恢復", job["error_detail"])

    def test_cancel_pending_job_prevents_future_claim(self):
        database = self._isolated_database()
        database.create_job("queue-job-4")

        self.assertTrue(database.request_job_cancel("queue-job-4"))
        self.assertIsNone(database.claim_next_pending_job())
        self.assertEqual(database.get_job("queue-job-4")["status"], "cancelled")

    def test_progress_fields_are_persisted_on_status_update(self):
        database = self._isolated_database()
        database.create_job("queue-job-5")

        database.update_job_status(
            "queue-job-5",
            "processing",
            "處理到一半",
            progress_current=1,
            progress_total=2,
        )
        job = database.get_job("queue-job-5")

        self.assertEqual(job["progress_current"], 1)
        self.assertEqual(job["progress_total"], 2)
        self.assertIsNotNone(job["updated_at"])


class UploadQueueRegressionTests(unittest.TestCase):
    def test_upload_route_enqueues_audio_job_instead_of_inline_background_task(self):
        source = (ROOT / "backend" / "main.py").read_text(encoding="utf-8")
        upload_body = source[
            source.index("async def upload_media") :
            source.index("# =============================================================================\n# 任務狀態查詢端點")
        ]

        self.assertIn("enqueue_audio_job", upload_body)
        self.assertNotIn("background_tasks.add_task", upload_body)

    def test_upload_route_saves_original_audio_to_source_audio_dir(self):
        import backend.main as main

        captured = {}
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            source_audio_dir = tmpdir_path / "source_audio"
            output_dir = tmpdir_path / "output"
            source_audio_dir.mkdir()
            output_dir.mkdir()

            def fake_enqueue_audio_job(**kwargs):
                captured.update(kwargs)

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_audio_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "enqueue_audio_job", side_effect=fake_enqueue_audio_job):
                response = asgi_request(
                    main.app,
                    "POST",
                    "/upload-media",
                    files={"file": ("meeting.mp3", BytesIO(b"ID3" + b"\0" * 32), "audio/mpeg")},
                    data={"recording_profile": "audio_compact"},
                )

            self.assertEqual(response.status_code, 202)
            self.assertIn("媒體檔已接收", response.json()["message"])
            saved_audio = list(source_audio_dir.glob("*.mp3"))
            self.assertEqual(len(saved_audio), 1)
            self.assertEqual(captured["audio_path"], saved_audio[0])
            self.assertEqual(captured["output_dir"], output_dir)
            self.assertEqual(captured["recording_profile"], "audio_compact")
            self.assertTrue(saved_audio[0].read_bytes().startswith(b"ID3"))

    def test_upload_route_reuses_existing_audio_with_same_sha256(self):
        import backend.main as main

        captured = {}
        media_bytes = b"ID3" + b"\0" * 32
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            source_audio_dir = tmpdir_path / "source_audio"
            output_dir = tmpdir_path / "output"
            source_audio_dir.mkdir()
            output_dir.mkdir()
            existing_audio = source_audio_dir / "existing-source.mp3"
            existing_audio.write_bytes(media_bytes)

            def fake_enqueue_audio_job(**kwargs):
                captured.update(kwargs)

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_audio_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "enqueue_audio_job", side_effect=fake_enqueue_audio_job):
                response = asgi_request(
                    main.app,
                    "POST",
                    "/upload-media",
                    files={"file": ("meeting.mp3", BytesIO(media_bytes), "audio/mpeg")},
                )

            self.assertEqual(response.status_code, 202)
            self.assertEqual(captured["audio_path"], existing_audio)
            self.assertEqual(list(source_audio_dir.glob("*.mp3")), [existing_audio])
            self.assertFalse(list(source_audio_dir.glob(".upload_*")))

    def test_upload_route_keeps_reused_audio_when_enqueue_fails(self):
        import backend.main as main

        media_bytes = b"ID3" + b"\0" * 32
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            source_audio_dir = tmpdir_path / "source_audio"
            output_dir = tmpdir_path / "output"
            source_audio_dir.mkdir()
            output_dir.mkdir()
            existing_audio = source_audio_dir / "existing-source.mp3"
            existing_audio.write_bytes(media_bytes)

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_audio_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "enqueue_audio_job", side_effect=RuntimeError("queue down")):
                response = asgi_request(
                    main.app,
                    "POST",
                    "/upload-media",
                    files={"file": ("meeting.mp3", BytesIO(media_bytes), "audio/mpeg")},
                )

            self.assertEqual(response.status_code, 500)
            self.assertTrue(existing_audio.exists())
            self.assertEqual(list(source_audio_dir.glob("*.mp3")), [existing_audio])
            self.assertFalse(list(source_audio_dir.glob(".upload_*")))

    def test_legacy_upload_audio_route_remains_supported(self):
        import backend.main as main

        captured = {}
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            source_audio_dir = tmpdir_path / "source_audio"
            output_dir = tmpdir_path / "output"
            source_audio_dir.mkdir()
            output_dir.mkdir()

            def fake_enqueue_audio_job(**kwargs):
                captured.update(kwargs)

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_audio_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "enqueue_audio_job", side_effect=fake_enqueue_audio_job):
                response = asgi_request(
                    main.app,
                    "POST",
                    "/upload-audio",
                    files={"file": ("meeting.mp3", BytesIO(b"ID3" + b"\0" * 32), "audio/mpeg")},
                )

            self.assertEqual(response.status_code, 202)
            self.assertIn("媒體檔已接收", response.json()["message"])
            self.assertTrue(captured["audio_path"].is_file())


class TempCleanupRegressionTests(unittest.TestCase):
    def test_stale_temp_cleanup_preserves_active_and_fresh_files(self):
        from backend.cleanup import cleanup_stale_temp_files

        with tempfile.TemporaryDirectory() as tmpdir:
            temp_dir = Path(tmpdir)
            stale = temp_dir / "old-unreferenced.wav"
            active = temp_dir / "active-retry.wav"
            fresh = temp_dir / "fresh.wav"
            nested_segment = temp_dir / "_seg_old_000.mp3"

            for path in (stale, active, fresh, nested_segment):
                path.write_bytes(b"audio")

            old_time = 1_700_000_000
            now = old_time + 7200
            for path in (stale, active, nested_segment):
                path.touch()
                import os
                os.utime(path, (old_time, old_time))
            import os
            os.utime(fresh, (now, now))

            deleted = cleanup_stale_temp_files(
                temp_dir=temp_dir,
                active_paths={active},
                max_age_seconds=3600,
                now=now,
            )

            self.assertEqual({path.name for path in deleted}, {"old-unreferenced.wav", "_seg_old_000.mp3"})
            self.assertFalse(stale.exists())
            self.assertTrue(active.exists())
            self.assertTrue(fresh.exists())

    def test_lifespan_runs_temp_cleanup_before_worker_start(self):
        source = (ROOT / "backend" / "main.py").read_text(encoding="utf-8")
        lifespan_body = source[
            source.index("async def lifespan") :
            source.index("app = FastAPI(")
        ]

        self.assertIn("cleanup_stale_temp_files_for_jobs", lifespan_body)
        self.assertLess(
            lifespan_body.index("cleanup_stale_temp_files_for_jobs"),
            lifespan_body.index("job_worker.start()"),
        )


class JobRetentionCleanupRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database

    def test_terminal_job_cleanup_deletes_only_jobs_older_than_retention(self):
        database = self._isolated_database()
        from backend.cleanup import cleanup_terminal_jobs

        now = datetime(2026, 1, 15, 12, 0, 0)
        old_completed_at = (now - timedelta(days=10)).strftime("%Y-%m-%d %H:%M:%S")
        recent_completed_at = (now - timedelta(days=2)).strftime("%Y-%m-%d %H:%M:%S")

        database.create_job("old-failed")
        database.update_job_status("old-failed", "failed", "失敗")
        database.create_job("recent-done")
        database.update_job_status("recent-done", "done", "完成")
        database.create_job("old-pending")

        with database.get_db() as conn:
            conn.execute(
                "UPDATE jobs SET completed_at=?, updated_at=? WHERE job_id='old-failed'",
                (old_completed_at, old_completed_at),
            )
            conn.execute(
                "UPDATE jobs SET completed_at=?, updated_at=? WHERE job_id='recent-done'",
                (recent_completed_at, recent_completed_at),
            )
            conn.execute(
                "UPDATE jobs SET queued_at=?, updated_at=? WHERE job_id='old-pending'",
                (old_completed_at, old_completed_at),
            )

        deleted = cleanup_terminal_jobs(max_age_days=7, now=now)

        self.assertEqual(deleted, 1)
        self.assertIsNone(database.get_job("old-failed"))
        self.assertIsNotNone(database.get_job("recent-done"))
        self.assertIsNotNone(database.get_job("old-pending"))

    def test_lifespan_runs_terminal_job_cleanup_before_worker_start(self):
        source = (ROOT / "backend" / "main.py").read_text(encoding="utf-8")
        lifespan_body = source[
            source.index("async def lifespan") :
            source.index("app = FastAPI(")
        ]

        self.assertIn("cleanup_terminal_jobs", lifespan_body)
        self.assertLess(
            lifespan_body.index("cleanup_terminal_jobs"),
            lifespan_body.index("job_worker.start()"),
        )


class JobDashboardRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database

    def test_jobs_endpoint_lists_jobs_and_filters_by_status(self):
        database = self._isolated_database()
        import backend.main as main

        database.create_job(
            "job-dashboard-pending",
            source="upload",
            payload={"audio_path": "/tmp/pending.wav", "output_dir": "/tmp"},
        )
        database.create_job(
            "job-dashboard-failed",
            source="line",
            payload={"message_id": "line-message", "user_id": "user-id"},
        )
        database.update_job_status(
            "job-dashboard-failed",
            "failed",
            "處理失敗",
            error_detail="測試錯誤",
        )

        response = asgi_request(main.app, "GET", "/jobs?status=failed&limit=10")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["total"], 1)
        self.assertEqual(payload["jobs"][0]["job_id"], "job-dashboard-failed")
        self.assertEqual(payload["jobs"][0]["status"], "failed")
        self.assertEqual(payload["jobs"][0]["source"], "line")
        self.assertEqual(payload["jobs"][0]["error_detail"], "測試錯誤")

    def test_failed_job_can_be_requeued_by_api(self):
        database = self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            audio_path = tmp_path / "retry.mp3"
            audio_path.write_bytes(b"ID3\x04\x00\x00\x00\x00\x00\x21audio")
            database.create_job(
                "job-retry-api",
                source="upload",
                payload={"audio_path": str(audio_path), "output_dir": str(tmp_path)},
            )
            database.update_job_status("job-retry-api", "failed", "處理失敗", error_detail="暫時錯誤")

            response = asgi_request(main.app, "POST", "/jobs/job-retry-api/retry")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "pending")
        self.assertEqual(payload["attempts"], 0)
        self.assertIsNone(payload["error_detail"])

    def test_retry_rejects_missing_audio_payload_file(self):
        database = self._isolated_database()
        import backend.main as main

        missing_audio = Path(tempfile.gettempdir()) / "meeting-assistant-missing-retry-source.mp3"
        if missing_audio.exists():
            missing_audio.unlink()
        database.create_job(
            "job-retry-missing-audio",
            source="upload",
            payload={"audio_path": str(missing_audio), "output_dir": tempfile.gettempdir()},
        )
        database.update_job_status("job-retry-missing-audio", "failed", "處理失敗", error_detail="原檔遺失")

        response = asgi_request(main.app, "POST", "/jobs/job-retry-missing-audio/retry")

        self.assertEqual(response.status_code, 409)
        self.assertEqual(database.get_job("job-retry-missing-audio")["status"], "failed")

    def test_job_events_endpoint_returns_status_timeline(self):
        database = self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            audio_path = tmp_path / "timeline.mp3"
            audio_path.write_bytes(b"ID3\x04\x00\x00\x00\x00\x00\x21audio")
            database.create_job(
                "job-events-api",
                payload={"audio_path": str(audio_path), "output_dir": str(tmp_path)},
            )
            database.update_job_status("job-events-api", "failed", "處理失敗", error_detail="temporary")
            database.requeue_failed_job("job-events-api")

            response = asgi_request(main.app, "GET", "/jobs/job-events-api/events")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["job_id"], "job-events-api")
        self.assertEqual(
            [event["event_type"] for event in payload["events"]],
            ["created", "status_failed", "requeued"],
        )
        self.assertEqual(payload["events"][1]["detail"], "temporary")

    def test_terminal_job_can_be_deleted_by_api(self):
        database = self._isolated_database()
        import backend.main as main

        database.create_job("job-delete-api")
        database.update_job_status("job-delete-api", "cancelled", "已取消")

        response = asgi_request(main.app, "DELETE", "/jobs/job-delete-api")

        self.assertEqual(response.status_code, 200)
        self.assertIsNone(database.get_job("job-delete-api"))

    def test_processing_job_cannot_be_deleted_by_api(self):
        database = self._isolated_database()
        import backend.main as main

        database.create_job("job-delete-processing", payload={"audio_path": "/tmp/a.mp3", "output_dir": "/tmp"})
        database.claim_next_pending_job()

        response = asgi_request(main.app, "DELETE", "/jobs/job-delete-processing")

        self.assertEqual(response.status_code, 409)
        self.assertIsNotNone(database.get_job("job-delete-processing"))


class MeetingRerunRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database

    def test_meeting_can_be_rerun_from_retained_source_audio(self):
        database = self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            output_dir = root / "output"
            source_dir.mkdir()
            output_dir.mkdir()
            audio_path = source_dir / "kept-source.mp3"
            audio_path.write_bytes(b"ID3\x04\x00\x00\x00\x00\x00\x21audio")
            output_path = output_dir / "meeting.md"
            output_path.write_text("# meeting", encoding="utf-8")
            meeting_id = database.save_meeting(
                title="重跑測試會議",
                date="2026/07/08",
                source_audio=audio_path.name,
                output_path=str(output_path),
                summary="摘要",
            )

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir):
                response = asgi_request(main.app, "POST", f"/meetings/{meeting_id}/rerun")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "pending")
        self.assertIn("原始媒體檔", payload["message"])
        job = database.get_job(payload["job_id"])
        self.assertIsNotNone(job)
        self.assertEqual(job["source"], "meeting_rerun")
        self.assertEqual(job["payload"]["audio_path"], str(audio_path))
        self.assertEqual(job["payload"]["output_dir"], str(output_dir))
        self.assertEqual(job["payload"]["meeting_title"], "重跑測試會議")

    def test_meeting_rerun_rejects_missing_source_audio(self):
        database = self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            output_dir = root / "output"
            source_dir.mkdir()
            output_dir.mkdir()
            output_path = output_dir / "meeting.md"
            output_path.write_text("# meeting", encoding="utf-8")
            meeting_id = database.save_meeting(
                title="缺檔會議",
                date="2026/07/08",
                source_audio="missing-source.mp3",
                output_path=str(output_path),
                summary="摘要",
            )

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir):
                response = asgi_request(main.app, "POST", f"/meetings/{meeting_id}/rerun")

        self.assertEqual(response.status_code, 409)
        self.assertIn("原始媒體檔", response.json()["detail"])
        self.assertEqual(database.count_jobs(), 0)


class MetricsRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database

    def test_metrics_endpoint_reports_job_counts_and_recent_errors(self):
        database = self._isolated_database()
        import backend.main as main

        database.create_job("metrics-pending")
        database.create_job("metrics-failed")
        database.update_job_status("metrics-failed", "failed", "處理失敗", error_detail="metrics error")
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "needs-review.md"
            output_path.write_text("metrics needs review", encoding="utf-8")
            database.save_meeting(
                title="Metrics Review",
                date="2026/07/12",
                source_audio="metrics-review.webm",
                output_path=str(output_path),
                summary="metrics needs review",
                quality_report={"score": 91, "label": "ok", "warnings": ["review this"]},
            )

            response = asgi_request(main.app, "GET", "/metrics")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["jobs"]["total"], 2)
        self.assertEqual(payload["jobs"]["by_status"]["pending"], 1)
        self.assertEqual(payload["jobs"]["by_status"]["failed"], 1)
        self.assertEqual(payload["meetings"]["total"], 1)
        self.assertEqual(payload["meetings"]["needs_review"], 1)
        self.assertEqual(payload["recent_errors"][0]["job_id"], "metrics-failed")
        self.assertEqual(payload["recent_errors"][0]["error_detail"], "metrics error")
        self.assertIn("storage", payload)

    def test_metrics_endpoint_reports_storage_usage(self):
        database = self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            output_dir = root / "output"
            source_dir.mkdir()
            output_dir.mkdir()
            (source_dir / "meeting-a.mp3").write_bytes(b"audio-a")
            (source_dir / "meeting-b.webm").write_bytes(b"video-bb")
            (source_dir / "meeting-c.m4a").write_bytes(b"ccc")
            (source_dir / ".upload_tmp.mp3").write_bytes(b"ignore")
            (source_dir / "note.txt").write_text("ignore", encoding="utf-8")
            (output_dir / "meeting-a.md").write_text("note-a", encoding="utf-8")
            (output_dir / "meeting-b.md").write_text("note-bb", encoding="utf-8")
            (output_dir / "meetings.db").write_bytes(b"ignore-db")
            audio_meeting_id = database.save_meeting(
                title="Audio A",
                date="2026/07/13",
                source_audio=str(source_dir / "meeting-a.mp3"),
                output_path=str(output_dir / "meeting-a.md"),
                summary="audio summary",
            )
            video_meeting_id = database.save_meeting(
                title="Video B",
                date="2026/07/13",
                source_audio="meeting-b.webm",
                output_path=str(output_dir / "meeting-b.md"),
                summary="video summary",
                quality_report={"recording": {"profile": "video_balanced"}},
            )

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "BACKUP_DIR", root / "backups"):
                response = asgi_request(main.app, "GET", "/metrics")
                inventory_response = asgi_request(main.app, "GET", "/source-media/inventory?limit=2")
                inventory_page_response = asgi_request(main.app, "GET", "/source-media/inventory?limit=1&offset=1")
                inventory_video_response = asgi_request(main.app, "GET", "/source-media/inventory/meeting-b.webm")
                inventory_media_response = asgi_request(main.app, "GET", "/source-media/inventory/meeting-c.m4a")
                inventory_download_response = asgi_request(main.app, "GET", "/source-media/inventory/meeting-c.m4a?download=1")
                hidden_media_response = asgi_request(main.app, "GET", "/source-media/inventory/.upload_tmp.mp3")
                linked_delete_response = asgi_request(main.app, "DELETE", "/source-media/inventory/meeting-b.webm")
                hidden_delete_response = asgi_request(main.app, "DELETE", "/source-media/inventory/.upload_tmp.mp3")
                delete_response = asgi_request(main.app, "DELETE", "/source-media/inventory/meeting-c.m4a")
                post_delete_inventory_response = asgi_request(main.app, "GET", "/source-media/inventory?limit=10")
                post_delete_metrics_response = asgi_request(main.app, "GET", "/metrics")
                deleted_file_exists = (source_dir / "meeting-c.m4a").exists()
                deleted_backup_path = Path(delete_response.json().get("backup_path", ""))
                deleted_backup_exists = deleted_backup_path.is_file()
                deleted_backup_content = deleted_backup_path.read_bytes() if deleted_backup_exists else b""
                deleted_backup_metadata_path = main._source_media_archive_metadata_path(deleted_backup_path)
                deleted_backup_metadata_exists = deleted_backup_metadata_path.is_file()
                deleted_backup_metadata = (
                    json.loads(deleted_backup_metadata_path.read_text(encoding="utf-8"))
                    if deleted_backup_metadata_exists
                    else {}
                )
                deleted_backup_metadata_bytes = (
                    deleted_backup_metadata_path.stat().st_size
                    if deleted_backup_metadata_exists
                    else 0
                )
                archive_response = asgi_request(main.app, "GET", "/source-media/archive?limit=10")
                restore_archive_id = archive_response.json()["files"][0]["archive_id"]
                archive_file_response = asgi_request(
                    main.app,
                    "GET",
                    "/source-media/archive/file",
                    params={"archive_id": restore_archive_id},
                )
                archive_download_response = asgi_request(
                    main.app,
                    "GET",
                    "/source-media/archive/file",
                    params={"archive_id": restore_archive_id, "download": "1"},
                )
                restore_response = asgi_request(
                    main.app,
                    "POST",
                    "/source-media/archive/restore",
                    params={"archive_id": restore_archive_id},
                )
                post_restore_inventory_response = asgi_request(main.app, "GET", "/source-media/inventory?limit=10")
                post_restore_metrics_response = asgi_request(main.app, "GET", "/metrics")
                restored_file_exists = (source_dir / "meeting-c.m4a").is_file()
                restored_backup_exists = deleted_backup_path.exists()
                restored_backup_metadata_exists = deleted_backup_metadata_path.exists()

        self.assertEqual(response.status_code, 200)
        storage = response.json()["storage"]
        self.assertEqual(storage["source_media_files"], 3)
        self.assertEqual(storage["source_media_bytes"], len(b"audio-a") + len(b"video-bb") + len(b"ccc"))
        self.assertEqual(storage["source_media_unlinked_files"], 1)
        self.assertEqual(storage["source_media_unlinked_bytes"], len(b"ccc"))
        self.assertEqual(storage["source_media_archived_files"], 0)
        self.assertEqual(storage["source_media_archived_bytes"], 0)
        self.assertEqual(
            [item["name"] for item in storage["source_media_largest_files"][:3]],
            ["meeting-b.webm", "meeting-a.mp3", "meeting-c.m4a"],
        )
        self.assertEqual(storage["source_media_largest_files"][0]["bytes"], len(b"video-bb"))
        self.assertEqual(storage["source_media_largest_files"][0]["source_media_type"], "video")
        self.assertEqual(storage["source_media_largest_files"][0]["linked_meeting_id"], video_meeting_id)
        self.assertEqual(storage["source_media_largest_files"][0]["linked_meeting_title"], "Video B")
        self.assertEqual(storage["source_media_largest_files"][1]["source_media_type"], "audio")
        self.assertEqual(storage["source_media_largest_files"][1]["linked_meeting_id"], audio_meeting_id)
        self.assertIsNone(storage["source_media_largest_files"][2]["linked_meeting_id"])
        self.assertEqual(storage["meeting_markdown_files"], 2)
        self.assertEqual(storage["meeting_markdown_bytes"], len("note-a".encode("utf-8")) + len("note-bb".encode("utf-8")))
        self.assertEqual(inventory_response.status_code, 200)
        inventory = inventory_response.json()
        self.assertEqual(inventory["total_files"], 3)
        self.assertEqual(inventory["limit"], 2)
        self.assertEqual(inventory["offset"], 0)
        self.assertEqual(inventory["total_bytes"], len(b"audio-a") + len(b"video-bb") + len(b"ccc"))
        self.assertEqual(inventory["unlinked_files"], 1)
        self.assertEqual(inventory["unlinked_bytes"], len(b"ccc"))
        self.assertEqual([item["name"] for item in inventory["files"]], ["meeting-b.webm", "meeting-a.mp3"])
        self.assertEqual(inventory["files"][0]["source_media_type"], "video")
        self.assertEqual(inventory["files"][1]["source_media_type"], "audio")
        self.assertEqual(inventory["files"][0]["linked_meeting_id"], video_meeting_id)
        self.assertEqual(inventory["files"][1]["linked_meeting_title"], "Audio A")
        self.assertEqual(inventory_page_response.status_code, 200)
        inventory_page = inventory_page_response.json()
        self.assertEqual(inventory_page["total_files"], 3)
        self.assertEqual(inventory_page["limit"], 1)
        self.assertEqual(inventory_page["offset"], 1)
        self.assertEqual([item["name"] for item in inventory_page["files"]], ["meeting-a.mp3"])
        self.assertEqual(inventory_video_response.status_code, 200)
        self.assertIn("video/webm", inventory_video_response.headers.get("content-type", ""))
        self.assertIn("inline", inventory_video_response.headers.get("content-disposition", ""))
        self.assertEqual(inventory_video_response.content, b"video-bb")
        self.assertEqual(inventory_media_response.status_code, 200)
        self.assertIn("audio/mp4", inventory_media_response.headers.get("content-type", ""))
        self.assertIn("inline", inventory_media_response.headers.get("content-disposition", ""))
        self.assertEqual(inventory_media_response.content, b"ccc")
        self.assertEqual(inventory_download_response.status_code, 200)
        self.assertIn("attachment", inventory_download_response.headers.get("content-disposition", ""))
        self.assertEqual(hidden_media_response.status_code, 400)
        self.assertEqual(linked_delete_response.status_code, 409)
        self.assertIn("仍連結到會議", linked_delete_response.json()["detail"])
        self.assertEqual(hidden_delete_response.status_code, 400)
        self.assertEqual(delete_response.status_code, 200)
        self.assertEqual(delete_response.json()["name"], "meeting-c.m4a")
        self.assertIn("source_media_deleted", delete_response.json()["backup_path"])
        self.assertTrue(deleted_backup_exists)
        self.assertEqual(deleted_backup_content, b"ccc")
        self.assertTrue(deleted_backup_metadata_exists)
        self.assertEqual(deleted_backup_metadata["original_name"], "meeting-c.m4a")
        self.assertEqual(deleted_backup_metadata["source_media_type"], "audio")
        self.assertEqual(deleted_backup_metadata["bytes"], len(b"ccc"))
        self.assertFalse(deleted_file_exists)
        post_delete_inventory = post_delete_inventory_response.json()
        self.assertEqual(post_delete_inventory["total_files"], 2)
        self.assertEqual(post_delete_inventory["unlinked_files"], 0)
        self.assertEqual([item["name"] for item in post_delete_inventory["files"]], ["meeting-b.webm", "meeting-a.mp3"])
        post_delete_storage = post_delete_metrics_response.json()["storage"]
        self.assertEqual(post_delete_storage["source_media_archived_files"], 1)
        self.assertEqual(post_delete_storage["source_media_archived_bytes"], len(b"ccc") + deleted_backup_metadata_bytes)
        self.assertEqual(archive_response.status_code, 200)
        archive_payload = archive_response.json()
        self.assertEqual(archive_payload["total_files"], 1)
        self.assertEqual(archive_payload["total_bytes"], len(b"ccc") + deleted_backup_metadata_bytes)
        self.assertEqual(archive_payload["limit"], 10)
        self.assertEqual(archive_payload["offset"], 0)
        self.assertEqual(archive_payload["files"][0]["name"], "meeting-c.m4a")
        self.assertEqual(archive_payload["files"][0]["bytes"], len(b"ccc"))
        self.assertEqual(archive_payload["files"][0]["metadata_bytes"], deleted_backup_metadata_bytes)
        self.assertEqual(archive_payload["files"][0]["source_media_type"], "audio")
        self.assertEqual(archive_file_response.status_code, 200)
        self.assertIn("audio/mp4", archive_file_response.headers.get("content-type", ""))
        self.assertIn("inline", archive_file_response.headers.get("content-disposition", ""))
        self.assertEqual(archive_file_response.content, b"ccc")
        self.assertEqual(archive_download_response.status_code, 200)
        self.assertIn("attachment", archive_download_response.headers.get("content-disposition", ""))
        self.assertEqual(archive_download_response.content, b"ccc")
        self.assertEqual(restore_response.status_code, 200)
        self.assertEqual(restore_response.json()["name"], "meeting-c.m4a")
        self.assertTrue(restored_file_exists)
        self.assertFalse(restored_backup_exists)
        self.assertFalse(restored_backup_metadata_exists)
        post_restore_inventory = post_restore_inventory_response.json()
        self.assertEqual(post_restore_inventory["total_files"], 3)
        self.assertEqual(post_restore_inventory["unlinked_files"], 1)
        post_restore_storage = post_restore_metrics_response.json()["storage"]
        self.assertEqual(post_restore_storage["source_media_archived_files"], 0)
        self.assertEqual(post_restore_storage["source_media_archived_bytes"], 0)

    def test_source_media_archive_preserves_webm_video_metadata_for_preview(self):
        self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            source_dir.mkdir()
            (source_dir / "screen.webm").write_bytes(b"webm-video")

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "BACKUP_DIR", root / "backups"):
                with mock.patch.object(main, "_storage_source_media_type", return_value="video"):
                    delete_response = asgi_request(main.app, "DELETE", "/source-media/inventory/screen.webm")
                backup_path = Path(delete_response.json().get("backup_path", ""))
                metadata_path = main._source_media_archive_metadata_path(backup_path)
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
                metadata_bytes = metadata_path.stat().st_size

                with mock.patch.object(main, "_storage_source_media_type", return_value="audio"), \
                     mock.patch.object(main, "_ffprobe_stream_types", return_value=set()):
                    archive_response = asgi_request(main.app, "GET", "/source-media/archive?limit=10")
                    archive_id = archive_response.json()["files"][0]["archive_id"]
                    archive_file_response = asgi_request(
                        main.app,
                        "GET",
                        "/source-media/archive/file",
                        params={"archive_id": archive_id},
                    )

        self.assertEqual(delete_response.status_code, 200)
        self.assertEqual(metadata["original_name"], "screen.webm")
        self.assertEqual(metadata["source_media_type"], "video")
        self.assertEqual(archive_response.status_code, 200)
        archive_payload = archive_response.json()
        self.assertEqual(archive_payload["total_bytes"], len(b"webm-video") + metadata_bytes)
        self.assertEqual(archive_payload["files"][0]["source_media_type"], "video")
        self.assertEqual(archive_payload["files"][0]["metadata_bytes"], metadata_bytes)
        self.assertEqual(archive_file_response.status_code, 200)
        self.assertIn("video/webm", archive_file_response.headers.get("content-type", ""))
        self.assertEqual(archive_file_response.content, b"webm-video")

    def test_source_media_archive_counts_orphan_metadata_storage(self):
        self._isolated_database()
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            source_dir = root / "source_audio"
            output_dir = root / "output"
            archive_dir = root / "backups" / "source_media_deleted" / "20260713"
            source_dir.mkdir()
            output_dir.mkdir()
            archive_dir.mkdir(parents=True)
            orphan_metadata = archive_dir / "171825_042088_screen.webm.json"
            orphan_metadata.write_text('{"source_media_type":"video"}', encoding="utf-8")
            orphan_metadata_bytes = orphan_metadata.stat().st_size
            ignored_metadata = archive_dir / "manual-note.txt.json"
            ignored_metadata.write_text("not a source media sidecar", encoding="utf-8")

            with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
                 mock.patch.object(main, "OUTPUT_DIR", output_dir), \
                 mock.patch.object(main, "BACKUP_DIR", root / "backups"):
                metrics_response = asgi_request(main.app, "GET", "/metrics")
                archive_response = asgi_request(main.app, "GET", "/source-media/archive?limit=10")

        self.assertEqual(metrics_response.status_code, 200)
        storage = metrics_response.json()["storage"]
        self.assertEqual(storage["source_media_archived_files"], 0)
        self.assertEqual(storage["source_media_archived_bytes"], orphan_metadata_bytes)
        self.assertEqual(archive_response.status_code, 200)
        archive_payload = archive_response.json()
        self.assertEqual(archive_payload["total_files"], 0)
        self.assertEqual(archive_payload["total_bytes"], orphan_metadata_bytes)
        self.assertEqual(archive_payload["files"], [])

    def test_metrics_endpoint_reports_ngrok_status(self):
        self._isolated_database()
        import backend.main as main

        ngrok_status = {
            "running": True,
            "public_url": "https://example.ngrok-free.app",
            "webhook_url": "https://example.ngrok-free.app/line-webhook",
            "message": "ngrok tunnel is forwarding to local server",
        }

        with mock.patch.object(main, "get_ngrok_status", return_value=ngrok_status):
            response = asgi_request(main.app, "GET", "/metrics")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ngrok"]["running"])
        self.assertEqual(payload["ngrok"]["webhook_url"], "https://example.ngrok-free.app/line-webhook")


class SearchRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database, Path(tmpdir.name)

    def test_search_meetings_uses_fts5_index_for_audio_filename(self):
        database, tmp_path = self._isolated_database()
        output_path = tmp_path / "indexed-meeting.md"
        output_path.write_text("完整逐字稿包含 ultrasonic validation", encoding="utf-8")

        database.save_meeting(
            title="Weekly Sync",
            date="2026/07/05",
            source_audio="sourceaudiospecial.wav",
            output_path=str(output_path),
            summary="討論例行事項",
        )

        results = database.search_meetings("sourceaudiospecial")

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["source_audio"], "sourceaudiospecial.wav")
        with database.get_db() as conn:
            table = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='meeting_fts'"
            ).fetchone()
        self.assertIsNotNone(table)

    def test_search_meetings_matches_chinese_substrings_and_full_transcript(self):
        database, tmp_path = self._isolated_database()
        output_path = tmp_path / "content-indexed-meeting.md"
        output_path.write_text(
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "這份 transcriptneedle 只存在於完整逐字稿內容。\n",
            encoding="utf-8",
        )

        meeting_id = database.save_meeting(
            title="不合格品會議20260708",
            date="2026/07/08",
            source_audio="quality-review.webm",
            output_path=str(output_path),
            summary="供應商外觀檢驗結果討論",
        )

        self.assertEqual([meeting_id], [row["id"] for row in database.search_meetings("不合格品")])
        self.assertEqual([meeting_id], [row["id"] for row in database.search_meetings("供應商")])
        self.assertEqual([meeting_id], [row["id"] for row in database.search_meetings("transcriptneedle")])

        with database.get_db() as conn:
            table = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='meeting_content_fts'"
            ).fetchone()
        self.assertIsNotNone(table)

    def test_repeated_searches_do_not_rebuild_or_grow_the_database(self):
        database, tmp_path = self._isolated_database()
        output_path = tmp_path / "stable-search.md"
        output_path.write_text("stable-search-content", encoding="utf-8")
        database.save_meeting(
            title="Stable Search",
            date="2026/07/08",
            source_audio="stable.wav",
            output_path=str(output_path),
            summary="stable-search-summary",
        )

        database.search_meetings("stable")
        before = database.DB_PATH.stat().st_size
        for _ in range(20):
            database.search_meetings("stable")
        after = database.DB_PATH.stat().st_size

        self.assertEqual(before, after)

    def test_list_and_search_include_quality_warning_count(self):
        database, tmp_path = self._isolated_database()
        output_path = tmp_path / "quality-badge.md"
        output_path.write_text("quality-badge-content", encoding="utf-8")
        quality_report = {
            "score": 72,
            "label": "需複核",
            "warnings": ["錄音音量偏低", "摘要未串聯待辦"],
        }
        meeting_id = database.save_meeting(
            title="Quality Badge",
            date="2026/07/12",
            source_audio="quality-badge.webm",
            output_path=str(output_path),
            summary="quality-badge-summary",
            quality_report=quality_report,
        )

        listed = next(row for row in database.list_meetings() if row["id"] == meeting_id)
        searched = database.search_meetings("Quality Badge")[0]

        self.assertEqual(listed["quality_score"], 72)
        self.assertEqual(listed["quality_label"], "需複核")
        self.assertEqual(listed["quality_warning_count"], 2)
        self.assertEqual(listed["quality_warning_preview"], "錄音音量偏低")
        self.assertEqual(searched["quality_score"], 72)
        self.assertEqual(searched["quality_label"], "需複核")
        self.assertEqual(searched["quality_warning_count"], 2)
        self.assertEqual(searched["quality_warning_preview"], "錄音音量偏低")

    def test_list_and_search_include_source_media_type(self):
        database, tmp_path = self._isolated_database()
        video_output = tmp_path / "video.md"
        audio_output = tmp_path / "audio.md"
        mp4_output = tmp_path / "mp4.md"
        legacy_output = tmp_path / "legacy-webm.md"
        video_output.write_text("video meeting", encoding="utf-8")
        audio_output.write_text("audio meeting", encoding="utf-8")
        mp4_output.write_text("mp4 meeting", encoding="utf-8")
        legacy_output.write_text("legacy webm meeting", encoding="utf-8")

        video_id = database.save_meeting(
            title="Video Meeting",
            date="2026/07/12",
            source_audio="recorded-screen.webm",
            output_path=str(video_output),
            summary="video meeting",
            quality_report={"recording": {"profile": "video_balanced"}},
        )
        audio_id = database.save_meeting(
            title="Audio Meeting",
            date="2026/07/12",
            source_audio="recorded-audio.webm",
            output_path=str(audio_output),
            summary="audio meeting",
            quality_report={"recording": {"profile": "audio_standard"}},
        )
        mp4_id = database.save_meeting(
            title="Uploaded MP4",
            date="2026/07/12",
            source_audio="uploaded.mp4",
            output_path=str(mp4_output),
            summary="mp4 meeting",
        )
        legacy_webm_id = database.save_meeting(
            title="Legacy WebM",
            date="2026/07/12",
            source_audio="legacy.webm",
            output_path=str(legacy_output),
            summary="legacy webm meeting",
        )

        listed_by_id = {row["id"]: row for row in database.list_meetings(limit=10)}

        self.assertEqual(listed_by_id[video_id]["source_media_type"], "video")
        self.assertEqual(listed_by_id[audio_id]["source_media_type"], "audio")
        self.assertEqual(listed_by_id[mp4_id]["source_media_type"], "video")
        self.assertIsNone(listed_by_id[legacy_webm_id]["source_media_type"])
        self.assertEqual(database.search_meetings("Video Meeting")[0]["source_media_type"], "video")
        self.assertEqual(database.search_meetings("Audio Meeting")[0]["source_media_type"], "audio")

        import backend.main as main

        source_dir = tmp_path / "source_audio"
        source_dir.mkdir()
        (source_dir / "legacy.webm").write_bytes(b"\x1a\x45\xdf\xa3")
        with mock.patch.object(main, "SOURCE_AUDIO_DIR", source_dir), \
             mock.patch.object(main, "_ffprobe_stream_types", return_value={"video", "audio"}):
            list_response = asgi_request(main.app, "GET", "/meetings?limit=10")
            legacy_search_response = asgi_request(main.app, "GET", "/meetings/search?q=Legacy%20WebM&limit=5")
        search_response = asgi_request(main.app, "GET", "/meetings/search?q=Video%20Meeting&limit=5")

        self.assertEqual(list_response.status_code, 200)
        api_by_id = {row["id"]: row for row in list_response.json()["records"]}
        self.assertEqual(api_by_id[video_id]["source_media_type"], "video")
        self.assertEqual(api_by_id[audio_id]["source_media_type"], "audio")
        self.assertEqual(api_by_id[legacy_webm_id]["source_media_type"], "video")
        self.assertEqual(search_response.status_code, 200)
        self.assertEqual(search_response.json()[0]["source_media_type"], "video")
        self.assertEqual(legacy_search_response.status_code, 200)
        self.assertEqual(legacy_search_response.json()[0]["source_media_type"], "video")

    def test_legacy_meeting_list_infers_quality_warning_count_from_markdown(self):
        database, tmp_path = self._isolated_database()
        output_path = tmp_path / "legacy-warning.md"
        repeated_turns = "\n".join(
            f"[00:0{index}] **[發言者 A]**：這一句不應該連續重複。"
            for index in range(4)
        )
        output_path.write_text(
            "## 一、討論摘要 (Discussion Summary)\n"
            "這是一段沒有 D 編號的摘要。\n"
            "## 二、最終決議 (Final Decisions)\n"
            "決議沒有 R 編號。\n"
            "## 三、待辦事項 (Action Items)\n"
            "| # | 任務描述 | 負責人 | 期限 | 優先級 |\n"
            "|---|---|---|---|---|\n"
            "| - | 整理追蹤表 | 發言者 A | 未提及 | 中 |\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "*(註：為節省篇幅，已省略逐字稿中重複內容)*\n"
            f"{repeated_turns}\n",
            encoding="utf-8",
        )
        meeting_id = database.save_meeting(
            title="Legacy Warning",
            date="2026/07/12",
            source_audio="legacy-warning.webm",
            output_path=str(output_path),
            summary="legacy-warning-summary",
        )

        listed = next(row for row in database.list_meetings() if row["id"] == meeting_id)
        searched = database.search_meetings("Legacy Warning")[0]

        self.assertIsNone(listed["quality_score"])
        self.assertIsNone(listed["quality_label"])
        self.assertGreaterEqual(listed["quality_warning_count"], 5)
        self.assertIn("舊紀錄需複核", listed["quality_warning_preview"])
        self.assertEqual(searched["quality_warning_count"], listed["quality_warning_count"])
        self.assertEqual(searched["quality_warning_preview"], listed["quality_warning_preview"])

    def test_needs_review_filter_scans_beyond_regular_list_limit(self):
        database, tmp_path = self._isolated_database()
        review_output = tmp_path / "older-review.md"
        normal_output = tmp_path / "newer-normal.md"
        review_output.write_text("shared-needs-review", encoding="utf-8")
        normal_output.write_text("shared-needs-review", encoding="utf-8")

        review_id = database.save_meeting(
            title="Older Review",
            date="2026/07/12",
            source_audio="older-review.webm",
            output_path=str(review_output),
            summary="shared-needs-review older",
            quality_report={"score": 95, "label": "ok", "warnings": ["check transcript"]},
        )
        normal_id = database.save_meeting(
            title="Newer Normal",
            date="2026/07/12",
            source_audio="newer-normal.webm",
            output_path=str(normal_output),
            summary="shared-needs-review newer",
            quality_report={"score": 98, "label": "ok", "warnings": []},
        )
        with database.get_db() as conn:
            conn.execute("UPDATE meetings SET created_at=? WHERE id=?", ("2026-07-12 08:00:00", review_id))
            conn.execute("UPDATE meetings SET created_at=? WHERE id=?", ("2026-07-12 09:00:00", normal_id))

        self.assertEqual(database.list_meetings(limit=1)[0]["id"], normal_id)

        filtered = database.list_meetings(limit=1, needs_review=True)

        self.assertEqual([row["id"] for row in filtered], [review_id])
        self.assertEqual(database.count_meetings(needs_review=True), 1)

    def test_search_and_api_can_filter_needs_review(self):
        database, tmp_path = self._isolated_database()
        review_output = tmp_path / "search-review.md"
        normal_output = tmp_path / "search-normal.md"
        review_output.write_text("shared-search-keyword", encoding="utf-8")
        normal_output.write_text("shared-search-keyword", encoding="utf-8")

        review_id = database.save_meeting(
            title="Search Review",
            date="2026/07/12",
            source_audio="search-review.webm",
            output_path=str(review_output),
            summary="shared-search-keyword older",
            quality_report={"score": 80, "label": "needs review", "warnings": []},
        )
        normal_id = database.save_meeting(
            title="Search Normal",
            date="2026/07/12",
            source_audio="search-normal.webm",
            output_path=str(normal_output),
            summary="shared-search-keyword newer",
            quality_report={"score": 98, "label": "ok", "warnings": []},
        )
        with database.get_db() as conn:
            conn.execute("UPDATE meetings SET created_at=? WHERE id=?", ("2026-07-12 08:00:00", review_id))
            conn.execute("UPDATE meetings SET created_at=? WHERE id=?", ("2026-07-12 09:00:00", normal_id))

        self.assertEqual(database.search_meetings("shared-search-keyword", limit=1)[0]["id"], normal_id)
        self.assertEqual(
            [row["id"] for row in database.search_meetings("shared-search-keyword", limit=1, needs_review=True)],
            [review_id],
        )
        filtered = database.search_meetings("shared-search-keyword", limit=1, needs_review=True)[0]
        self.assertEqual(filtered["quality_warning_preview"], "品質分數 80 低於 85，建議複核")

        import backend.main as main

        list_response = asgi_request(main.app, "GET", "/meetings?limit=1&needs_review=true")
        search_response = asgi_request(
            main.app,
            "GET",
            "/meetings/search?q=shared-search-keyword&limit=1&needs_review=true",
        )

        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(list_response.json()["total"], 1)
        self.assertEqual(list_response.json()["records"][0]["id"], review_id)
        self.assertEqual(list_response.json()["records"][0]["quality_warning_preview"], "品質分數 80 低於 85，建議複核")
        self.assertEqual(search_response.status_code, 200)
        self.assertEqual([row["id"] for row in search_response.json()], [review_id])
        self.assertEqual(search_response.json()[0]["quality_warning_preview"], "品質分數 80 低於 85，建議複核")


class MeetingEvidenceRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database, Path(tmpdir.name)

    def test_supplementary_evidence_is_appended_to_meeting_markdown(self):
        database, tmp_path = self._isolated_database()
        from backend import evidence

        meeting_path = tmp_path / "meeting.md"
        meeting_path.write_text(
            "## 📋 一、討論摘要 (Discussion Summary)\n原摘要\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n[00:00] 內容\n",
            encoding="utf-8",
        )
        meeting_id = database.save_meeting(
            title="補充資料測試",
            date="2026/07/05",
            source_audio="meeting.mp3",
            output_path=str(meeting_path),
            summary="原摘要",
        )
        evidence_path = tmp_path / "quote.png"
        evidence_path.write_bytes(b"png-bytes")

        with mock.patch.object(
            evidence,
            "generate_evidence_markdown",
            return_value=(
                "### 資料 1：quote.png\n"
                "- 系統判斷：與採購決策高度相關\n"
                "- 擷取重點：單價為 12,000 元\n"
                "- 原始檔案：quote.png\n"
            ),
        ):
            result = evidence.analyze_and_append_evidence(
                meeting_id=meeting_id,
                source_path=evidence_path,
                original_filename="quote.png",
                note=None,
                model="test-model",
            )

        updated = meeting_path.read_text(encoding="utf-8")
        self.assertIn("## 📎 五、補充資料與佐證", updated)
        self.assertIn("單價為 12,000 元", updated)
        self.assertTrue(Path(result["attachment_path"]).is_file())
        self.assertIn("quote.png", result["evidence_markdown"])

    def test_evidence_upload_endpoint_appends_analysis_and_returns_updated_content(self):
        database, tmp_path = self._isolated_database()
        import backend.main as main

        meeting_path = tmp_path / "meeting.md"
        meeting_path.write_text("## 📋 一、討論摘要 (Discussion Summary)\n原摘要\n", encoding="utf-8")
        meeting_id = database.save_meeting(
            title="API 補充資料測試",
            date="2026/07/05",
            source_audio="meeting.mp3",
            output_path=str(meeting_path),
            summary="原摘要",
        )

        fake_result = {
            "status": "success",
            "meeting_id": meeting_id,
            "file_name": "evidence.txt",
            "attachment_path": str(tmp_path / "attachments" / "evidence.txt"),
            "evidence_markdown": "### 資料 1：evidence.txt\n- 擷取重點：補充內容\n",
            "full_content": "## 📎 五、補充資料與佐證\n補充內容",
        }

        with mock.patch.object(main, "analyze_and_append_evidence", return_value=fake_result):
            response = asgi_request(
                main.app,
                "POST",
                f"/meetings/{meeting_id}/evidence",
                files={"file": ("evidence.txt", b"hello", "text/plain")},
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["meeting_id"], meeting_id)
        self.assertIn("補充內容", payload["full_content"])


class DocxExporterRegressionTests(unittest.TestCase):
    def test_exporter_converts_action_items_markdown_table_to_word_table(self):
        from docx import Document
        from backend.exporter import export_meeting_to_docx

        markdown = """---
title: 會議記錄 - 測試會議
---

## 📋 一、討論摘要 (Discussion Summary)

這是一段含有 **粗體重點** 的摘要。

## 📌 三、待辦事項 (Action Items)

| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |
|---|---------|---------|---------|--------|------|--------|
| A1 | D1 | R1 | **追蹤** 供應商報價與回覆 | 王經理 | 7月底 | 高 |
| A2 | D2 | 未提及 | 整理測試資料 | QA | 未定 | 中 |
"""

        with tempfile.TemporaryDirectory() as tmpdir:
            template_path = Path(tmpdir) / "template.docx"
            template_doc = Document()
            template_doc.add_table(rows=6, cols=2)
            template_doc.save(str(template_path))
            template_env = mock.patch.dict(os.environ, {"MEETING_DOCX_TEMPLATE_PATH": str(template_path)})
            template_env.start()
            self.addCleanup(template_env.stop)
            output_path = Path(tmpdir) / "meeting.docx"
            ok = export_meeting_to_docx(
                {
                    "date": "2026/07/06",
                    "title": "測試會議",
                    "full_content": markdown,
                },
                str(output_path),
            )

            self.assertTrue(ok)
            doc = Document(str(output_path))
            content_cell = doc.tables[0].cell(5, 0)
            self.assertEqual(len(content_cell.tables), 1)

            action_table = content_cell.tables[0]
            self.assertEqual(action_table.cell(0, 3).text, "任務描述")
            self.assertEqual(action_table.cell(1, 1).text, "D1")
            self.assertEqual(action_table.cell(1, 2).text, "R1")
            self.assertEqual(action_table.cell(1, 3).text, "追蹤 供應商報價與回覆")
            self.assertTrue(action_table.cell(1, 3).paragraphs[0].runs[0].bold)
            self.assertNotIn("**", content_cell.text)
            self.assertNotIn("|---", content_cell.text)


class JobQueueWorkerRegressionTests(unittest.TestCase):
    def _isolated_database(self):
        import backend.database as database

        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        patcher = mock.patch.object(database, "DB_PATH", Path(tmpdir.name) / "meetings.db")
        patcher.start()
        self.addCleanup(patcher.stop)
        database.init_db()
        return database, Path(tmpdir.name)

    def test_audio_jobs_default_to_five_attempts(self):
        database, tmpdir = self._isolated_database()
        import backend.job_queue as job_queue

        audio_path = tmpdir / "default-attempts.mp3"
        audio_path.write_bytes(b"audio")
        job_queue.enqueue_audio_job(
            "worker-default-attempts-job",
            audio_path=audio_path,
            output_dir=tmpdir,
            model="test-model",
        )

        self.assertEqual(database.get_job("worker-default-attempts-job")["max_attempts"], 5)

    def test_audio_worker_keeps_source_file_for_retry_and_terminal_failure(self):
        database, tmpdir = self._isolated_database()
        import backend.job_queue as job_queue

        audio_path = tmpdir / "retry-source.mp3"
        audio_path.write_bytes(b"audio")
        job_queue.enqueue_audio_job(
            "worker-retry-job",
            audio_path=audio_path,
            output_dir=tmpdir,
            model="test-model",
            max_attempts=2,
        )
        worker = job_queue.JobQueueWorker(poll_interval=0.01)

        first_claim = database.claim_next_pending_job()
        with mock.patch.object(job_queue, "process_audio_task", return_value=None):
            worker.process_job(first_claim)

        self.assertEqual(database.get_job("worker-retry-job")["status"], "pending")
        self.assertTrue(audio_path.exists())

        second_claim = database.claim_next_pending_job()
        with mock.patch.object(job_queue, "process_audio_task", return_value=None):
            worker.process_job(second_claim)

        self.assertEqual(database.get_job("worker-retry-job")["status"], "failed")
        self.assertTrue(audio_path.exists())

    def test_audio_worker_keeps_source_file_after_success(self):
        database, tmpdir = self._isolated_database()
        import backend.job_queue as job_queue

        audio_path = tmpdir / "success-source.mp3"
        output_path = tmpdir / "meeting.md"
        audio_path.write_bytes(b"audio")
        output_path.write_text("done", encoding="utf-8")
        job_queue.enqueue_audio_job(
            "worker-success-job",
            audio_path=audio_path,
            output_dir=tmpdir,
            model="test-model",
        )
        worker = job_queue.JobQueueWorker(poll_interval=0.01)

        def mark_done(**kwargs):
            database.update_job_status(
                kwargs["job_id"],
                "done",
                "完成",
                output_path=str(output_path),
            )
            return output_path

        claim = database.claim_next_pending_job()
        with mock.patch.object(job_queue, "process_audio_task", side_effect=mark_done):
            worker.process_job(claim)

        self.assertEqual(database.get_job("worker-success-job")["status"], "done")
        self.assertTrue(audio_path.exists())


class LineRegressionTests(unittest.TestCase):
    def test_enqueue_line_audio_job_is_idempotent_by_message_id(self):
        import backend.database as database
        import backend.job_queue as job_queue

        with tempfile.TemporaryDirectory() as tmpdir:
            with mock.patch.object(database, "DB_PATH", Path(tmpdir) / "meetings.db"):
                database.init_db()

                job_queue.enqueue_line_audio_job(
                    job_id="line-job-first",
                    message_id="same-line-message",
                    user_id="user-id",
                    model="test-model",
                )
                job_queue.enqueue_line_audio_job(
                    job_id="line-job-duplicate",
                    message_id="same-line-message",
                    user_id="user-id",
                    model="test-model",
                )

                jobs = database.list_jobs(limit=10)

        self.assertEqual(len(jobs), 1)
        self.assertEqual(jobs[0]["job_id"], "line-job-first")

    def test_push_text_sends_long_outputs_in_five_message_batches(self):
        import backend.line_handler as line_handler

        class FakeLineApi:
            def __init__(self):
                self.requests = []

            def push_message(self, request):
                self.requests.append(request)

        api = FakeLineApi()
        line_handler.push_text(api, "user-id", "a" * (4900 * 6))

        self.assertEqual(len(api.requests), 2)
        self.assertEqual([len(request.messages) for request in api.requests], [5, 1])

    def test_download_line_audio_waits_when_line_content_is_still_processing(self):
        import backend.line_handler as line_handler

        class FakeResponse:
            def __init__(self, status_code, content=b"", payload=None):
                self.status_code = status_code
                self.content = content
                self._payload = payload or {}

            def raise_for_status(self):
                if self.status_code >= 400:
                    raise RuntimeError(f"HTTP {self.status_code}")

            def json(self):
                return self._payload

        calls = []
        content_calls = 0

        def fake_get(url, **kwargs):
            nonlocal content_calls
            calls.append(url)
            if url.endswith("/content/transcoding"):
                return FakeResponse(200, payload={"status": "succeeded"})
            content_calls += 1
            if content_calls == 1:
                return FakeResponse(202)
            return FakeResponse(200, content=b"audio")

        with mock.patch.object(line_handler.requests, "get", side_effect=fake_get):
            audio = line_handler.download_line_audio("message-id")

        self.assertEqual(audio, b"audio")
        self.assertTrue(any(url.endswith("/content/transcoding") for url in calls))

    def test_line_file_message_enqueues_supported_audio_file(self):
        import backend.main as main
        import backend.line_handler as line_handler
        from linebot.v3.webhooks import (
            DeliveryContext,
            EventMode,
            FileMessageContent,
            MessageEvent,
            UserSource,
        )

        event = MessageEvent(
            source=UserSource(userId="user-id"),
            timestamp=1,
            mode=EventMode.ACTIVE,
            webhookEventId="event-id",
            deliveryContext=DeliveryContext(isRedelivery=False),
            replyToken="reply-token",
            message=FileMessageContent(
                id="file-message-id",
                fileName="meeting.mp3",
                fileSize=1234,
            ),
        )

        class FakeParser:
            def parse(self, body, signature):
                return [event]

        enqueued = []

        def fake_enqueue(**kwargs):
            enqueued.append(kwargs)

        with mock.patch.object(line_handler, "get_webhook_parser", return_value=FakeParser()), \
             mock.patch.object(line_handler, "get_line_api", return_value=object()), \
             mock.patch.object(line_handler, "reply_text"), \
             mock.patch.object(main, "enqueue_line_audio_job", side_effect=fake_enqueue):
            response = asgi_request(
                main.app,
                "POST",
                "/line-webhook",
                headers={"X-Line-Signature": "sig"},
                content=b'{"events":[]}',
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(enqueued[0]["message_id"], "file-message-id")
        self.assertEqual(enqueued[0]["user_id"], "user-id")
        self.assertEqual(enqueued[0]["file_name"], "meeting.mp3")

    def test_line_text_status_query_replies_with_latest_user_job(self):
        import backend.main as main
        import backend.database as database
        import backend.line_handler as line_handler
        from linebot.v3.webhooks import (
            DeliveryContext,
            EventMode,
            MessageEvent,
            TextMessageContent,
            UserSource,
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            with mock.patch.object(database, "DB_PATH", Path(tmpdir) / "meetings.db"):
                database.init_db()
                database.create_job(
                    "line-status-job",
                    task_type="line_audio_processing",
                    source="line",
                    payload={
                        "message_id": "line-message-id",
                        "user_id": "user-id",
                        "model": "test-model",
                    },
                    message="LINE 媒體已接收，已排入可靠處理佇列。",
                )
                database.update_job_status(
                    "line-status-job",
                    "processing",
                    "📝 正在轉錄第 1/2 段音訊...",
                    progress_current=1,
                    progress_total=2,
                )

                event = MessageEvent(
                    source=UserSource(userId="user-id"),
                    timestamp=1,
                    mode=EventMode.ACTIVE,
                    webhookEventId="event-id",
                    deliveryContext=DeliveryContext(isRedelivery=False),
                    replyToken="reply-token",
                    message=TextMessageContent(id="text-message-id", text="狀態", quoteToken="quote-token"),
                )

                class FakeParser:
                    def parse(self, body, signature):
                        return [event]

                replies = []
                with mock.patch.object(line_handler, "get_webhook_parser", return_value=FakeParser()), \
                     mock.patch.object(line_handler, "get_line_api", return_value=object()), \
                     mock.patch.object(line_handler, "reply_text", side_effect=lambda api, token, text: replies.append(text)), \
                     mock.patch.object(line_handler, "process_line_text_in_background") as chat_handler:
                    response = asgi_request(
                        main.app,
                        "POST",
                        "/line-webhook",
                        headers={"X-Line-Signature": "sig"},
                        content=b'{"events":[]}',
                    )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(chat_handler.called)
        self.assertTrue(any("line-sta" in reply for reply in replies))
        self.assertTrue(any("處理中" in reply for reply in replies))
        self.assertTrue(any("1/2" in reply for reply in replies))

    def test_line_audio_flow_creates_job_and_pushes_returned_output_path(self):
        import backend.database as database
        import backend.line_handler as line_handler

        job_id = "line-regression-job"
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            output_path = tmpdir_path / "meeting_notes_line.md"
            output_path.write_text("會議記錄內容", encoding="utf-8")

            pushed_messages = []

            with mock.patch.object(database, "DB_PATH", tmpdir_path / "meetings.db"):
                database.init_db()
                with database.get_db() as conn:
                    conn.execute("DELETE FROM jobs WHERE job_id=?", (job_id,))

                with mock.patch.object(line_handler, "get_line_api", return_value=object()), \
                     mock.patch.object(line_handler, "download_line_audio", return_value=b"audio"), \
                     mock.patch.object(line_handler, "SOURCE_AUDIO_DIR", tmpdir_path), \
                     mock.patch.object(line_handler, "process_audio_task", return_value=output_path) as process_mock, \
                     mock.patch.object(line_handler, "push_text", side_effect=lambda api, user_id, text: pushed_messages.append(text)):
                    line_handler.process_line_audio_in_background(
                        job_id=job_id,
                        message_id="message-id",
                        user_id="user-id",
                        model="gemini-3.1-flash-lite",
                    )

                saved_audio = list(tmpdir_path.glob("line-reg_*.m4a"))
                self.assertEqual(len(saved_audio), 1)
                self.assertEqual(saved_audio[0].read_bytes(), b"audio")
                self.assertFalse(process_mock.call_args.kwargs["cleanup_source_audio"])

                self.assertIsNotNone(database.get_job(job_id))
                self.assertTrue(any("會議記錄內容" in msg for msg in pushed_messages))

                with database.get_db() as conn:
                    conn.execute("DELETE FROM jobs WHERE job_id=?", (job_id,))

    def test_line_audio_flow_reuses_existing_audio_with_same_sha256(self):
        import backend.database as database
        import backend.line_handler as line_handler

        job_id = "line-dedup-job"
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            output_path = tmpdir_path / "meeting_notes_line.md"
            output_path.write_text("line result", encoding="utf-8")
            existing_audio = tmpdir_path / "existing-source.m4a"
            existing_audio.write_bytes(b"audio")

            with mock.patch.object(database, "DB_PATH", tmpdir_path / "meetings.db"):
                database.init_db()
                with database.get_db() as conn:
                    conn.execute("DELETE FROM jobs WHERE job_id=?", (job_id,))

                with mock.patch.object(line_handler, "get_line_api", return_value=object()), \
                     mock.patch.object(line_handler, "download_line_audio", return_value=b"audio"), \
                     mock.patch.object(line_handler, "SOURCE_AUDIO_DIR", tmpdir_path), \
                     mock.patch.object(line_handler, "process_audio_task", return_value=output_path) as process_mock, \
                     mock.patch.object(line_handler, "push_text"):
                    line_handler.process_line_audio_in_background(
                        job_id=job_id,
                        message_id="message-id",
                        user_id="user-id",
                        model="gemini-3.1-flash-lite",
                    )

                self.assertEqual(process_mock.call_args.kwargs["audio_path"], existing_audio)
                self.assertEqual(list(tmpdir_path.glob("*.m4a")), [existing_audio])
                self.assertFalse(list(tmpdir_path.glob(".upload_*")))

                with database.get_db() as conn:
                    conn.execute("DELETE FROM jobs WHERE job_id=?", (job_id,))

    def test_line_delivery_message_excludes_full_transcript_and_points_to_output_path(self):
        import backend.line_handler as line_handler

        markdown = """---
title: 會議記錄
---

## 📋 一、討論摘要 (Discussion Summary)
摘要內容。

## ✅ 二、最終決議 (Final Decisions)
決議內容。

## 📌 三、待辦事項 (Action Items)
| # | 任務描述 | 負責人 | 期限 | 優先級 |
|---|---------|--------|------|--------|
| 1 | 整理需求 | 王經理 | 下週 | 高 |

## 📝 四、完整逐字稿 (Verbatim Transcript)
[00:00] **[發言者 A]**：這是一大段完整逐字稿，不應該推回 LINE。
"""

        output_path = Path("/tmp/output/meeting_notes.md")
        message = line_handler.build_line_meeting_delivery_message(
            markdown,
            output_path,
        )

        self.assertIn("摘要內容", message)
        self.assertIn("決議內容", message)
        self.assertIn("整理需求", message)
        self.assertIn("完整逐字稿已保存", message)
        self.assertIn(str(output_path), message)
        self.assertNotIn("這是一大段完整逐字稿", message)


class UiRegressionTests(unittest.TestCase):
    def test_regression_tests_do_not_import_deprecated_fastapi_testclient(self):
        source = (ROOT / "tests" / "test_regressions.py").read_text(encoding="utf-8")
        deprecated_import = "fastapi" + ".testclient"

        self.assertNotIn(deprecated_import, source)

    def test_desktop_gui_handles_backend_failed_status(self):
        source = (ROOT / "gui" / "app.py").read_text(encoding="utf-8")

        self.assertIn('status == "failed"', source)
        self.assertNotIn('status == "error"', source)

    def test_web_ui_can_cancel_active_backend_jobs(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn("/jobs/${jobId}/cancel", html)
        self.assertIn("activeUploadJobId", html)
        self.assertIn("activeRecordingJobId", html)

    def test_web_ui_has_job_dashboard(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="job-dashboard"', html)
        self.assertIn('id="job-dashboard-summary"', html)
        self.assertIn('id="job-dashboard-toggle"', html)
        self.assertIn('aria-expanded="false"', html)
        self.assertIn("function toggleJobDashboard", html)
        self.assertIn("function updateJobDashboardSummary", html)
        self.assertIn("async function loadJobs", html)
        self.assertIn("/jobs?limit=20", html)
        self.assertIn("任務狀態", html)

    def test_web_ui_has_metrics_panel_and_job_actions(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="ops-dashboard"', html)
        self.assertIn("async function loadMetrics", html)
        self.assertIn("/metrics", html)
        self.assertIn("/health", html)
        self.assertIn("health.status === 'ok'", html)
        self.assertNotIn("healthEl.textContent = failed ? '需注意' : '正常'", html)
        self.assertIn("/jobs/${jobId}/retry", html)
        self.assertIn("/jobs/${jobId}", html)
        self.assertIn("維運狀態", html)
        self.assertIn('id="ops-needs-review"', html)
        self.assertIn("data.meetings?.needs_review", html)
        self.assertIn('id="ops-source-storage"', html)
        self.assertIn('id="ops-source-storage-tile"', html)
        self.assertIn('id="source-storage-modal"', html)
        self.assertIn('id="source-storage-preview"', html)
        self.assertIn('id="source-storage-list"', html)
        self.assertIn('id="source-storage-search"', html)
        self.assertIn('id="source-storage-filter"', html)
        self.assertIn('id="source-storage-media-filter"', html)
        self.assertIn('id="source-storage-sort"', html)
        self.assertIn('id="source-storage-refresh"', html)
        self.assertIn('id="source-storage-reset"', html)
        self.assertIn('id="source-storage-status"', html)
        self.assertIn('id="source-storage-more"', html)
        self.assertIn('id="source-storage-load-more"', html)
        self.assertIn('aria-live="polite"', html)
        self.assertIn('oninput="renderSourceStorageInventory()"', html)
        self.assertIn('onchange="renderSourceStorageInventory()"', html)
        self.assertIn("data.storage", html)
        self.assertIn("storage.source_media_files", html)
        self.assertIn("storage.source_media_bytes", html)
        self.assertIn("storage.source_media_unlinked_files", html)
        self.assertIn("storage.source_media_unlinked_bytes", html)
        self.assertIn("storage.source_media_archived_files", html)
        self.assertIn("storage.source_media_archived_bytes", html)
        self.assertIn("storage.source_media_largest_files", html)
        self.assertIn("function sourceStorageTitle", html)
        self.assertIn("function sourceArchiveRetentionLabel", html)
        self.assertIn("runtimeConfig.source_media_archive_retention_days", html)
        self.assertIn("已移除備份保留", html)
        self.assertIn("已移除備份自動清理停用", html)
        self.assertIn("totalManagedBytes", html)
        self.assertIn("managedFileLabel", html)
        self.assertIn("sourceStorageEl.classList.toggle('attention', unlinkedFiles > 0)", html)
        self.assertIn("sourceStorageEl.classList.remove('attention')", html)
        self.assertIn("已移除備份", html)
        self.assertIn("最大檔案：", html)
        self.assertIn("未連結原始檔：", html)
        self.assertIn("file.linked_meeting_title", html)
        self.assertIn("formatBytes(sourceBytes)", html)
        self.assertIn("原始檔", html)
        self.assertIn("const SOURCE_STORAGE_FETCH_LIMIT = 500", html)
        self.assertIn("function openSourceStorageInventory", html)
        self.assertIn("function closeSourceStorageInventory", html)
        self.assertIn("function previewSourceMedia", html)
        self.assertIn("function closeSourceMediaPreview", html)
        self.assertIn("function setSourceStorageStatus", html)
        self.assertIn("function sourceStorageExtension", html)
        self.assertIn("function sourceStorageMediaType", html)
        self.assertIn("function sourceStorageAudioFallbackCapable", html)
        self.assertIn("function renderSourceStorageFile", html)
        self.assertIn("let sourceStorageInventoryState", html)
        self.assertIn("function sourceStorageSearchText", html)
        self.assertIn("function sourceStorageFileMatches", html)
        self.assertIn("function sourceStorageLiveFilterMatches", html)
        self.assertIn("function sourceStorageMediaTypeMatches", html)
        self.assertIn("function sourceStorageSortFiles", html)
        self.assertIn("function sourceStorageHiddenCount", html)
        self.assertIn("function sourceStorageHiddenNotice", html)
        self.assertIn("function sourceStorageLoadedArchiveBytes", html)
        self.assertIn("function sourceStorageArchiveUnlistedBytes", html)
        self.assertIn("function sourceStorageStatusSummary", html)
        self.assertIn("function syncSourceStorageLoadMoreButton", html)
        self.assertIn("async function fetchSourceStorageInventoryPage", html)
        self.assertIn("async function loadMoreSourceStorageInventory", html)
        self.assertIn("async function refreshSourceStorageInventory", html)
        self.assertIn("function resetSourceStorageFilters", html)
        self.assertIn("function renderSourceStorageInventory", html)
        self.assertIn("function renderSourceStorageInventory() {\n  const list = document.getElementById('source-storage-list');\n  if (!list || !sourceStorageInventoryState.loaded) return;\n  closeSourceMediaPreview();", html)
        self.assertIn("sourceStorageInventoryState = {", html)
        self.assertIn("loadedAt: new Date().toISOString()", html)
        self.assertIn("renderSourceStorageInventory();", html)
        self.assertIn("setSourceStorageStatus('正在更新原始檔清單...');", html)
        self.assertIn("最後更新：${formatDate(sourceStorageInventoryState.loadedAt)}", html)
        self.assertIn("目前載入 ${totalLoaded} / ${totalAvailable} 個項目", html)
        self.assertIn("尚有未載入：${parts.join('、')}（每次載入 ${SOURCE_STORAGE_FETCH_LIMIT}）", html)
        self.assertIn("載入更多（尚有 ${hiddenTotal} 個）", html)
        self.assertIn("sourceStorageInventoryState.files.length", html)
        self.assertIn("sourceStorageInventoryState.archivedFiles.length", html)
        self.assertIn("const archiveUnlistedBytes = sourceStorageArchiveUnlistedBytes(archivePayload, sourceStorageInventoryState.archivedFiles);", html)
        self.assertIn("return total + Math.max(0, mediaBytes) + Math.max(0, metadataBytes);", html)
        self.assertIn("if (hiddenArchive > 0) return 0;", html)
        self.assertIn("未列出備份容量 ${formatBytes(archiveUnlistedBytes)}", html)
        self.assertIn("source-storage-badge warning", html)
        self.assertIn("setSourceStorageStatus(`更新失敗：${err.message}`, true);", html)
        self.assertIn("sourceStorageMediaTypeMatches(file, mediaFilter)", html)
        self.assertIn("sourceStorageSortFiles(", html)
        self.assertIn("refresh.textContent = '更新中';", html)
        self.assertIn("await openSourceStorageInventory();", html)
        self.assertIn("await loadMetrics();", html)
        self.assertIn("if (sort) sort.value = 'modified_desc';", html)
        self.assertIn("目前顯示 ${shownCount} 個", html)
        self.assertIn("function deleteUnlinkedSourceMedia", html)
        self.assertIn("const sourceUrl = `${API}/source-media/inventory/${encodeURIComponent(filename)}`", html)
        self.assertIn("const downloadUrl = `${sourceUrl}?download=1`", html)
        self.assertIn("const mediaType = sourceStorageMediaType(file)", html)
        self.assertIn("sourceStorageMediaType({ name, source_media_type: mediaType })", html)
        self.assertIn("sourceStorageExtension(file)", html)
        self.assertIn("const videoExtensions = new Set(['.webm', '.mp4', '.mov', '.mkv', '.avi', '.mpeg', '.mpg', '.wmv'])", html)
        self.assertIn("const audioExtensions = new Set(['.mp3', '.wav', '.m4a', '.aac', '.ogg', '.opus', '.flac', '.wma'])", html)
        self.assertIn("source-storage-badge media-type", html)
        self.assertIn("/source-media/inventory?limit=${SOURCE_STORAGE_FETCH_LIMIT}&offset=${liveOffset}", html)
        self.assertIn("/source-media/archive?limit=${SOURCE_STORAGE_FETCH_LIMIT}&offset=${archiveOffset}", html)
        self.assertIn("/source-media/archive/file?archive_id=", html)
        self.assertIn("/source-media/archive/restore?archive_id=", html)
        self.assertIn("const metadataBytes = Number(file.metadata_bytes || 0);", html)
        self.assertIn("中繼資料 ${formatBytes(metadataBytes)}", html)
        self.assertIn("/source-media/inventory/${encodeURIComponent(filename)}", html)
        self.assertIn("previewSourceMedia(this.dataset.sourceUrl, this.dataset.sourceType, this.dataset.sourceFile)", html)
        self.assertIn("source-storage-preview-title", html)
        self.assertIn("source-storage-preview-unsupported", html)
        self.assertIn("const downloadUrl = url.includes('?') ? `${url}&download=1` : `${url}?download=1`;", html)
        self.assertIn("const safeDownloadUrl = escapeHtml(downloadUrl);", html)
        self.assertIn('href="${safeDownloadUrl}" download', html)
        self.assertIn('aria-label="在新分頁開啟${mediaTitle}"', html)
        self.assertIn('aria-label="下載${mediaTitle}"', html)
        self.assertIn('aria-label="關閉原始檔預覽"', html)
        self.assertIn('role="status" aria-live="polite"', html)
        self.assertIn("const resolvedMediaType = sourceStorageMediaType({ name, source_media_type: mediaType });", html)
        self.assertIn("requestedMode === 'audio' && canAudioFallback ? 'audio' : resolvedMediaType", html)
        self.assertIn("sourceStorageAudioFallbackCapable({ name, source_media_type: mediaType })", html)
        self.assertIn("const mediaActionLabel = mediaTypeLabel || '原始檔';", html)
        self.assertIn('aria-label="預覽${mediaActionLabel}"', html)
        self.assertIn('aria-label="在新分頁開啟${mediaActionLabel}"', html)
        self.assertIn('aria-label="下載${mediaActionLabel}"', html)
        self.assertIn('aria-label="開啟連結會議"', html)
        self.assertIn('aria-label="刪除未連結原始檔"', html)
        self.assertIn('aria-label="還原原始檔備份"', html)
        self.assertIn('data-source-mode="audio"', html)
        self.assertIn('data-source-mode="video"', html)
        self.assertIn('aria-label="切換為音訊預覽"', html)
        self.assertIn('aria-label="切換為影片預覽"', html)
        self.assertIn("🔊 音訊預覽", html)
        self.assertIn("🎥 影片預覽", html)
        self.assertIn("playerMode === 'audio'", html)
        self.assertIn("無法判斷此檔案是否可直接預覽", html)
        self.assertIn("<video controls preload=\"metadata\"", html)
        self.assertIn('class="source-storage-media-player"', html)
        self.assertIn("enhanceSourceStoragePreviewPlayer(preview);", html)
        self.assertIn("function enhanceSourceStoragePreviewPlayer", html)
        self.assertIn("querySelector('video.source-storage-media-player, audio')", html)
        self.assertIn("<audio controls preload=\"metadata\"", html)
        self.assertIn("method: 'DELETE'", html)
        self.assertIn("method: 'POST'", html)
        self.assertIn("payload.backup_path", html)
        self.assertIn("已從清單移除", html)
        self.assertIn("function restoreSourceMediaArchive", html)
        self.assertIn("已移除備份", html)
        self.assertIn('target="_blank"', html)
        self.assertIn(">新分頁</a>", html)
        self.assertIn('download aria-label="下載${mediaActionLabel}">下載</a>', html)
        self.assertIn("openMeetingFromSourceStorage", html)
        self.assertIn("未連結會議", html)
        self.assertIn("請先確認這不是仍需保留的證據檔", html)
        self.assertIn("source-media-action danger", html)
        self.assertIn("function showNeedsReviewMeetings", html)
        self.assertIn("loadMeetings(search.value.trim())", html)
        self.assertIn("需複核", html)
        self.assertIn('id="ops-ngrok"', html)
        self.assertIn('id="ops-ngrok-tile"', html)
        self.assertIn("ngrokTile.title = ngrokDetail", html)
        self.assertIn("LINE/ngrok", html)
        self.assertIn("data.ngrok", html)

    def test_web_ui_loads_runtime_config_and_prevents_oversized_uploads(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="upload-limit-hint"', html)
        self.assertIn("let runtimeConfig", html)
        self.assertIn("async function loadRuntimeConfig", html)
        self.assertIn("async function initializeDashboard", html)
        self.assertIn("await loadRuntimeConfig()", html)
        self.assertIn("Promise.allSettled", html)
        self.assertIn("initializeDashboard()", html)
        self.assertIn("/config", html)
        self.assertIn("/upload-media", html)
        self.assertNotIn("${API}/upload-audio", html)
        self.assertIn("source_media_archive_retention_days", html)
        self.assertIn("selectedFile.size > runtimeConfig.max_upload_bytes", html)
        self.assertIn("formatBytes(runtimeConfig.max_upload_bytes)", html)
        self.assertIn("function supportedUploadExtensions", html)
        self.assertIn("function updateUploadAcceptList", html)
        self.assertIn('accept=".mp3,.wav,.m4a,.aac,.ogg,.flac,.webm,.mp4,.mov,.avi,.mkv,.mpeg,.mpg,.wmv"', html)
        self.assertIn("const FALLBACK_UPLOAD_EXTENSIONS = ['.mp3', '.wav', '.m4a', '.aac', '.ogg', '.flac', '.webm', '.mp4', '.mov', '.avi', '.mkv', '.mpeg', '.mpg', '.wmv']", html)
        self.assertIn("const FALLBACK_UPLOAD_ACCEPT = FALLBACK_UPLOAD_EXTENSIONS.join(',')", html)
        self.assertIn("input.accept = extensions.length ? extensions.join(',') : FALLBACK_UPLOAD_ACCEPT", html)
        self.assertIn("updateUploadAcceptList()", html)
        self.assertIn("function uploadFileExtension", html)
        self.assertIn("function isSupportedUploadFile", html)
        self.assertIn("function uploadUnsupportedMessage", html)
        self.assertIn("if (!isSupportedUploadFile(file))", html)
        self.assertIn("if (!isSupportedUploadFile(selectedFile))", html)
        self.assertIn("不支援的媒體格式", html)
        self.assertIn("上傳音訊/影片", html)
        self.assertIn("上傳音訊/影片並產生會議記錄", html)
        self.assertIn("請先選擇一個音訊或影片檔！", html)
        self.assertIn("let uploadSelectionError = ''", html)
        self.assertIn("let uploadBusy = false", html)
        self.assertIn("uploadSelectionError = uploadUnsupportedMessage(file)", html)
        self.assertIn("alert(uploadSelectionError || \"請先選擇一個音訊或影片檔！\")", html)
        self.assertIn("fileInput.value = ''", html)
        self.assertIn("function updateUploadActionState", html)
        self.assertIn("btnUpload.disabled = uploadBusy || !selectedFile", html)
        self.assertIn("if (uploadBusy && activeUploadJobId)", html)
        self.assertIn("uploadBusy = true", html)

    def test_web_ui_has_readable_api_errors_and_job_event_timeline(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn("async function apiErrorMessage", html)
        self.assertIn("payload.detail", html)
        self.assertIn('id="job-event-panel"', html)
        self.assertIn("async function loadJobEvents", html)
        self.assertIn("/jobs/${jobId}/events", html)
        self.assertIn("事件", html)

    def test_web_ui_can_upload_supplementary_evidence_for_a_meeting(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="evidence-file-input"', html)
        self.assertIn("async function uploadEvidence", html)
        self.assertIn("/meetings/${id}/evidence", html)
        self.assertIn("補充資料", html)

    def test_web_ui_can_rerun_a_meeting_from_detail_view(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="rerun-meeting-button"', html)
        self.assertIn("async function rerunMeeting", html)
        self.assertIn("/meetings/${id}/rerun", html)
        self.assertIn("重跑", html)

    def test_web_ui_has_wide_reading_mode_for_meeting_detail(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="reading-mode-button"', html)
        self.assertIn("function toggleReadingMode", html)
        self.assertIn("function syncReadingModeButton", html)
        self.assertIn("reading-mode", html)
        self.assertIn("document-view", html)
        self.assertIn("@media (max-width: 1100px)", html)

    def test_web_ui_prevents_detail_header_text_stacking(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn("grid-template-columns: minmax(0, 1fr);", html)
        self.assertIn("overflow-wrap: break-word;", html)
        self.assertIn("word-break: keep-all;", html)
        self.assertIn("text-wrap: balance;", html)
        self.assertNotIn("overflow-wrap: anywhere", html)
        self.assertIn("display: flex; align-items: center; flex-wrap: wrap; gap: 6px 12px;", html)
        self.assertIn(".card-meta span", html)
        self.assertIn("min-width: 0; max-width: 100%;", html)
        self.assertIn(".detail-actions .btn-primary", html)
        self.assertIn("white-space: nowrap;", html)
        self.assertIn(".source-storage-preview-header", html)
        self.assertIn(".source-storage-preview-actions", html)
        self.assertRegex(html, r"\.audio-evidence-header \{[^}]*flex-wrap: wrap;")
        self.assertRegex(html, r"\.audio-evidence-label \{[^}]*flex: 1 1 240px;[^}]*max-width: 100%;")
        self.assertRegex(html, r"\.source-media-actions \{[^}]*max-width: 100%;")
        self.assertRegex(html, r"\.source-media-action \{[^}]*max-width: 100%;")
        self.assertRegex(html, r"\.source-storage-preview-header \{[^}]*flex-wrap: wrap;")
        self.assertRegex(html, r"\.source-storage-preview-actions \{[^}]*flex: 0 1 auto;[^}]*max-width: 100%;")
        self.assertRegex(html, r"\.source-storage-preview-title \{[^}]*flex: 1 1 240px;[^}]*max-width: 100%;")
        self.assertRegex(html, r"\.source-storage-preview-title \{[^}]*overflow-wrap: break-word;[^}]*word-break: keep-all;")
        self.assertRegex(html, r"\.source-storage-name \{[^}]*overflow-wrap: break-word;[^}]*word-break: keep-all;")
        self.assertRegex(html, r"\.source-storage-preview-unsupported \{[^}]*overflow-wrap: break-word;[^}]*word-break: keep-all;")
        self.assertIn(".source-storage-item", html)
        self.assertIn("align-items: stretch;", html)
        self.assertIn(".source-storage-status", html)
        self.assertIn(".source-storage-status.error", html)
        self.assertIn(".source-storage-badge.warning", html)
        self.assertIn(".source-storage-more", html)
        self.assertIn(".source-storage-more[hidden]", html)
        self.assertIn(".source-storage-load-more", html)
        self.assertIn(".source-storage-controls", html)
        self.assertIn("grid-template-columns: minmax(0, 1fr) repeat(3, minmax(140px, 180px)) auto auto;", html)
        self.assertIn(".source-storage-refresh", html)
        self.assertIn(".source-storage-reset", html)
        self.assertIn("grid-template-columns: 1fr;", html)

    def test_web_ui_can_record_screen_audio_with_toggleable_microphone(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="tab-screen"', html)
        self.assertIn('id="rec-screen-options"', html)
        self.assertIn('id="rec-include-mic" checked', html)
        self.assertIn('id="btn-rec-mic"', html)
        self.assertIn('id="rec-size-estimate" aria-live="polite"', html)
        self.assertIn('id="rec-status" role="status" aria-live="polite"', html)
        self.assertIn("錄製結果會顯示於此", html)
        self.assertIn("switchRecMode('screen')", html)
        self.assertIn("鏡頭錄影", html)
        self.assertIn("螢幕錄影/分頁音訊", html)
        self.assertIn("錄影平衡", html)
        self.assertIn('onchange="handleRecordingProfileChange()"', html)
        self.assertIn("let lastAudioRecordingProfileId = 'audio_standard';", html)
        self.assertIn("let activeRecordingMode = null;", html)
        self.assertIn("let activeRecordingProfileId = null;", html)
        self.assertIn("function handleRecordingProfileChange", html)
        self.assertIn("lastAudioRecordingProfileId = select.value === 'audio_compact'", html)
        self.assertIn("DEFAULT_RECORDING_PROFILE_LABELS", html)
        self.assertIn("function recordingProfileBaseLabel", html)
        self.assertIn("runtimeConfig.recording_profiles?.[profile]?.label", html)
        self.assertIn("function recordingProfileOptionLabel", html)
        self.assertIn("function updateRecordingProfileOptionLabels", html)
        self.assertIn("updateRecordingProfileOptionLabels();", html)
        self.assertIn("option.textContent = recordingProfileOptionLabel(option.value);", html)
        self.assertIn("const profileLabel = recordingProfileBaseLabel(getRecordingProfileId());", html)
        self.assertIn("hint.textContent = `${prefix}${current}每小時預估容量", html)
        self.assertIn("const desired = videoMode ? 'video_balanced' : lastAudioRecordingProfileId;", html)
        self.assertIn("function updateRecordingModeOptions", html)
        self.assertIn("function setRecordingConfigurationLocked", html)
        self.assertIn("document.getElementById('tab-audio').disabled = locked", html)
        self.assertIn("document.getElementById('rec-quality-profile').disabled = locked", html)
        self.assertIn("function setRecordingSubmissionLocked", html)
        self.assertIn("document.getElementById('rec-title').disabled = locked", html)
        self.assertIn(".form-input:disabled", html)
        self.assertIn("if (document.getElementById('rec-quality-profile').disabled) return;", html)
        self.assertIn("function recordingModeLabel", html)
        self.assertIn("function recordingPermissionHint", html)
        self.assertIn("function setRecordingResult", html)
        self.assertIn(".rec-result.error", html)
        self.assertIn("recResult.classList.toggle('error'", html)
        self.assertIn("if (mode === 'screen') return '螢幕錄製';", html)
        self.assertIn("if (mode === 'video') return '鏡頭錄影';", html)
        self.assertIn("`🔴 ${recordingModeLabel()}中...`", html)
        self.assertIn("function toggleRecordingMic", html)
        self.assertIn("function updateRecordingMicToggle", html)
        self.assertIn("function canMuteRecordingMic", html)
        self.assertIn("function restoreRecorderIdleState", html)
        self.assertIn("function discardActiveLocalRecording", html)
        self.assertIn("discardActiveLocalRecording();", html)
        self.assertIn("關閉視窗會取消本次錄製且不會送出 AI 分析", html)
        self.assertIn("recorder.onstop = null;", html)
        self.assertIn("restoreRecorderIdleState();", html)
        self.assertIn("錄製未開始，請重新確認權限後再試", html)
        self.assertIn("錄製已取消，內容未送出", html)
        self.assertIn("recPreview.classList.remove('show');", html)
        self.assertIn("getDisplayMedia", html)
        self.assertIn("audio: true", html)
        self.assertIn("const includeMic = Boolean(recIncludeMic?.checked)", html)
        self.assertIn("if (includeMic)", html)
        self.assertIn("getUserMedia({ audio: microphoneConstraints })", html)
        self.assertIn("sampleRate: profile.audio_sample_rate", html)
        self.assertIn("recMicAudioTracks = micAudioTracks", html)
        self.assertIn("track.enabled = !recMicMuted", html)
        self.assertIn("if (!recMicMuted && !canMuteRecordingMic()) return;", html)
        self.assertIn("recMode !== 'screen' || recScreenAudioCaptured", html)
        self.assertIn("線上會議其他人的聲音不會被錄到", html)
        self.assertIn("線上會議其他人聲音可能未錄到", html)
        self.assertIn("未偵測到音訊，請重新分享並勾選分享音訊", html)
        self.assertIn("麥克風已關閉", html)
        self.assertIn("mediaRecorder.requestData()", html)
        self.assertIn("Unable to flush recorder data before stop.", html)
        self.assertIn("!recordedChunks.length || file.size <= 0", html)
        self.assertIn("const completedModeLabel = recordingModeLabel(completedMode);", html)
        self.assertIn("setRecordingResult(`❌ ${completedModeLabel}失敗`, 'error')", html)
        self.assertIn("${completedModeLabel}失敗：沒有取得任何錄製資料", html)
        self.assertIn("setRecordingResult(`❌ ${completedModeLabel}檔案過大`, 'error')", html)
        self.assertIn("${completedModeLabel}檔案超過上限", html)
        self.assertIn("setRecordingResult(`✅ ${completedModeLabel}完成！正在送出 AI 分析，請稍候...`)", html)
        self.assertIn("${completedModeLabel}完成", html)
        self.assertIn("setRecordingResult(`❌ ${completedModeLabel}上傳失敗`, 'error')", html)
        self.assertIn("activeRecordingProfileId = getRecordingProfileId();", html)
        self.assertIn("const completedProfileId = activeRecordingProfileId || getRecordingProfileId();", html)
        self.assertIn("formData.append('recording_profile', completedProfileId);", html)
        self.assertIn("setRecordingSubmissionLocked(true);", html)
        self.assertIn("setRecordingSubmissionLocked(false);", html)
        self.assertNotIn("formData.append('recording_profile', getRecordingProfileId());", html)
        self.assertIn("setRecordingResult('❌ AI 分析失敗', 'error')", html)
        self.assertIn("setRecordingResult('已取消處理。', 'error')", html)
        self.assertIn("Optional microphone capture failed; continuing with screen audio.", html)
        self.assertIn("目前沒有可錄製的聲音", html)
        self.assertIn("分享分頁或視窗時需勾選分享音訊", html)
        self.assertIn("攝影機與麥克風權限", html)
        self.assertIn("麥克風權限", html)
        self.assertIn("${recordingPermissionHint()}", html)
        self.assertIn("...displayStream.getVideoTracks()", html)
        self.assertIn("displayStream.getAudioTracks()", html)
        self.assertIn("...micAudioTracks", html)
        self.assertIn("createMediaStreamDestination", html)

    def test_frontend_smoke_script_checks_static_ui_and_upload_guard(self):
        smoke_script = ROOT / "scripts" / "smoke_e2e.sh"

        self.assertTrue(smoke_script.is_file())
        self.assertTrue(os.access(smoke_script, os.X_OK))

        script = smoke_script.read_text(encoding="utf-8")
        self.assertIn("BASE_URL", script)
        self.assertIn("/history", script)
        self.assertIn("ops-dashboard", script)
        self.assertIn("/upload-media", script)
        self.assertNotIn("/upload-audio", script)
        self.assertIn("fake.mp3", script)
        self.assertIn("415", script)

    def test_desktop_gui_client_uses_primary_media_upload_endpoint(self):
        client_source = (ROOT / "gui" / "api_client.py").read_text(encoding="utf-8")

        self.assertIn("/upload-media", client_source)
        self.assertNotIn("/upload-audio", client_source)

    def test_python_docx_is_declared_as_runtime_dependency(self):
        requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8")

        self.assertIn("python-docx", requirements)

    def test_audioop_lts_is_only_installed_for_python_313_or_newer(self):
        requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8")

        self.assertIn('audioop-lts>=0.2.2; python_version >= "3.13"', requirements)


class TestSuiteQualityRegressionTests(unittest.TestCase):
    def test_root_test_modules_do_not_print_at_import_time(self):
        for filename in ("test_gemini.py", "test_regex.py"):
            source = (ROOT / filename).read_text(encoding="utf-8")

            self.assertNotIn("print(", source, msg=f"{filename} should be quiet during discovery")


class StartupScriptRegressionTests(unittest.TestCase):
    def test_startup_kills_existing_meeting_assistant_listener_before_launch(self):
        import start

        calls = []

        class Result:
            def __init__(self, stdout="", returncode=0):
                self.stdout = stdout
                self.returncode = returncode

        def fake_run(args, **kwargs):
            calls.append(args)
            if args[0] == "lsof":
                return Result("123\n456\n")
            if args[0] == "ps" and args[2] == "123":
                return Result("/usr/bin/python -m uvicorn backend.main:app --port 8001\n")
            if args[0] == "ps" and args[2] == "456":
                return Result("/usr/bin/python unrelated_server.py\n")
            if args[:2] == ["kill", "-0"]:
                return Result(returncode=1)
            return Result()

        with mock.patch.object(start.platform, "system", return_value="Darwin"), \
             mock.patch.object(start.os, "getpid", return_value=999), \
             mock.patch.object(start.subprocess, "run", side_effect=fake_run), \
             mock.patch.object(start.time, "sleep"), \
             mock.patch("builtins.print"):
            start.terminate_existing_server(8001)

        self.assertIn(["kill", "-TERM", "123"], calls)
        self.assertNotIn(["kill", "-TERM", "456"], calls)

    def test_startup_script_uses_cleanup_before_uvicorn_launch(self):
        source = (ROOT / "start.py").read_text(encoding="utf-8")
        main_block = source[source.index('if __name__ == "__main__":') :]

        self.assertIn("terminate_existing_server(SERVER_PORT)", main_block)
        self.assertLess(
            main_block.index("terminate_existing_server(SERVER_PORT)"),
            main_block.index("subprocess.run(["),
        )

    def test_startup_can_build_mobile_history_url_with_api_key(self):
        import start

        url = start.mobile_history_url("192.168.1.20", 8001, "mobile secret")

        self.assertEqual(url, "http://192.168.1.20:8001/history?api_key=mobile%20secret")

    def test_startup_prints_direct_mobile_access_url_for_same_network(self):
        import start

        with mock.patch.dict(start.os.environ, {"APP_API_KEY": "mobile-key"}, clear=False), \
             mock.patch.object(start, "local_lan_ip", return_value="192.168.1.20"), \
             mock.patch("builtins.print") as printed:
            start.print_access_urls()

        output = "\n".join(str(call.args[0]) for call in printed.call_args_list if call.args)
        self.assertIn("手機 / 平板：http://192.168.1.20:8001/history", output)
        self.assertNotIn("手機 / 平板：http://192.168.1.20:8001/history?api_key=", output)

    def test_startup_generates_temporary_api_key_when_missing(self):
        import start

        with mock.patch.dict(start.os.environ, {}, clear=True), \
             mock.patch.object(start.secrets, "token_urlsafe", return_value="generated-mobile-key"), \
             mock.patch("builtins.print"):
            key = start.ensure_app_api_key()
            env_key = start.os.environ["APP_API_KEY"]

        self.assertEqual(key, "generated-mobile-key")
        self.assertEqual(env_key, "generated-mobile-key")

    def test_startup_main_generates_api_key_before_mobile_urls_and_backend_launch(self):
        source = (ROOT / "start.py").read_text(encoding="utf-8")
        main_block = source[source.index('if __name__ == "__main__":') :]

        self.assertIn("ensure_app_api_key()", main_block)
        self.assertLess(
            main_block.index("ensure_app_api_key()"),
            main_block.index("print_access_urls()"),
        )
        self.assertLess(
            main_block.index("ensure_app_api_key()"),
            main_block.index("subprocess.run(["),
        )

    def test_startup_main_prints_mobile_access_urls(self):
        source = (ROOT / "start.py").read_text(encoding="utf-8")
        main_block = source[source.index('if __name__ == "__main__":') :]

        self.assertIn("print_access_urls()", main_block)

    def test_startup_launches_ngrok_with_configured_static_url(self):
        import start

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            pid_file = tmp_path / "ngrok.pid"
            log_file = tmp_path / "ngrok.log"

            process = mock.Mock()
            process.pid = 321

            with mock.patch.dict(
                start.os.environ,
                {
                    "MEETING_ASSISTANT_NGROK": "1",
                    "MEETING_ASSISTANT_NGROK_URL": "https://example.ngrok-free.app",
                },
                clear=False,
            ), \
                 mock.patch.object(start, "NGROK_PID_FILE", pid_file), \
                 mock.patch.object(start, "NGROK_LOG_FILE", log_file), \
                 mock.patch.object(start.shutil, "which", return_value="/usr/local/bin/ngrok"), \
                 mock.patch.object(start, "terminate_existing_ngrok"), \
                 mock.patch.object(start.subprocess, "Popen", return_value=process) as popen, \
                 mock.patch("builtins.print"):
                returned = start.start_ngrok_tunnel(8001, wait_for_status=False)
                pid_text = pid_file.read_text(encoding="utf-8")

        self.assertIs(returned, process)
        command = popen.call_args.args[0]
        self.assertEqual(command[:3], ["ngrok", "http", "8001"])
        self.assertIn("--url=https://example.ngrok-free.app", command)
        self.assertIn("--log=stdout", command)
        self.assertEqual(pid_text, "321")

    def test_startup_finds_winget_ngrok_when_not_on_path(self):
        import start

        with tempfile.TemporaryDirectory() as tmpdir:
            local_app_data = Path(tmpdir)
            ngrok_path = (
                local_app_data
                / "Microsoft"
                / "WinGet"
                / "Packages"
                / "Ngrok.Ngrok_Microsoft.Winget.Source_8wekyb3d8bbwe"
                / "ngrok.exe"
            )
            ngrok_path.parent.mkdir(parents=True)
            ngrok_path.write_bytes(b"fake exe")

            with mock.patch.dict(start.os.environ, {"LOCALAPPDATA": str(local_app_data)}, clear=False), \
                 mock.patch.object(start.shutil, "which", return_value=None):
                command = start._ngrok_command()

        self.assertEqual(command, str(ngrok_path))


class FreeOptimizationRegressionTests(unittest.TestCase):
    def test_segment_cache_is_reused_across_jobs_for_identical_audio(self):
        import backend.tasks as tasks

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"same-audio")
            context = tasks._segment_cache_context(audio_path, "gemini-test", 1, 10)
            tasks._save_segment_transcript_cache(
                root / "output",
                "job-one",
                0,
                context,
                "[00:00] **[發言者 A]**：共用快取內容。",
            )

            cached = tasks._load_segment_transcript_cache(
                root / "output", "job-two", 0, context
            )

        self.assertIn("共用快取內容", cached)

    def test_smart_segment_boundary_prefers_nearby_silence(self):
        from pydub import AudioSegment
        from pydub.generators import Sine
        from backend.tasks import _smart_segment_boundaries

        tone = Sine(440).to_audio_segment(duration=9000)
        audio = tone + AudioSegment.silent(duration=2000) + tone
        boundaries = _smart_segment_boundaries(audio, segment_ms=10000)

        self.assertEqual(boundaries[0], 0)
        self.assertEqual(boundaries[-1], len(audio))
        self.assertGreaterEqual(boundaries[1], 9500)
        self.assertLessEqual(boundaries[1], 10500)

    def test_audio_preflight_rejects_effectively_silent_recording(self):
        import backend.tasks as tasks
        from pydub import AudioSegment

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            audio_path = root / "silent.wav"
            export_handle = AudioSegment.silent(duration=31000).export(audio_path, format="wav")
            export_handle.close()

            with self.assertRaisesRegex(RuntimeError, "幾乎沒有可辨識聲音"):
                tasks._prepare_audio_for_transcription(audio_path, root / "temp", "silent-job")

    def test_summary_payload_repairs_ids_refs_and_timecodes_locally(self):
        from backend.tasks import _summary_response_to_markdown

        payload = {
            "discussion_summary": [
                {"id": "D9", "topic": "測試", "summary": "內容", "evidence_timecodes": ["04:40"]}
            ],
            "final_decisions": [
                {"id": "R8", "related_discussions": ["D9", "D1"], "decision": "通過", "status": "confirmed"}
            ],
            "action_items": [
                {"id": "A7", "related_discussions": ["D1"], "related_decisions": ["R8", "R1"], "task": "追蹤"}
            ],
        }
        transcript = (
            "[00:00] **[發言者 A]**：開始。\n"
            "[05:00] **[發言者 B]**：確認。"
        )

        markdown = _summary_response_to_markdown(json.dumps(payload, ensure_ascii=False), transcript)

        self.assertIn("### D1. 測試", markdown)
        self.assertIn("佐證時間：05:00", markdown)
        self.assertIn("| R1 | D1 |", markdown)
        self.assertIn("| A1 | D1 | R1 |", markdown)
        self.assertNotIn("D9", markdown)
        self.assertNotIn("R8", markdown)

    def test_summary_grounding_uses_actual_meeting_date_and_explicit_speaker_owner(self):
        from backend.tasks import _infer_meeting_date, _summary_response_to_markdown

        meeting_date = _infer_meeting_date(
            "現場錄製_2026-7-10 下午1-17-27",
            Path("099a6561_20260710_131729.webm"),
        )
        transcript = (
            "[00:00] **[發言者 A]**：熱處理參數暫定 140 到 145 度，後續再調整。\n"
            "[00:15] **[發言者 A]**：今天是 7 月 10 號，那我會再問一下品保，改成下禮拜一去做端相。"
        )
        payload = {
            "discussion_summary": [
                {"topic": "熱處理", "summary": "討論熱處理參數。", "evidence_timecodes": ["00:00"]}
            ],
            "final_decisions": [
                {
                    "related_discussions": ["D1"],
                    "decision": "熱處理參數暫定 140 到 145 度，後續再調整。",
                    "basis": "會議中提出暫定參數。",
                    "status": "confirmed",
                    "evidence_timecodes": ["00:00"],
                }
            ],
            "action_items": [
                {
                    "related_discussions": ["D1"],
                    "task": "向品保確認端相安排",
                    "owner": "品保",
                    "due": "2026/07/14",
                    "due_source": "下禮拜一",
                    "source_timecodes": ["00:15"],
                    "priority": "中",
                }
            ],
        }

        markdown = _summary_response_to_markdown(
            json.dumps(payload, ensure_ascii=False),
            transcript,
            meeting_date,
        )

        self.assertEqual(meeting_date.isoformat(), "2026-07-10")
        self.assertIn("| R1 | D1 |", markdown)
        self.assertIn("| pending |", markdown)
        self.assertIn("佐證：00:00", markdown)
        self.assertIn("| 發言者 A | 2026/07/13（原文：下禮拜一） |", markdown)

    def test_summary_prompt_requires_fact_ledger_and_preserves_relative_due_source(self):
        from backend.tasks import _build_summary_prompt

        prompt = _build_summary_prompt(
            "[00:00] **[發言者 A]**：下週一追蹤。",
            datetime(2026, 7, 10).date(),
        )

        self.assertIn("實際會議日期：2026/07/10（星期五）", prompt)
        self.assertIn("silently build a fact ledger", prompt)
        self.assertIn("due_source 必須保留逐字稿原句", prompt)
        self.assertIn("我會問品保", prompt)

    def test_meeting_quality_report_round_trips_through_database(self):
        import backend.database as database

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_path = root / "meeting.md"
            output_path.write_text("meeting", encoding="utf-8")
            quality_report = {"score": 95, "label": "良好", "segments": [{"index": 0}]}
            with mock.patch.object(database, "DB_PATH", root / "meetings.db"):
                database.init_db()
                meeting_id = database.save_meeting(
                    "品質測試",
                    "2026/07/12",
                    "meeting.webm",
                    str(output_path),
                    "摘要",
                    job_id="quality-job",
                    quality_report=quality_report,
                )
                record = database.get_meeting(meeting_id)

        self.assertEqual(record["job_id"], "quality-job")
        self.assertEqual(record["quality_score"], 95)
        self.assertEqual(record["quality_report"]["segments"][0]["index"], 0)

    def test_meeting_source_audio_endpoint_streams_retained_file(self):
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            audio_path = Path(tmpdir) / "source.mp3"
            audio_path.write_bytes(b"ID3\x04\x00\x00audio")
            record = {
                "id": 5,
                "title": "音檔證據",
                "source_audio": str(audio_path),
            }
            with mock.patch.object(main, "get_meeting", return_value=record):
                response = asgi_request(main.app, "GET", "/meetings/5/source-audio")
                download_response = asgi_request(main.app, "GET", "/meetings/5/source-audio?download=1")
                media_response = asgi_request(main.app, "GET", "/meetings/5/source-media")
                media_head_response = asgi_request(main.app, "HEAD", "/meetings/5/source-media")
                media_download_head_response = asgi_request(main.app, "HEAD", "/meetings/5/source-media?download=1")

        self.assertEqual(response.status_code, 200)
        self.assertIn("audio/mpeg", response.headers.get("content-type", ""))
        self.assertIn("inline", response.headers.get("content-disposition", ""))
        self.assertEqual(response.headers.get("accept-ranges"), "bytes")
        self.assertEqual(response.content, b"ID3\x04\x00\x00audio")
        self.assertEqual(download_response.status_code, 200)
        self.assertIn("attachment", download_response.headers.get("content-disposition", ""))
        self.assertEqual(media_response.status_code, 200)
        self.assertEqual(media_response.content, response.content)
        self.assertEqual(media_head_response.status_code, 200)
        self.assertIn("audio/mpeg", media_head_response.headers.get("content-type", ""))
        self.assertIn("inline", media_head_response.headers.get("content-disposition", ""))
        self.assertEqual(media_head_response.headers.get("accept-ranges"), "bytes")
        self.assertEqual(media_head_response.content, b"")
        self.assertEqual(media_download_head_response.status_code, 200)
        self.assertIn("attachment", media_download_head_response.headers.get("content-disposition", ""))

    def test_webm_source_audio_endpoint_uses_recording_profile_media_type(self):
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            video_path = Path(tmpdir) / "screen.webm"
            video_path.write_bytes(b"webm")
            record = {
                "id": 6,
                "title": "錄影證據",
                "source_audio": str(video_path),
                "quality_report": {"recording": {"profile": "video_balanced"}},
            }
            with mock.patch.object(main, "get_meeting", return_value=record):
                video_response = asgi_request(main.app, "GET", "/meetings/6/source-audio")
                video_head_response = asgi_request(main.app, "HEAD", "/meetings/6/source-media")

            record["quality_report"] = {"recording": {"profile": "audio_standard"}}
            with mock.patch.object(main, "get_meeting", return_value=record):
                audio_response = asgi_request(main.app, "GET", "/meetings/6/source-audio")

        self.assertEqual(video_response.status_code, 200)
        self.assertIn("video/webm", video_response.headers.get("content-type", ""))
        self.assertEqual(video_head_response.status_code, 200)
        self.assertIn("video/webm", video_head_response.headers.get("content-type", ""))
        self.assertEqual(video_head_response.content, b"")
        self.assertEqual(audio_response.status_code, 200)
        self.assertIn("audio/webm", audio_response.headers.get("content-type", ""))

    def test_source_media_type_uses_ffprobe_when_profile_is_unknown(self):
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            source_path = Path(tmpdir) / "recording.webm"
            source_path.write_bytes(b"webm")
            record = {"source_audio": str(source_path), "quality_report": {}}

            with mock.patch.object(main.shutil, "which", return_value="ffprobe"), \
                 mock.patch.object(
                     main.subprocess,
                     "run",
                     return_value=subprocess.CompletedProcess(
                         args=[],
                         returncode=0,
                         stdout='{"streams":[{"codec_type":"video"},{"codec_type":"audio"}]}',
                         stderr="",
                     ),
                 ):
                self.assertEqual(main._source_media_type(record), "video")

            main.FFPROBE_STREAM_CACHE.clear()
            with mock.patch.object(main.shutil, "which", return_value="ffprobe"), \
                 mock.patch.object(
                     main.subprocess,
                     "run",
                     return_value=subprocess.CompletedProcess(
                         args=[],
                         returncode=0,
                         stdout='{"streams":[{"codec_type":"audio"}]}',
                         stderr="",
                     ),
                 ):
                self.assertEqual(main._source_media_type(record), "audio")

    def test_meeting_detail_exposes_source_media_type(self):
        import backend.main as main

        with tempfile.TemporaryDirectory() as tmpdir:
            source_path = Path(tmpdir) / "screen.webm"
            source_path.write_bytes(b"video")
            record = {
                "id": 7,
                "title": "錄影會議",
                "date": "2026/07/12",
                "source_audio": str(source_path),
                "output_path": "meeting.md",
                "summary": "summary",
                "job_id": None,
                "quality_score": None,
                "quality_label": None,
                "created_at": datetime(2026, 7, 12, 9, 0, 0),
                "full_content": "## 一、討論摘要 (Discussion Summary)\nD1 測試",
                "quality_report": {
                    "recording": {
                        "profile": "video_balanced",
                        "source_audio_sha256": "abc123def4567890",
                    }
                },
            }
            with mock.patch.object(main, "get_meeting", return_value=record):
                response = asgi_request(main.app, "GET", "/meetings/7")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["source_media_type"], "video")
        self.assertEqual(payload["recording_profile"], "video_balanced")
        self.assertEqual(payload["source_media_size_bytes"], 5)
        self.assertEqual(payload["source_media_sha256"], "abc123def4567890")

    def test_meeting_rerun_api_can_force_only_one_segment(self):
        import backend.main as main

        record = {
            "id": 8,
            "title": "局部重跑測試",
            "quality_report": {"segments": [{"index": 0}, {"index": 1}]},
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            audio_path = Path(tmpdir) / "meeting.webm"
            audio_path.write_bytes(b"audio")
            meeting_path = Path(tmpdir) / "meeting.md"
            meeting_path.write_text(
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "### 【第 1 段｜00:00 – 10:00】\n[00:00] **[發言者 A]**：第一段。\n"
                "### 【第 2 段｜10:00 – 20:00】\n[10:00] **[發言者 B]**：第二段。",
                encoding="utf-8",
            )
            record["output_path"] = str(meeting_path)
            with mock.patch.object(main, "get_meeting", return_value=record), \
                 mock.patch.object(main, "_resolve_meeting_source_audio", return_value=audio_path), \
                 mock.patch.object(main, "enqueue_audio_job") as enqueue:
                response = asgi_request(
                    main.app,
                    "POST",
                    "/meetings/8/rerun",
                    json={"segments": [1]},
                )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(enqueue.call_args.kwargs["force_segment_indices"], [1])
        self.assertEqual(enqueue.call_args.kwargs["transcript_reuse_source_path"], meeting_path)
        self.assertIn("第 2 段", response.json()["message"])

    def test_meeting_rerun_api_can_rebuild_summary_without_forcing_transcription(self):
        import backend.main as main

        record = {
            "id": 8,
            "title": "摘要重整測試",
            "quality_report": {"segments": [{"index": 0}, {"index": 1}]},
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            audio_path = Path(tmpdir) / "meeting.webm"
            audio_path.write_bytes(b"audio")
            meeting_path = Path(tmpdir) / "meeting.md"
            meeting_path.write_text("## 📝 四、完整逐字稿 (Verbatim Transcript)\n[00:00] **[發言者 A]**：測試。", encoding="utf-8")
            record["output_path"] = str(meeting_path)
            with mock.patch.object(main, "get_meeting", return_value=record), \
                 mock.patch.object(main, "_resolve_meeting_source_audio", return_value=audio_path), \
                 mock.patch.object(main, "enqueue_audio_job") as enqueue:
                response = asgi_request(
                    main.app,
                    "POST",
                    "/meetings/8/rerun",
                    json={"summary_only": True},
                )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(enqueue.call_args.kwargs["force_segment_indices"], [])
        self.assertEqual(enqueue.call_args.kwargs["source"], "meeting_summary_rerun")
        self.assertEqual(enqueue.call_args.kwargs["summary_source_path"], meeting_path)
        self.assertIn("摘要重整", response.json()["message"])

    def test_high_quality_summary_api_enables_second_model_verification(self):
        import backend.main as main

        record = {"id": 9, "title": "高手質測試", "quality_report": None}
        with tempfile.TemporaryDirectory() as tmpdir:
            audio_path = Path(tmpdir) / "meeting.webm"
            audio_path.write_bytes(b"audio")
            meeting_path = Path(tmpdir) / "meeting.md"
            meeting_path.write_text(
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n[00:00] **[發言者 A]**：測試。",
                encoding="utf-8",
            )
            record["output_path"] = str(meeting_path)
            with mock.patch.object(main, "get_meeting", return_value=record), \
                 mock.patch.object(main, "_resolve_meeting_source_audio", return_value=audio_path), \
                 mock.patch.object(main, "enqueue_audio_job") as enqueue:
                response = asgi_request(
                    main.app,
                    "POST",
                    "/meetings/9/rerun",
                    json={"summary_only": True, "high_quality": True},
                )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(enqueue.call_args.kwargs["high_quality_summary"])
        self.assertEqual(enqueue.call_args.kwargs["source"], "meeting_summary_high_quality")
        self.assertIn("第二模型", response.json()["message"])

    def test_high_quality_summary_calls_verifier_and_uses_verified_result(self):
        from backend import tasks

        initial = {
            "discussion_summary": [{"topic": "初稿", "summary": "第一階段", "evidence_timecodes": ["00:00"]}],
            "final_decisions": [],
            "action_items": [],
        }
        verified = {
            "discussion_summary": [{"topic": "查核後", "summary": "第二模型已查核", "evidence_timecodes": ["00:00"]}],
            "final_decisions": [],
            "action_items": [],
        }

        class FakeModels:
            def __init__(self):
                self.calls = []

            def generate_content(self, **kwargs):
                self.calls.append(kwargs)
                payload = initial if len(self.calls) == 1 else verified
                return type("Response", (), {"text": json.dumps(payload, ensure_ascii=False)})()

        fake_client = type("Client", (), {"models": FakeModels()})()
        transcript = "[00:00] **[發言者 A]**：第二模型已查核。"

        with mock.patch.object(tasks, "update_job_status"):
            content, used_model = tasks._generate_meeting_content_from_transcript(
                fake_client,
                full_transcript=transcript,
                job_id="high-quality-job",
                summary_primary_model="gemma-primary",
                summary_secondary_model="gemini-verifier",
                summary_verifier_model="gemini-3.5-flash",
                meeting_date=datetime(2026, 7, 10).date(),
                high_quality=True,
            )

        self.assertEqual([call["model"] for call in fake_client.models.calls], ["gemma-primary", "gemini-3.5-flash"])
        self.assertIn("第二模型已查核", content)
        self.assertNotIn("第一階段", content)
        self.assertEqual(used_model, "gemma-primary+verified:gemini-3.5-flash")

    def test_partial_rerun_reuses_unselected_segments_from_existing_record(self):
        from backend import tasks

        summary_payload = {
            "discussion_summary": [{"topic": "局部重跑", "summary": "完成", "evidence_timecodes": ["00:00"]}],
            "final_decisions": [],
            "action_items": [],
        }

        class FakeModels:
            def generate_content(self, **kwargs):
                return type("Response", (), {"text": json.dumps(summary_payload, ensure_ascii=False)})()

        class FakeClient:
            def __init__(self, *args, **kwargs):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            audio_path = root / "meeting.webm"
            audio_path.write_bytes(b"audio")
            slices = []
            for index in range(3):
                segment_path = root / f"_seg_meeting_{index:03d}.mp3"
                segment_path.write_bytes(b"segment")
                slices.append(tasks.AudioSlice(segment_path, index * 600, (index + 1) * 600))
            source_path = root / "existing.md"
            source_path.write_text(
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "### 【第 1 段｜00:00 – 10:00】\n[00:00] **[發言者 A]**：保留第一段。\n"
                "### 【第 2 段｜10:00 – 20:00】\n[10:00] **[發言者 B]**：舊第二段。\n"
                "### 【第 3 段｜20:00 – 30:00】\n[20:00] **[發言者 C]**：保留第三段。",
                encoding="utf-8",
            )
            output_dir = root / "output"

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", side_effect=FakeClient), \
                 mock.patch.object(tasks, "_prepare_audio_for_transcription", return_value=(audio_path, {})), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=slices), \
                 mock.patch.object(tasks, "_load_segment_transcript_cache", return_value=None), \
                 mock.patch.object(tasks, "_transcribe_segment_with_recovery", return_value="[10:00] **[發言者 B]**：新第二段。") as transcribe, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="partial-rerun-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    force_segment_indices=[1],
                    transcript_reuse_source_path=source_path,
                )

            self.assertIsNotNone(output_path)
            self.assertEqual(transcribe.call_count, 1)
            output_text = output_path.read_text(encoding="utf-8")
            self.assertIn("保留第一段", output_text)
            self.assertIn("新第二段", output_text)
            self.assertIn("保留第三段", output_text)
            self.assertNotIn("舊第二段", output_text)

    def test_old_meeting_detail_recovers_segment_controls_from_transcript(self):
        import backend.main as main

        record = {
            "id": 12,
            "title": "舊紀錄",
            "date": "2026/07/08",
            "source_audio": "old.webm",
            "output_path": "old.md",
            "summary": "摘要",
            "job_id": None,
            "quality_score": None,
            "quality_label": None,
            "created_at": "2026-07-08 10:00:00",
            "quality_report": None,
            "full_content": (
                "## 一、討論摘要 (Discussion Summary)\n摘要\n"
                "## 二、最終決議 (Final Decisions)\n決議\n"
                "## 三、待辦事項 (Action Items)\n| # | 任務描述 | 負責人 | 期限 | 優先級 |\n|---|---|---|---|---|\n| A1 | 無 | 無 | 無 | 中 |\n"
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "### 【第 1 段｜00:00 – 10:00】\n[00:00] **[發言者 A]**：第一段。\n"
                "### 【第 2 段｜10:00 – 20:00】\n[10:00] **[發言者 B]**：第二段。"
            ),
        }
        with mock.patch.object(main, "get_meeting", return_value=record):
            response = asgi_request(main.app, "GET", "/meetings/12")

        self.assertEqual(response.status_code, 200)
        report = response.json()["quality_report"]
        self.assertEqual(len(report["segments"]), 2)
        self.assertEqual(report["segments"][1]["start_seconds"], 600)
        self.assertIn("已重建分段", report["label"])

    def test_meeting_detail_flags_unsafe_legacy_transcript_in_quality_report(self):
        import backend.main as main

        repeated_turns = "".join(
            f"[00:{index:02d}] **[發言者 A]**：這一句不應該連續重複。\n"
            for index in range(4)
        )
        record = {
            "id": 13,
            "title": "異常舊紀錄",
            "date": "2026/07/08",
            "source_audio": "bad.webm",
            "output_path": "bad.md",
            "summary": "摘要",
            "job_id": None,
            "quality_score": None,
            "quality_label": None,
            "created_at": "2026-07-08 10:00:00",
            "quality_report": None,
            "full_content": (
                "## 一、討論摘要 (Discussion Summary)\n摘要\n"
                "## 二、最終決議 (Final Decisions)\n決議\n"
                "## 三、待辦事項 (Action Items)\n| # | 任務描述 | 負責人 | 期限 | 優先級 |\n|---|---|---|---|---|\n| A1 | 無 | 無 | 無 | 中 |\n"
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "### 【第 1 段｜00:00 – 10:00】\n"
                "*(註：為節省篇幅，已省略逐字稿中重複內容)*\n"
                f"{repeated_turns}"
            ),
        }
        with mock.patch.object(main, "get_meeting", return_value=record):
            response = asgi_request(main.app, "GET", "/meetings/13")

        self.assertEqual(response.status_code, 200)
        report = response.json()["quality_report"]
        warnings = "\n".join(report["warnings"])
        self.assertIn("逐字稿品質警示", warnings)
        self.assertIn("省略", warnings)
        self.assertIn("重複", warnings)
        self.assertEqual(report["label"], "舊紀錄，已重建分段")

    def test_meeting_detail_flags_unlinked_legacy_summary_in_quality_report(self):
        import backend.main as main

        record = {
            "id": 14,
            "title": "未串聯舊紀錄",
            "date": "2026/07/08",
            "source_audio": "unlinked.webm",
            "output_path": "unlinked.md",
            "summary": "摘要",
            "job_id": None,
            "quality_score": None,
            "quality_label": None,
            "created_at": "2026-07-08 10:00:00",
            "quality_report": None,
            "full_content": (
                "## 一、討論摘要 (Discussion Summary)\n"
                "這是一段沒有 D 編號的摘要。\n"
                "## 二、最終決議 (Final Decisions)\n"
                "| # | 關聯討論 | 決議 | 依據 | 狀態 |\n"
                "|---|---|---|---|---|\n"
                "| R2 | D9 | 先執行改善 | 00:10 | confirmed |\n"
                "## 三、待辦事項 (Action Items)\n"
                "| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n"
                "|---|---|---|---|---|---|---|\n"
                "| - | D9 | R3 | 整理追蹤表 | 發言者 A | 未提及 | 中 |\n"
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "### 【第 1 段｜00:00 – 10:00】\n[00:00] **[發言者 A]**：逐字稿正常。\n"
            ),
        }
        with mock.patch.object(main, "get_meeting", return_value=record):
            response = asgi_request(main.app, "GET", "/meetings/14")

        self.assertEqual(response.status_code, 200)
        report = response.json()["quality_report"]
        warnings = "\n".join(report["warnings"])
        self.assertIn("摘要品質警示", warnings)
        self.assertIn("D 編號", warnings)
        self.assertIn("A 編號", warnings)
        self.assertIn("D9", warnings)
        self.assertIn("R3", warnings)

    def test_manual_summary_edit_preserves_transcript_and_ai_original(self):
        import backend.database as database
        import backend.main as main
        from backend.tasks import _extract_transcript_section_body

        original = (
            "---\ntitle: 測試\n---\n\n"
            "## 一、討論摘要 (Discussion Summary)\n\n### D1. 原始摘要\n- 摘要：AI 原稿\n\n"
            "## 二、最終決議 (Final Decisions)\n\n| # | 關聯討論 | 決議 | 依據 | 狀態 |\n|---|---|---|---|---|\n| R1 | D1 | 未提及 | 未提及 | pending |\n\n"
            "## 三、待辦事項 (Action Items)\n\n| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n|---|---|---|---|---|---|---|\n| A1 | D1 | R1 | 未提及 | 未提及 | 未提及 | 中 |\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n[00:00] **[發言者 A]**：逐字稿不得改變。\n\n"
            "## 📎 五、補充資料與佐證 (Supplementary Evidence)\n\n### 資料：spec.pdf\n- 系統判斷：應保留。\n"
        )
        edited_summary = (
            "## 一、討論摘要 (Discussion Summary)\n\n### D1. 人工修訂\n- 摘要：修訂後摘要\n\n"
            "## 二、最終決議 (Final Decisions)\n\n| # | 關聯討論 | 決議 | 依據 | 狀態 |\n|---|---|---|---|---|\n| R1 | D1 | 維持測試 | 00:00 | confirmed |\n\n"
            "## 三、待辦事項 (Action Items)\n\n| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n|---|---|---|---|---|---|---|\n| A1 | D1 | R1 | 整理紀錄 | 發言者 A | 未提及 | 中 |"
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_path = root / "meeting.md"
            output_path.write_text(original, encoding="utf-8")
            with mock.patch.object(database, "DB_PATH", root / "meetings.db"):
                database.init_db()
                meeting_id = database.save_meeting(
                    "編輯測試", "2026/07/12", "meeting.webm", str(output_path), "AI 原稿"
                )
                response = asgi_request(
                    main.app,
                    "PUT",
                    f"/meetings/{meeting_id}/summary",
                    json={"summary_markdown": edited_summary},
                )
                revisions = database.list_meeting_revisions(meeting_id)

            saved = output_path.read_text(encoding="utf-8")

        self.assertEqual(response.status_code, 200)
        self.assertIn("修訂後摘要", saved)
        self.assertEqual(
            _extract_transcript_section_body(saved),
            _extract_transcript_section_body(original),
        )
        self.assertIn("## 📎 五、補充資料與佐證", saved)
        self.assertIn("spec.pdf", saved)
        self.assertEqual(len(revisions), 1)
        self.assertEqual(revisions[0]["content"], original)

    def test_manual_transcript_edit_preserves_summary_and_revision_history(self):
        import backend.database as database
        import backend.main as main
        from backend.tasks import _extract_transcript_section_body

        original = (
            "## 📋 一、討論摘要 (Discussion Summary)\n\n### D1. 原始摘要\n- 摘要：摘要不變\n\n"
            "## ✅ 二、最終決議 (Final Decisions)\n\n| # | 關聯討論 | 決議 | 依據 | 狀態 |\n|---|---|---|---|---|\n| R1 | D1 | 保留 | 00:00 | confirmed |\n\n"
            "## 📌 三、待辦事項 (Action Items)\n\n| # | 關聯討論 | 關聯決議 | 任務描述 | 負責人 | 期限 | 優先級 |\n|---|---|---|---|---|---|---|\n| A1 | D1 | R1 | 保留摘要 | 發言者 A | 未提及 | 中 |\n\n"
            "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
            "### 【第 1 段｜00:00 – 10:00】\n[00:00] **[發言者 A]**：舊逐字稿。\n\n"
            "## 📎 五、補充資料與佐證 (Supplementary Evidence)\n\n### 資料：photo.png\n- 系統判斷：應保留。\n"
        )
        edited_transcript = (
            "### 【第 1 段｜00:00 – 10:00】\n"
            "[00:00] **[發言者 A]**：修正後逐字稿，保留時間戳與發言者。\n"
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            output_path = root / "meeting.md"
            output_path.write_text(original, encoding="utf-8")
            with mock.patch.object(database, "DB_PATH", root / "meetings.db"):
                database.init_db()
                meeting_id = database.save_meeting(
                    "逐字稿編輯測試", "2026/07/12", "meeting.webm", str(output_path), "摘要不變"
                )
                response = asgi_request(
                    main.app,
                    "PUT",
                    f"/meetings/{meeting_id}/transcript",
                    json={"transcript_markdown": edited_transcript},
                )
                revisions = database.list_meeting_revisions(meeting_id)

            saved = output_path.read_text(encoding="utf-8")

        self.assertEqual(response.status_code, 200)
        self.assertIn("摘要不變", saved)
        self.assertIn("修正後逐字稿", _extract_transcript_section_body(saved))
        self.assertNotIn("舊逐字稿", _extract_transcript_section_body(saved))
        self.assertIn("## 📎 五、補充資料與佐證", saved)
        self.assertIn("photo.png", saved)
        self.assertEqual(len(revisions), 1)
        self.assertEqual(revisions[0]["source"], "manual_transcript_edit")
        self.assertEqual(revisions[0]["content"], original)

    def test_summary_only_processing_never_calls_transcription_model(self):
        from backend import tasks

        summary_payload = {
            "discussion_summary": [
                {"topic": "測試", "summary": "已沿用逐字稿。", "evidence_timecodes": ["00:00"]}
            ],
            "final_decisions": [],
            "action_items": [],
        }

        class FakeModels:
            def generate_content(self, **kwargs):
                return type(
                    "Response",
                    (),
                    {"text": json.dumps(summary_payload, ensure_ascii=False)},
                )()

        class FakeClient:
            def __init__(self, *args, **kwargs):
                self.models = FakeModels()

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            audio_path = root / "099a6561_20260710_131729.webm"
            audio_path.write_bytes(b"audio")
            source_path = root / "existing.md"
            source_path.write_text(
                "## 📝 四、完整逐字稿 (Verbatim Transcript)\n"
                "[00:00] **[發言者 A]**：保留這份完整逐字稿。\n",
                encoding="utf-8",
            )
            output_dir = root / "output"

            with mock.patch.dict(os.environ, {"GEMINI_API_KEY": "test-key"}), \
                 mock.patch.object(tasks.genai, "Client", side_effect=FakeClient), \
                 mock.patch.object(tasks, "_prepare_audio_for_transcription", return_value=(audio_path, {})), \
                 mock.patch.object(tasks, "_split_audio_to_segments", return_value=[audio_path]), \
                 mock.patch.object(tasks, "_transcribe_segment") as transcribe, \
                 mock.patch.object(tasks, "is_job_cancel_requested", return_value=False), \
                 mock.patch.object(tasks, "update_job_status"), \
                 mock.patch.object(tasks, "save_meeting"):
                output_path = tasks.process_audio_task(
                    job_id="summary-only-job",
                    audio_path=audio_path,
                    output_dir=output_dir,
                    meeting_title="現場錄製_2026-7-10 下午1-17-27",
                    summary_source_path=source_path,
                )

            self.assertIsNotNone(output_path)
            transcribe.assert_not_called()
            output_text = output_path.read_text(encoding="utf-8")
            self.assertIn("已沿用逐字稿", output_text)
            self.assertIn("保留這份完整逐字稿", output_text)

    def test_web_ui_shows_quality_report_and_segment_rerun_controls(self):
        html = (ROOT / "static" / "index.html").read_text(encoding="utf-8")

        self.assertIn('id="search-status"', html)
        self.assertIn("AbortController", html)
        self.assertIn("找到 ${fetchedRecords.length} 筆${reviewOnly ? '需複核且相關' : '相關'}會議記錄", html)
        self.assertIn("function renderQualityReport", html)
        self.assertIn("quality_report", html)
        self.assertIn("quality-warning", html)
        self.assertIn("report.warnings", html)
        self.assertIn("function renderQualityActions", html)
        self.assertIn("quality-rerun-summary-button", html)
        self.assertIn("quality-rerun-summary-high-quality-button", html)
        self.assertIn("quality-rerun-full-button", html)
        self.assertIn("function renderCardQuality", html)
        self.assertIn("function isNeedsReviewRecord", html)
        self.assertIn("function sourceMediaIcon", html)
        self.assertIn("${sourceMediaIcon(r)} ${escapeHtml(r.source_audio)}", html)
        self.assertIn('id="needs-review-filter"', html)
        self.assertIn("review-filter", html)
        self.assertIn("quality_warning_count", html)
        self.assertIn("quality_warning_preview", html)
        self.assertIn("card-quality-chip", html)
        self.assertIn("card-review-reason", html)
        self.assertIn("原因：${escapeHtml(warningPreview)}", html)
        self.assertIn("需複核 ${warningCount}", html)
        self.assertIn("const reviewParam = reviewOnly ? '&needs_review=true' : ''", html)
        self.assertIn("/meetings/search?q=${encodeURIComponent(query)}&limit=100${reviewParam}", html)
        self.assertIn("/meetings?limit=50${reviewParam}", html)
        self.assertIn("startsWith('摘要品質警示')", html)
        self.assertIn("startsWith('逐字稿品質警示')", html)
        self.assertIn("triggerButtonId", html)
        self.assertIn("rerun-segment-${index}", html)
        self.assertIn("JSON.stringify({ segments: [segmentIndex] })", html)
        self.assertIn('id="rerun-summary-button"', html)
        self.assertIn("summary_only: true, high_quality: highQuality", html)
        self.assertIn('id="rerun-summary-high-quality-button"', html)
        self.assertIn('id="edit-summary-button"', html)
        self.assertIn('id="edit-transcript-button"', html)
        self.assertIn('id="transcript-editor-modal"', html)
        self.assertIn('id="source-media-player"', html)
        self.assertIn("/source-media", html)
        self.assertNotIn("source-audio`", html)
        self.assertIn("source-media-actions", html)
        self.assertIn(".source-media-actions", html)
        self.assertIn(".source-media-player-status", html)
        self.assertIn(".source-media-player-status.error", html)
        self.assertIn("flex-wrap: wrap;", html)
        self.assertIn("justify-content: flex-end;", html)
        self.assertIn("justify-content: flex-start;", html)
        self.assertIn("source-media-facts", html)
        self.assertIn("source-media-fact", html)
        self.assertIn("function sourceMediaKindLabel", html)
        self.assertIn("function recordingProfileLabel", html)
        self.assertIn("return recordingProfileBaseLabel(profile);", html)
        self.assertNotIn("video_balanced: '錄影模式'", html)
        self.assertIn("function sourceHashPreview", html)
        self.assertIn("function renderSourceMediaFacts", html)
        self.assertIn("SHA256 ${hash}", html)
        self.assertIn("meeting?.recording_profile", html)
        self.assertIn("meeting?.source_media_size_bytes", html)
        self.assertIn("meeting?.source_media_sha256", html)
        self.assertIn("recording.source_audio_size_bytes", html)
        self.assertIn("recording.source_audio_sha256", html)
        self.assertIn("download=1", html)
        self.assertIn("↗ 開啟", html)
        self.assertIn("⬇ 下載", html)
        self.assertIn("function isVideoSource", html)
        self.assertIn("function sourceMediaVideoPreviewCapable", html)
        self.assertIn("function sourceMediaPlayerMode", html)
        self.assertIn("function switchSourceMediaPlayer", html)
        self.assertIn("function primeSourceVideoPreview", html)
        self.assertIn("function setSourceMediaPlayerStatus", html)
        self.assertIn("function sourceMediaHasAudioFallbackAction", html)
        self.assertIn("function updateSourceMediaPlayerStatus", html)
        self.assertIn("function bindSourceMediaPlayerStatus", html)
        self.assertIn("function enhanceSourceMediaPlayer", html)
        self.assertIn("video.dataset.previewPrimed", html)
        self.assertIn("loadedmetadata", html)
        self.assertIn("player.error", html)
        self.assertIn("!player.videoWidth || !player.videoHeight", html)
        self.assertIn("沒有偵測到可顯示的畫面", html)
        self.assertIn("sourceMediaHasAudioFallbackAction(statusEl)", html)
        self.assertIn("button[aria-label=\"切換為音訊預覽\"], button[data-source-mode=\"audio\"]", html)
        self.assertIn("可先按「音訊預覽」確認聲音", html)
        self.assertIn("請使用「新分頁」或「下載」確認原始檔。", html)
        self.assertIn("無法在瀏覽器中預覽", html)
        self.assertIn("player.dataset.statusBound", html)
        self.assertIn("querySelector('video.source-storage-media-player, audio')", html)
        self.assertIn("preferredPreviewSecond", html)
        self.assertIn("Math.max(0, video.duration - 0.05)", html)
        self.assertIn("video.currentTime = previewSecond;", html)
        self.assertIn("enhanceSourceMediaPlayer();", html)
        self.assertIn("source_media_type", html)
        self.assertIn("if (sourceType === 'audio') return false;", html)
        self.assertIn("meeting?.recording_profile || meeting?.quality_report?.recording?.profile", html)
        self.assertIn("new Set(['.webm', '.mp4', '.mov', '.mkv', '.avi', '.mpeg', '.mpg', '.wmv'])", html)
        self.assertIn("new Set(['.webm', '.mp4'])", html)
        self.assertIn('id="source-media-evidence"', html)
        self.assertIn('id="source-media-player-status"', html)
        self.assertIn('id="source-media-player-status" class="source-media-player-status" role="status" aria-live="polite"', html)
        self.assertIn("🎥 影片預覽", html)
        self.assertIn("🔊 音訊預覽", html)
        self.assertIn("<video id=\"source-media-player\"", html)
        self.assertIn("playsinline", html)
        self.assertIn("display: grid;", html)
        self.assertIn("grid-template-columns: minmax(0, 1fr);", html)
        self.assertIn("function enhanceTranscriptTimecodes", html)
        self.assertIn("function seekSourceAudio", html)
        self.assertIn("/transcript", html)
        self.assertIn("manual_transcript_edit", html)
        self.assertIn('id="revision-history-button"', html)
        self.assertIn("function saveSummaryEdit", html)
        self.assertIn("function saveTranscriptEdit", html)
        self.assertIn('id="rec-quality-profile"', html)
        self.assertIn("audioBitsPerSecond", html)
        self.assertIn("videoBitsPerSecond", html)
        self.assertIn("recording_profile", html)


if __name__ == "__main__":
    unittest.main()
